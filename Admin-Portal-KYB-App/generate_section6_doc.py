"""
Generate Section 6 specification document as a .docx file.
Run: python3 generate_section6_doc.py
Output: Admin-Portal-KYB-App/Section6_Portfolio_Anomaly_Scanner.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ──────────────────────────────────────────────────────────
C_BLUE       = RGBColor(0x3B, 0x82, 0xF6)   # primary blue
C_PURPLE     = RGBColor(0x8B, 0x5C, 0xF6)   # purple accent
C_GREEN      = RGBColor(0x22, 0xC5, 0x5E)   # green / clean
C_AMBER      = RGBColor(0xF5, 0x9E, 0x0B)   # amber / medium
C_ORANGE     = RGBColor(0xF9, 0x73, 0x16)   # orange / high
C_RED        = RGBColor(0xEF, 0x44, 0x44)   # red / critical
C_SLATE      = RGBColor(0x1E, 0x29, 0x3B)   # dark slate
C_LIGHT      = RGBColor(0x94, 0xA3, 0xB8)   # light text
C_BLACK      = RGBColor(0x0F, 0x17, 0x2A)   # near-black

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get(f"{edge}_val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get(f"{edge}_sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get(f"{edge}_color", "CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_body(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
    return p


def add_source_box(doc, sources: list):
    """Add a shaded source-reference box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run("📂 Source files: " + " · ".join(sources))
    run.font.size = Pt(8.5)
    run.italic    = True
    run.font.color.rgb = C_PURPLE


def add_severity_badge(doc, severity: str, name: str, description: str):
    colors = {
        "CRITICAL": ("EF4444", "🔴"),
        "HIGH":     ("F97316", "🟠"),
        "MEDIUM":   ("F59E0B", "🟡"),
        "LOW":      ("22C55E", "🟢"),
        "NOTICE":   ("3B82F6", "🔵"),
    }
    hex_c, icon = colors.get(severity, ("94A3B8", "⚪"))
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # badge cell
    c0 = tbl.rows[0].cells[0]
    c0.width = Inches(1.1)
    set_cell_bg(c0, hex_c)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f"{icon} {severity}")
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # description cell
    c1 = tbl.rows[0].cells[1]
    set_cell_bg(c1, "1E293B")
    p1 = c1.paragraphs[0]
    r1a = p1.add_run(f"{name}  ")
    r1a.bold = True
    r1a.font.size = Pt(9.5)
    r1a.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1b = p1.add_run(description)
    r1b.font.size = Pt(9)
    r1b.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    doc.add_paragraph()   # spacer


def add_matrix_table(doc, headers, rows_data, col_colors=None):
    """Generic coloured matrix table."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, "0F172A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
    # Data rows
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            bg = "1E293B" if i % 2 == 0 else "0F172A"
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    doc.add_paragraph()


def add_divider(doc):
    p = doc.add_paragraph("─" * 90)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.color.rgb = C_LIGHT
        run.font.size = Pt(7)


# ════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ═══════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "🔬 Section 6 — Portfolio Anomaly, Contradiction &\nCross-Reference Scanner", level=1, color=C_BLUE)
add_heading(doc, "KYB Intelligence Hub — Home Tab Enhancement Specification", level=2, color=C_LIGHT)

add_body(doc,
    f"Document generated: {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
    italic=True, color=C_LIGHT, size=9)

doc.add_paragraph()
add_body(doc,
    "PURPOSE",
    bold=True, color=C_PURPLE, size=11)
add_body(doc,
    "This document specifies every analysis, cross-tabulation, anomaly rule, data source, "
    "and UI element that will be built in Section 6 of the Home tab. Section 6 surfaces the "
    "'weird things' — businesses where two or more KYB signals contradict each other, where "
    "cross-section relationships reveal structural gaps, and where the portfolio's collective "
    "pattern exposes systemic issues that single-business views miss.",
    size=10)

add_body(doc,
    "KEY DESIGN RULES",
    bold=True, color=C_PURPLE, size=11)
add_bullet(doc, "Zero new Redshift queries at render time — all signals from stats_df and funnel_df already in memory.")
add_bullet(doc, "Every rule maps to check_agent_v2.DETERMINISTIC_CHECKS, index.ts, verification_results.sql, or aiNaicsEnrichment.ts.")
add_bullet(doc, "No LLM calls. No guessing. No hallucination. Only deterministic cross-field logic.")
add_bullet(doc, "Every finding shows business ID drilldowns with the exact contradicting signal columns.")
add_bullet(doc, "Every finding cites the exact source file and line number that defines why the combination is anomalous.")

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════
# STRUCTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "Structure Overview", level=1, color=C_BLUE)
add_body(doc,
    "Section 6 is organised into 9 sub-sections. All are computed from data already loaded "
    "for Sections 1–5 — no additional database round-trips.", size=10)

structure = [
    ("6.0", "Portfolio Health Scorecard",              "Summary banner — 5 KPI cards giving an instant portfolio health read before any analysis."),
    ("6.1", "Registry × TIN Cross-Tabulation",         "S1 × S2 — the fundamental KYB matrix. 3×3 heatmap + 5 named diagonal segments."),
    ("6.2", "Registration Depth Cascade",              "S1 internal — Domestic vs Foreign, tax-haven concentration, state match/mismatch analysis."),
    ("6.3", "Data Integrity Contradictions",           "5 signal combinations that should be logically impossible — bug indicators."),
    ("6.4", "Identity Completeness Triangle",          "S1 × S2 × S3 — TIN × IDV × is_sole_prop three-way relationship."),
    ("6.5", "Registry × Risk Signal Matrix",           "S1 × S4 — entity status vs watchlist/BK/liens. 6 high-severity named combinations."),
    ("6.6", "NAICS Classification Completeness",       "S3 × S1 × Firmographic — the three classification gap types (G1/G2/G3)."),
    ("6.7", "Score vs KYB Purity Matrix",              "S5 × S1+S2+S3+S4 — the underwriting oversight cross-reference. 4 flagged override cells."),
    ("6.8", "Anomaly Accumulation Triage",             "Synthesis — per-business anomaly count, Top 10, and co-occurrence heatmap."),
]

add_matrix_table(doc,
    ["Sub-section", "Name", "What it answers"],
    structure)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.0  PORTFOLIO HEALTH SCORECARD
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.0  Portfolio Health Scorecard", level=1, color=C_BLUE)
add_body(doc,
    "The opening banner of Section 6. Shown before any drill-down analysis. Gives the "
    "underwriting manager a single-line read of how healthy the portfolio's KYB data is "
    "from a structural consistency standpoint.", size=10)

add_heading(doc, "5 KPI Cards", level=2, color=C_PURPLE)
kpis = [
    ("✅ 0 Anomalies",    "Businesses with zero cross-field contradictions.",                          "Green",    "#22C55E"),
    ("⚠️ 1 Anomaly",     "Isolated issue. Usually resolvable with one action.",                       "Amber",    "#F59E0B"),
    ("🔴 2–3 Anomalies", "Compounded risk. Needs investigation before any decision.",                 "Orange",   "#F97316"),
    ("🔴 4+ Anomalies",  "Systematic KYB failure. Full manual review required.",                      "Red",      "#EF4444"),
    ("⚡ CRITICAL Combos","Count of the most dangerous cross-signal pairs across the portfolio.",      "Purple",   "#8B5CF6"),
]
add_matrix_table(doc,
    ["KPI Card", "What it counts", "Colour", "Hex"],
    kpis)

add_body(doc,
    "Below the 5 cards: a horizontal severity distribution bar (same visual language as the "
    "score gauge in Section 5) showing what % of the portfolio falls in each anomaly band. "
    "This acts as the navigator — the bands tell you which sub-sections to focus on.", size=10)

add_source_box(doc, ["stats_df (all signals)", "funnel_df (SOS/TIN funnel)", "_home_ws_clean (score decisions)"])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.1  REGISTRY × TIN
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.1  Registry × TIN Cross-Tabulation  (S1 × S2)", level=1, color=C_BLUE)
add_body(doc,
    "The most fundamental KYB cross-reference. The SOS and TIN verification paths are "
    "completely independent in the codebase — SOS from Middesk/OpenCorporates (index.ts:1371), "
    "TIN from Middesk IRS direct lookup (index.ts:429). They should correlate: a legitimately "
    "operating entity should pass both. When they diverge, the divergence is diagnostic.", size=10)

add_heading(doc, "The 3×3 Heatmap Table", level=2, color=C_PURPLE)
add_body(doc,
    "Rows = Registry Status.  Columns = TIN Status.  Each cell = business count + % + severity colour.", size=10)
add_matrix_table(doc,
    ["Registry Status \\ TIN Status", "TIN Pass", "TIN Fail", "TIN Not Checked"],
    [
        ["Registry Active (sos_active=true)",   "✅ Baseline clean",      "⚠️ DBA/name issue",       "❓ EIN check missing"],
        ["Registry Inactive (sos_active=false)", "🟠 Dissolved + valid EIN","🔴 Dissolved + mismatch", "🔴 Dissolved, EIN unknown"],
        ["No Registry (sos_match_boolean=null)", "🟠 EIN ok, entity unverifiable","🔴 Double failure","🔴 Completely unverified"],
    ])

add_heading(doc, "5 Named Diagonal Segments — each with business ID drilldown", level=2, color=C_PURPLE)

segs = [
    ("1", "Registry Active + TIN Pass",
     "CLEAN",
     "The portfolio baseline. Both the entity's SOS filing (Middesk/OC) and the IRS EIN check (Middesk) confirmed. "
     "These businesses have the strongest dual-verification signal. Count and % shown.",
     "sos_active=true AND tin_match_boolean=true",
     ""),

    ("2", "Registry Active + TIN Fail",
     "MEDIUM",
     "Most common anomaly. Entity IS legally registered (SOS active) but the EIN-to-name IRS check failed. "
     "Root cause: applicant submitted a DBA/trade name instead of the legal name on the EIN certificate "
     "(IRS Form CP-575 or 147C). Not necessarily fraud. Action: request IRS EIN confirmation letter.",
     "sos_active=true AND tin_match_status=failure",
     "check_agent_v2: sos_active_tin_failed (MEDIUM)  ·  index.ts:429 tin_match, index.ts:482 tin_match_boolean"),

    ("3", "Registry Found + TIN Not Checked",
     "MEDIUM",
     "Entity verified to exist (SOS found) but no IRS confirmation at all. Sub-split: "
     "(a) tin_submitted non-null → EIN given, Middesk TIN task never completed. "
     "(b) tin_submitted null → no EIN given at onboarding. "
     "Action differs: (a) investigate Middesk task status; (b) request EIN from business.",
     "sos_match_boolean=true AND tin_match_boolean IS NULL",
     "index.ts:482 — tin_match_boolean only written when tin_match has a value"),

    ("4", "Registry Inactive + TIN Pass",
     "HIGH",
     "EIN is valid (IRS confirmed) but the entity is NOT in good standing per SOS — dissolved, "
     "administratively revoked, or voluntarily wound down. Critical nuance: an EIN is never cancelled "
     "when a company dissolves. This combination means the entity may be operating under a dead legal entity.",
     "sos_active=false AND tin_match_boolean=true",
     "check_agent_v2: sos_inactive_tin_ok (HIGH)  ·  index.ts:1434 sos_active derivation"),

    ("5", "No Registry + TIN Pass",
     "HIGH",
     "IRS confirms the EIN, but Middesk/OpenCorporates cannot find the SOS filing. "
     "Usually: entity too new (<2 weeks, SOS propagation lag), incorporated in a slow-propagating state, "
     "or operating under a DBA name not matching SOS records. Entity existence is still unverifiable "
     "despite having a valid EIN.",
     "sos_match_boolean=false/null AND tin_match_boolean=true",
     "index.ts:1371 sos_match per-vendor logic  ·  index.ts:1421 sos_match_boolean derivation"),
]

for num, name, sev, desc, signal, source in segs:
    add_body(doc, f"Segment {num}: {name}", bold=True, color=C_AMBER, size=10)
    add_severity_badge(doc, sev if sev != "CLEAN" else "LOW", name, "")
    add_body(doc, desc, size=9.5)
    add_code(doc, f"Signal: {signal}")
    if source:
        add_source_box(doc, [source])
    doc.add_paragraph()

add_heading(doc, "Data Sources", level=2, color=C_PURPLE)
add_bullet(doc, "sos_match_boolean — funnel_df (index.ts:1421: engine.getResolvedFact('sos_match')?.value === 'success')")
add_bullet(doc, "sos_active — stats_df (index.ts:1426: filings.some(f => f.active===true || f.status==='active'))")
add_bullet(doc, "tin_match_boolean — stats_df (index.ts:482: dependent fact)")
add_bullet(doc, "tin_match_status — stats_df (index.ts:429: Middesk reviewTasks[key='tin'].status)")
add_bullet(doc, "tin_submitted — stats_df (index.ts:399: masked EIN from onboarding form)")
add_source_box(doc, ["integration-service/lib/facts/kyb/index.ts:399-491 (TIN)", "index.ts:1371-1435 (SOS)"])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.2  REGISTRATION DEPTH CASCADE
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.2  Registration Depth Cascade — Domestic & Foreign Analysis  (S1 internal)", level=1, color=C_BLUE)
add_body(doc,
    "Section 1 shows four KPI counts in a line. This section shows the RELATIONSHIPS between "
    "those counts — how businesses flow from one verification level to the next, and where they "
    "drop off. The explicit Domestic vs Foreign registration cross-analysis is the centrepiece.", size=10)

add_heading(doc, "The 3-Level Cascade Funnel", level=2, color=C_PURPLE)
add_body(doc,
    "Each arrow shows the drop-off count and rate — e.g. 'Of 30 Registry Found, 29 (97%) are "
    "Domestic. Of those 29, 28 (97%) match the operating state.'", size=10)
add_code(doc,
    "All Onboarded (N)\n"
    "    ↓\n"
    "Registry Found (sos_match_boolean=true)     →  No Registry (sos_match_boolean=false/null)\n"
    "    ↓\n"
    "Domestic Reg Found (formation_state NOT IN {DE,NV,WY,SD,MT,NM})  →  Foreign-Only (tax-haven)\n"
    "    ↓\n"
    "State Match (formation_state = primary_address.state)   →  State Mismatch (multi-state/error)")

add_heading(doc, "Panel A — Domestic Registration Deep-Dive", level=2, color=C_PURPLE)
add_body(doc,
    "Businesses where sos_domestic_verification=1 "
    "(clients.verification_results, verification_results.sql:36-39 — built when "
    "key='sos_domestic' AND sublabel='Domestic Active').", size=10)

dom_analysis = [
    "Formation state distribution — horizontal bar chart: which states dominate domestic registrations (CA, TX, FL, NY, etc.)?",
    "Formation state vs operating state agreement rate — within the domestic group, what % have formation_state = primary_address.state?",
    "Business age distribution — histogram of age in years (derived from formation_date), showing portfolio maturity.",
    "Domestic + State Match rate — these businesses have the strongest registry signal: incorporated AND operating in the same state.",
    "Drilldown: all Domestic businesses table showing business_id, formation_state, operating_state, sos_match_verif, sos_domestic_verif.",
]
for a in dom_analysis:
    add_bullet(doc, a)

add_heading(doc, "Panel B — Foreign-Only Registration (Tax-Haven) Deep-Dive", level=2, color=C_PURPLE)
add_body(doc,
    "Businesses where sos_match_boolean=true AND formation_state IN {DE, NV, WY, SD, MT, NM}. "
    "Middesk's address-based SOS search found a FOREIGN QUALIFICATION filing in the operating state "
    "but the DOMESTIC incorporation filing in the tax-haven state was NOT retrieved. "
    "This is the check_agent_v2 rule 'tax_haven_state' (NOTICE) and "
    "'formation_operating_mismatch' (MEDIUM).", size=10)

for_analysis = [
    "Tax-haven distribution — which states? (DE almost always dominates for LLCs/corps seeking legal advantages; WY for newer privacy structures; NV for older ones.)",
    "Operating state distribution within the foreign-only group — where are these DE/NV entities actually operating? (TX, CA, FL, NY are the most common.)",
    "Entity resolution gap explanation — for each foreign-only business: the domestic filing exists in the tax-haven state but Middesk could not retrieve it via address-based search. sos_domestic_verif=0 for all of these.",
    "Cross-reference with sos_match_verif (Submitted Active): is the foreign qualification confirmed as active?",
    "Drilldown: all Foreign-Only businesses table showing business_id, formation_state, operating_state, sos_match_verif, sos_domestic_verif.",
    "Action recommendation: verify (1) the domestic filing in DE/NV state SOS portal; (2) confirm a foreign qualification exists in the operating state.",
]
for a in for_analysis:
    add_bullet(doc, a)

add_heading(doc, "Panel C — State Mismatch Analysis (Domestic but formation ≠ operating)", level=2, color=C_PURPLE)
add_body(doc,
    "Within the domestic-registered group, businesses where formation_state ≠ primary_address.state. "
    "These trigger check_agent_v2 rule 'formation_operating_mismatch' (MEDIUM).", size=10)

add_body(doc,
    "Three sub-categories shown with business counts and drilldowns:", bold=True, size=10)
mismatch_cats = [
    ("Multi-state operation",
     "Legitimately incorporated in State A, operating in State B. Very common for growing businesses. "
     "Foreign qualification required in the operating state. Example: California LLC operating in Texas."),
    ("HQ vs formation state",
     "Business entered its main office state on the onboarding form but is incorporated elsewhere. "
     "Common: incorporated in Delaware, HQ in New York. primary_address.state = 'NY', formation_state = 'DE'. "
     "primary_address.state is SUBMITTED (not verified) — it reflects what the business told us."),
    ("Potential data entry error",
     "Formation_state and operating state differ by only one similar state, or the business is very local "
     "and there is no plausible reason for multi-state structure. Needs manual verification."),
]
for cat, desc in mismatch_cats:
    add_body(doc, f"  • {cat}:", bold=True, size=10)
    add_body(doc, f"    {desc}", size=9.5)

add_body(doc,
    "State-pair frequency table: formation_state → operating_state pairs sorted by count. "
    "Shows the most common geographic patterns across the portfolio.", size=10)

add_heading(doc, "Data Sources", level=2, color=C_PURPLE)
add_bullet(doc, "sos_match_boolean — funnel_df (index.ts:1421)")
add_bullet(doc, "sos_match_verif, sos_domestic_verif — stats_df (clients.verification_results — verification_results.sql:36-51)")
add_bullet(doc, "formation_state — funnel_df + stats_df (Middesk businessEntityVerification.formation_state)")
add_bullet(doc, "operating_state — funnel_df.operating_state (primary_address.value.state, extracted in _load_kyb_funnel_for_bids)")
add_bullet(doc, "formation_date — stats_df (for age_business calculation)")
add_bullet(doc, "Tax-haven set: {DE, NV, WY, SD, MT, NM} — same as check_agent_v2.TAX_HAVENS and Section 1 _SEG_CALC")
add_source_box(doc, [
    "warehouse-service/.../verification_results.sql:36-39 (sos_domestic_verification)",
    "integration-service/lib/facts/kyb/index.ts:1371-1435 (sos_match, sos_active)",
    "check_agent_v2.py: tax_haven_state (NOTICE), formation_operating_mismatch (MEDIUM)",
])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.3  DATA INTEGRITY CONTRADICTIONS
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.3  Data Integrity Contradictions — The Impossible States", level=1, color=C_BLUE)
add_body(doc,
    "These are signal combinations that should NEVER exist if the pipeline is working correctly. "
    "They are not risk signals — they are potential bugs, data freshness gaps, or pipeline "
    "ordering issues. Each maps to a specific logical impossibility defined in check_agent_v2.", size=10)

contras = [
    (
        "Contradiction 1",
        "SOS Active=True AND SOS Match=False",
        "HIGH",
        "sos_active='true' AND sos_match_boolean='false'.",
        "sos_active is a DEPENDENT fact derived FROM sos_filings (index.ts:1426). "
        "If sos_match_boolean=false, the vendor returned 'failure' and no sos_filings array was written. "
        "Therefore sos_active should be undefined — not 'true'. This combination is logically impossible "
        "from the same fact-write cycle. Indicates stale cached filing data or a race condition.",
        "Query raw sos_filings from PostgreSQL RDS port 5432 (too large for Redshift). "
        "File a bug report if count > 0.",
        "index.ts:1426 (sos_active) vs index.ts:1421 (sos_match_boolean)  ·  check_agent_v2: sos_active_match_contradiction (HIGH)",
    ),
    (
        "Contradiction 2",
        "sos_match_verif=1 AND sos_active=false",
        "MEDIUM",
        "clients.verification_results.sos_match_verification=1 AND stats_df.sos_active='false'.",
        "sos_match_verification=1 means sublabel='Submitted Active' was recorded in business_entity_review_task "
        "at verification time (verification_results.sql:42-45). "
        "sos_active=false means the sos_filings array currently shows all filings as inactive (index.ts:1434). "
        "These are computed at different times — the most likely explanation is a legitimate state change: "
        "the entity was active when verified but has since been revoked/dissolved.",
        "Check formation_state and the date of verification vs current SOS status. "
        "If the entity changed status recently, this is expected. If not, investigate data staleness.",
        "verification_results.sql:42-45 (sos_match_verification)  ·  index.ts:1434 (sos_active=false path)",
    ),
    (
        "Contradiction 3",
        "sos_domestic_verif=1 AND sos_match_verif=0",
        "MEDIUM",
        "sos_domestic_verification=1 (Domestic Active confirmed) AND sos_match_verification=0 (no Submitted Active).",
        "From verification_results.sql: sos_domestic_verification=1 when key='sos_domestic' AND sublabel='Domestic Active'. "
        "sos_match_verification=1 when key='sos_match' AND sublabel='Submitted Active'. "
        "Logically, a Domestic Active filing should also produce a Submitted Active result for the same business. "
        "This discrepancy suggests the two review tasks were written from different verification runs "
        "or the stored procedure sp_truncate_and_insert_verification_results() has a data quality issue.",
        "Run the diagnostic SQL to inspect raw business_entity_review_task rows for both keys. "
        "Check if the rows are from different business_entity_verification IDs.",
        "verification_results.sql:36-45 (both flag derivations)  ·  business_entity_review_task table",
    ),
    (
        "Contradiction 4",
        "tin_match_status=success AND tin_match_boolean=false",
        "CRITICAL",
        "Raw tin_match.value.status='success' AND tin_match_boolean='false'.",
        "tin_match_boolean is a DEPENDENT fact with the exact derivation (index.ts:482-490): "
        "engine.getResolvedFact('tin_match')?.value === 'success' "
        "|| engine.getResolvedFact('tin_match')?.value?.status === 'success'. "
        "If tin_match.value.status='success', tin_match_boolean MUST be true. "
        "If count > 0, this is a backend fact derivation bug in integration-service.",
        "File an immediate bug report against integration-service lib/facts/kyb/index.ts lines 488-490. "
        "This is the highest-confidence bug indicator in the dataset.",
        "index.ts:482-490 (tin_match_boolean DEPENDENT fact exact derivation)  ·  check_agent_v2: tin_bool_status_mismatch (CRITICAL)",
    ),
    (
        "Contradiction 5",
        "kyb_complete=true AND (sos_active≠true OR tin_match≠true)",
        "HIGH",
        "kyb_complete='true' AND (sos_active is not 'true' OR tin_match_boolean is not 'true').",
        "kyb_complete=true is supposed to require both business_verified=true AND screened_people non-empty. "
        "Finding kyb_complete=true when SOS is inactive or TIN failed means the completion logic "
        "did not enforce these preconditions — or the signals changed after kyb_complete was written.",
        "Verify screened_people fact via PostgreSQL RDS. Check if kyb_complete was set manually "
        "via a manual override (request_type='fact_override'). Review the business timeline.",
        "check_agent_v2: kyb_complete_no_idv (HIGH)  ·  index.ts: kyb_complete derivation",
    ),
]

for num, name, sev, signal, desc, action, source in contras:
    add_body(doc, f"{num}: {name}", bold=True, color=C_RED if sev=="CRITICAL" else C_ORANGE, size=11)
    add_severity_badge(doc, sev, name, "")
    add_body(doc, "Signal:", bold=True, size=9.5)
    add_code(doc, signal)
    add_body(doc, "Why this is impossible:", bold=True, size=9.5)
    add_body(doc, desc, size=9.5)
    add_body(doc, "Action:", bold=True, size=9.5)
    add_body(doc, action, size=9.5)
    add_source_box(doc, [source])

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.4  IDENTITY COMPLETENESS TRIANGLE
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.4  Identity Completeness Triangle  (TIN × IDV × is_sole_prop)", level=1, color=C_BLUE)
add_body(doc,
    "Worth AI uses three overlapping identity checks, and is_sole_prop changes which ones apply "
    "(index.ts:552: is_sole_prop = true when tin_submitted=null AND idv_passed_boolean=true). "
    "For sole props, IDV is the correct path. For non-sole-props, BOTH TIN and IDV should be verified.", size=10)

add_heading(doc, "6 Meaningful Segments", level=2, color=C_PURPLE)
id_segs = [
    ("1", "Non-sole-prop + TIN Pass + IDV Pass",
     "CLEAN",
     "Complete verification. Both entity EIN (IRS confirmed) and beneficial owner (Plaid biometrics) verified. The gold standard KYB completion state."),
    ("2", "Non-sole-prop + TIN Pass + IDV Fail/Missing",
     "HIGH",
     "Entity EIN confirmed by IRS but beneficial owner NOT verified. KYB is structurally incomplete. "
     "check_agent_v2: no_idv_not_sole_prop (HIGH). Action: send IDV link to owner."),
    ("3", "Non-sole-prop + TIN Fail + IDV Pass",
     "MEDIUM",
     "Beneficial owner confirmed via Plaid biometrics but entity EIN has an IRS name mismatch. "
     "Usually means DBA submitted instead of legal name on EIN certificate. "
     "check_agent_v2: sos_active_tin_failed (MEDIUM). Action: request IRS CP-575 or 147C."),
    ("4", "Non-sole-prop + TIN Fail + IDV Fail",
     "CRITICAL",
     "Both identity verification paths failed for a non-sole-prop entity. "
     "Neither the entity EIN nor the owner identity can be confirmed. Highest identity risk state. "
     "Do NOT approve. Full manual review required."),
    ("5", "Sole-prop + IDV Pass",
     "CLEAN",
     "Correct verification path for sole proprietorships. The EIN check is NOT applicable "
     "because sole props use the owner's SSN, not an employer EIN. IDV is the definitive check. "
     "TIN failure (if present) is expected and should not be flagged."),
    ("6", "Sole-prop + TIN Fail + IDV Pass",
     "NOTICE",
     "Structurally clean for a sole prop — the TIN failure is expected behaviour. "
     "check_agent_v2: sole_prop_tin_verification (NOTICE) surfaces this explicitly so "
     "underwriters understand the TIN failure is NOT a red flag for this entity type."),
]

for num, name, sev, desc in id_segs:
    add_body(doc, f"Segment {num}: {name}", bold=True, size=10,
             color=C_GREEN if "CLEAN" in sev else C_RED if "CRITICAL" in sev else C_AMBER)
    add_severity_badge(doc, sev if sev != "CLEAN" else "LOW", "", "")
    add_body(doc, desc, size=9.5)
    doc.add_paragraph()

add_heading(doc, "KYB Completion Funnel", level=2, color=C_PURPLE)
add_body(doc,
    "A stacked bar chart showing what % of the portfolio has completed each verification layer, "
    "and which layer is the most common bottleneck. "
    "Three layers: Registry (S1) → TIN (S2) → IDV (S3). "
    "The bottleneck layer tells the platform team where to focus integration improvements.", size=10)

add_source_box(doc, [
    "index.ts:552 (is_sole_prop DEPENDENT fact)",
    "index.ts:482 (tin_match_boolean)",
    "index.ts:541 (idv_passed_boolean)",
    "check_agent_v2: sole_prop_tin_verification (NOTICE), no_idv_not_sole_prop (HIGH), no_kyb_complete_sos_ok (NOTICE)",
])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.5  REGISTRY × RISK SIGNAL MATRIX
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.5  Registry × Risk Signal Matrix  (S1 × S4)", level=1, color=C_BLUE)
add_body(doc,
    "Entity existence status changes what a risk signal MEANS. A watchlist hit on a verified, "
    "active entity is actionable. A watchlist hit on an entity with no registry is qualitatively "
    "different — you cannot safely attribute the hit OR rule it out. This matrix makes that "
    "distinction visible at portfolio scale.", size=10)

add_heading(doc, "The 4×4 Matrix: Registry State × Risk Type", level=2, color=C_PURPLE)
add_matrix_table(doc,
    ["Registry State", "Watchlist Hit", "Bankruptcies", "Judgements/Liens", "Adverse Media"],
    [
        ["Active (verified)",       "Actionable — investigate",    "Credit risk — post-discharge?",  "Financial distress signal", "Reputational risk"],
        ["Inactive (dissolved)",    "HIGH risk — entity not in standing", "May have dissolved DUE TO BK", "Very high — compounded",    "HIGH risk"],
        ["Foreign-Only (tax-haven)","High — entity resolution gap", "Need formation state BK check", "Elevated",                  "Elevated"],
        ["No Registry",             "Cannot act safely",           "Cannot attribute safely",         "Cannot attribute safely",   "Cannot confirm entity"],
    ])

add_heading(doc, "6 High-Severity Named Findings", level=2, color=C_PURPLE)

risk_combos = [
    (
        "1", "Watchlist Hit + No Registry", "HIGH",
        "watchlist_hits > 0 AND sos_match_boolean=false/null",
        "Cannot verify entity existence AND has a compliance screening hit. "
        "Acting on the watchlist hit when entity identity is uncertain risks misattributing "
        "the hit to the wrong business. check_agent_v2: watchlist_low_sos_confidence (HIGH).",
        "Enhanced due diligence: first verify entity identity (request formation documents, EIN letter) "
        "before acting on watchlist hit. Do not decline solely on unverified hits.",
    ),
    (
        "2", "Watchlist Hit + SOS Inactive", "HIGH",
        "watchlist_hits > 0 AND sos_active=false",
        "Entity is dissolved/revoked AND has a watchlist hit. Possible: entity dissolved to evade "
        "sanctions listing, or dissolution is coincidental. Either way: highest compliance risk state.",
        "Same-day escalation to compliance team. Document the dissolution date vs watchlist listing date.",
    ),
    (
        "3", "Multiple Public Records Simultaneously", "HIGH",
        "num_bankruptcies > 0 AND (num_judgements > 0 OR num_liens > 0)",
        "Entity has bankruptcies AND judgements or liens at the same time. Multiple public record "
        "categories active simultaneously indicates systemic financial distress, not an isolated event. "
        "check_agent_v2: multiple_public_records (HIGH).",
        "Full public records review required before any credit decision. "
        "Aggregate total exposure across all public record types.",
    ),
    (
        "4", "Bankruptcy + SOS Inactive", "HIGH",
        "num_bankruptcies > 0 AND sos_active=false",
        "Entity may have dissolved BECAUSE of the bankruptcy — very different from an entity that "
        "survived bankruptcy and is still active. check_agent_v2: bankruptcy_active_sos notes "
        "that post-discharge entities carry elevated credit risk; dissolved entities are higher still.",
        "Check bankruptcy filing date vs dissolution date. Post-discharge survivor vs "
        "dissolved-during-proceedings entities require different treatment.",
    ),
    (
        "5", "Adverse Media + No Formal Watchlist Hit", "NOTICE",
        "adverse_media_hits > 0 AND watchlist_hits = 0",
        "Negative news/media coverage exists but no formal PEP or sanctions listing. "
        "Adverse media is intentionally EXCLUDED from consolidated watchlist.value "
        "(rules.ts:294 — filteredMetadata filters out WATCHLIST_HIT_TYPE.ADVERSE_MEDIA) "
        "because it is not a regulatory requirement. But it represents reputational risk. "
        "check_agent_v2: adverse_media_no_watchlist (NOTICE).",
        "Review adverse media content. Assess whether the coverage relates to the applicant "
        "entity or a name-similar entity. Consider enhanced monitoring.",
    ),
    (
        "6", "High Revenue + No Registry", "MEDIUM",
        "revenue > 0 AND sos_match_boolean=false/null",
        "Firmographic data (ZoomInfo/Equifax entity-matching XGBoost model) suggests a real, "
        "operating business, but SOS verification failed. The vendor matched a firmographic "
        "record to this business but Middesk/OC could not find the SOS filing. Possible entity "
        "resolution error: the matched firmographic record may belong to a different legal entity. "
        "check_agent_v2: no_firmographic_data does the inverse; this combination does the converse.",
        "Cross-reference the revenue source vendor (ZoomInfo pid=24 vs Equifax pid=17). "
        "Verify whether the matched entity's EIN matches the submitted EIN.",
    ),
]

for num, name, sev, signal, desc, action in risk_combos:
    add_body(doc, f"Finding {num}: {name}", bold=True, color=C_RED if sev=="HIGH" else C_AMBER, size=10)
    add_severity_badge(doc, sev, name, "")
    add_code(doc, f"Signal: {signal}")
    add_body(doc, desc, size=9.5)
    add_body(doc, f"Action: {action}", italic=True, size=9.5)
    doc.add_paragraph()

add_source_box(doc, [
    "stats_df: sos_match_boolean, sos_active, watchlist_hits, num_bankruptcies, num_judgements, num_liens, adverse_media, revenue",
    "integration-service/lib/facts/rules.ts:294 (adverse media exclusion from watchlist)",
    "check_agent_v2: watchlist_low_sos_confidence (HIGH), adverse_media_no_watchlist (NOTICE), multiple_public_records (HIGH), bankruptcy_active_sos (MEDIUM)",
])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.6  NAICS CLASSIFICATION COMPLETENESS
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.6  NAICS Classification Completeness  (S3 × S1 × Firmographic)", level=1, color=C_BLUE)
add_body(doc,
    "NAICS=561499 (hardcoded NAICS_OF_LAST_RESORT in aiNaicsEnrichment.ts:63) is not just a bad "
    "classification — it propagates downstream into the Worth Score Company Profile category "
    "and into the MCC code. This section diagnoses WHY 561499 occurs across the portfolio.", size=10)

add_heading(doc, "The Three Classification Gap Types", level=2, color=C_PURPLE)
gaps = [
    ("G1 — No NAICS Signal at All",
     "naics_code IS NULL",
     "MEDIUM",
     "No vendor (ZI/EFX/OC/Middesk/AI) provided any NAICS classification. "
     "Rarest case. Usually means the business is extremely new or entity matching failed completely."),
    ("G2 — Website Present but NAICS=561499",
     "website IS NOT NULL AND naics_code = '561499'",
     "MEDIUM",
     "The AI enrichment (AINaicsEnrichment, aiNaicsEnrichment.ts) received the website URL "
     "(params.website in the prompt) but either: (a) the web_search tool timed out (5s limit, "
     "aiNaicsEnrichment.ts:151-152), (b) the page was unreadable/blocked, or (c) the page had "
     "no industry-indicative content. check_agent_v2: website_no_naics (MEDIUM). "
     "This is the most FIXABLE gap — re-running classification with the website URL explicitly "
     "provided would likely resolve it."),
    ("G3 — No Website + NAICS=561499",
     "website IS NULL AND naics_code = '561499'",
     "LOW",
     "The most common case. No website was submitted and no vendor had a prior NAICS code. "
     "The AI classified from business name + address only and could not determine industry. "
     "This is Gap G1 in the aiNaicsEnrichment naming. check_agent_v2: naics_fallback (LOW)."),
]

for gname, signal, sev, desc in gaps:
    add_body(doc, gname, bold=True, size=10, color=C_AMBER)
    add_severity_badge(doc, sev, gname, "")
    add_code(doc, f"Signal: {signal}")
    add_body(doc, desc, size=9.5)
    doc.add_paragraph()

add_heading(doc, "561499 Distribution Analysis — 4 Sub-analyses", level=2, color=C_PURPLE)
dist_analysis = [
    "Total 561499 count vs healthy NAICS codes — what % of the portfolio is unclassified?",
    "Formation state distribution of 561499 businesses — do they cluster in certain states (suggesting a systematic entity-matching gap)?",
    "Score decision cross-reference — how many APPROVE/REVIEW/DECLINE decisions have NAICS=561499? A business with 561499 that still gets APPROVE means the Plaid/financial signals are very strong. A business with 561499 AND DECLINE is likely being penalised by the Company Profile category in the score model.",
    "MCC contamination — count of businesses where mcc_code is the fallback (5614), derived from NAICS=561499 via the rel_naics_mcc lookup. These businesses have an inaccurate merchant category code affecting payments processing decisions.",
]
for a in dist_analysis:
    add_bullet(doc, a)

add_source_box(doc, [
    "integration-service/lib/aiEnrichment/aiNaicsEnrichment.ts:63 (NAICS_OF_LAST_RESORT='561499')",
    "aiNaicsEnrichment.ts:151-152 (web_search tool, 5s timeout)",
    "check_agent_v2: naics_fallback (LOW), website_no_naics (MEDIUM), mcc_naics_mismatch (LOW)",
])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.7  SCORE vs KYB PURITY MATRIX
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.7  Score vs KYB Purity Matrix  (S5 × S1+S2+S3+S4)", level=1, color=C_BLUE)
add_body(doc,
    "The underwriting oversight cross-reference. The Worth Score (weighted_score_850) is an ML model output. "
    "It can APPROVE a business with active watchlist hits if financial signals are strong enough, and can "
    "DECLINE a business where all KYB checks pass if financial data is missing. This matrix makes those "
    "overrides visible to the underwriting and compliance teams.", size=10)

add_heading(doc, "KYB Purity Classification — 4 Tiers", level=2, color=C_PURPLE)
purity_tiers = [
    ("KYB Clean",
     "sos_match_boolean=true AND sos_active=true AND tin_match_boolean=true AND watchlist_hits=0 AND num_bankruptcies=0",
     "All four primary KYB verifications pass with no risk signals."),
    ("KYB Minor Issues",
     "Any one of: TIN fail, NAICS=561499, revenue missing, formation/operating state mismatch",
     "One verification gap or data quality issue. Typically resolvable."),
    ("KYB Major Issues",
     "SOS inactive, TIN fail + no registry, bankruptcies>0, multiple public records categories",
     "Significant verification failure or active financial distress signal."),
    ("KYB Critical",
     "watchlist_hits > 0, no registry at all (sos_match_boolean=false/null), or multiple CRITICAL anomaly triggers",
     "Compliance signal or complete entity unverifiability. Requires compliance escalation."),
]

for tier, signal, desc in purity_tiers:
    add_body(doc, tier, bold=True, size=10, color=C_BLUE)
    add_code(doc, f"Condition: {signal}")
    add_body(doc, desc, size=9.5)
    doc.add_paragraph()

add_heading(doc, "The 4×4 Matrix", level=2, color=C_PURPLE)
add_matrix_table(doc,
    ["KYB Purity \\ Score Decision", "APPROVE", "FURTHER REVIEW", "DECLINE", "Not Scored"],
    [
        ["KYB Clean",          "✅ Consistent",          "Possible upgrade",     "⚠️ Score vs KYB gap",  "❓ Monitor"],
        ["KYB Minor Issues",   "Monitor",                "✅ Consistent",        "✅ Consistent",         "❓ Monitor"],
        ["KYB Major Issues",   "⚠️ Review override",    "✅ Consistent",        "✅ Consistent",         "🔴 Urgent"],
        ["KYB Critical",       "🔴 Compliance review",  "⚠️ Monitor closely",   "✅ Consistent",         "🔴 Highest priority"],
    ])

add_heading(doc, "4 Flagged Override Cells — each shown as a named finding card with business ID drilldown", level=2, color=C_PURPLE)
overrides = [
    (
        "APPROVE + KYB Critical", "CRITICAL",
        "The ML model approved a business that has active watchlist hits, multiple bankruptcies, "
        "or no registry record at all. This occurs when the financial model (PyTorch, Plaid banking data) "
        "is very strong and overrides the compliance/KYB signals. These businesses need a compliance "
        "team review REGARDLESS of the score. The score is not wrong — but compliance review is required.",
        "Flag all APPROVE decisions where KYB Critical conditions are present. "
        "Compliance team must sign off before disbursement.",
    ),
    (
        "DECLINE + KYB Clean", "HIGH",
        "The model declined a business where all four KYB sections show no issues. "
        "This almost always means the financial sub-model (PyTorch) had no input data: "
        "no Plaid banking connected, no revenue from ZI/EFX, and formation_date missing causing "
        "age_business=0. The business may be legitimate but underserved by the model's data coverage. "
        "These are prime candidates for manual underwriting review.",
        "Identify which model input features are null (revenue, formation_date, Plaid data). "
        "Consider requesting direct financial statements as an alternative data path.",
    ),
    (
        "Not Scored + KYB Critical", "CRITICAL",
        "A business with high-risk signals (watchlist hit, no registry, multiple public records) "
        "has NEVER been scored. This is the highest-priority blind spot for the underwriting team. "
        "The risk signals are known but there is no model assessment at all.",
        "Prioritise for immediate manual review. Escalate to the AI-score-service team "
        "to understand why scoring has not run. If >48h old, manual underwriting decision required.",
    ),
    (
        "Not Scored + KYB Clean", "NOTICE",
        "Business has clean KYB signals but no score. Usually means it is very new (<48h) "
        "or the Plaid banking connection has not been established yet.",
        "Monitor. If business is >48h old and Plaid is connected, escalate to AI-score-service.",
    ),
]

for name, sev, desc, action in overrides:
    add_body(doc, name, bold=True, size=10, color=C_RED if sev in ("CRITICAL","HIGH") else C_AMBER)
    add_severity_badge(doc, sev, name, "")
    add_body(doc, desc, size=9.5)
    add_body(doc, f"Action: {action}", italic=True, size=9.5)
    doc.add_paragraph()

add_source_box(doc, [
    "stats_df + funnel_df (KYB Purity classification signals)",
    "_home_ws_clean (score decision from Section 5 _load_home_scores)",
    "aiscore.py: score_300_850 = final_proba × 550 + 300",
    "check_agent_v2: all DETERMINISTIC_CHECKS relevant to KYB Critical classification",
])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# 6.8  ANOMALY ACCUMULATION TRIAGE
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "6.8  Anomaly Accumulation Triage — Portfolio Synthesis", level=1, color=C_BLUE)
add_body(doc,
    "Each of sub-sections 6.1–6.7 identifies specific anomaly types. This final section counts "
    "how many anomalies each business accumulates across ALL types and ranks them. "
    "This is distinct from the Section 4 Red Flag Distribution (which is based on risk signal scores). "
    "This ranking is based on STRUCTURAL DATA CONSISTENCY failures — a business can rank high here "
    "even with zero watchlist hits, just from contradictory SOS/TIN/IDV signals.", size=10)

add_heading(doc, "Per-Business Anomaly Score — Weighted by Severity", level=2, color=C_PURPLE)
add_body(doc,
    "For each business in the portfolio, compute a weighted anomaly score across all 6.1–6.7 types:", size=10)
weights = [
    ("CRITICAL anomaly", "4 points", "e.g. tin_match_status=success AND tin_match_boolean=false"),
    ("HIGH anomaly",     "3 points", "e.g. Watchlist Hit + No Registry; Registry Inactive + TIN Pass"),
    ("MEDIUM anomaly",   "2 points", "e.g. Registry Active + TIN Fail; State mismatch without foreign qual."),
    ("NOTICE/LOW",       "1 point",  "e.g. Tax-haven formation state; NAICS=561499 G3 gap"),
]
add_matrix_table(doc,
    ["Severity", "Points", "Example"],
    weights)

add_heading(doc, "Anomaly Count Histogram", level=2, color=C_PURPLE)
add_body(doc,
    "Bar chart: X-axis = anomaly count per business (0, 1, 2, 3, 4, 5+). "
    "Y-axis = number of businesses. "
    "This shows: what % of the portfolio is structurally clean (0 anomalies), "
    "whether the portfolio has a long tail of high-anomaly businesses, "
    "and the shape of the distribution — concentrated at 0-1 (healthy) vs spread across 2+ (systemic issues).", size=10)

add_heading(doc, "Top 10 Highest-Anomaly Businesses", level=2, color=C_PURPLE)
add_body(doc,
    "A ranked table of the 10 businesses with the most simultaneous anomalies, showing:", size=10)
top10_cols = [
    "Business ID",
    "Total anomaly count (unweighted)",
    "Weighted anomaly score",
    "Anomaly type flags — one column per anomaly type, checkmark if triggered",
    "Investigate button — clicking sets this business ID in the filter bar (same _pending_bid mechanism as Home tab)",
]
for col in top10_cols:
    add_bullet(doc, col)

add_body(doc,
    "Note: This ranking differs from the existing 'Top 10 Businesses Needing Attention' "
    "(Section 6 of the current Home tab, based on red flag score). "
    "A business can rank high here purely from structural KYB contradictions — "
    "e.g. the SOS/TIN/IDV impossible-state contradictions — "
    "even with a clean watchlist and no public records.", size=10, italic=True)

add_heading(doc, "Anomaly Co-Occurrence Heatmap", level=2, color=C_PURPLE)
add_body(doc,
    "A symmetric matrix showing, for each pair of anomaly types, "
    "how many businesses trigger BOTH simultaneously.", size=10)
add_body(doc,
    "The diagonal = individual anomaly count. "
    "Off-diagonal cells = co-occurrence count. "
    "Color-coded by intensity (darker = more businesses with both anomalies simultaneously).", size=10)
add_body(doc,
    "Key patterns to look for:", bold=True, size=10)
cooc_patterns = [
    "NAICS=561499 AND Revenue Missing — these cluster together because both result from the same root cause: ZI/EFX entity matching failed, leaving both firmographic signals empty.",
    "No Registry AND TIN Not Checked — entity existence unverified AND no EIN check. The most fundamentally incomplete onboarding state.",
    "Tax-Haven Formation State AND State Mismatch — by definition all tax-haven formations that are not operating in the formation state will also trigger the mismatch rule.",
    "Registry Inactive AND Multiple Public Records — entity under financial distress may have had SOS revoked for non-payment of fees during the same period as the financial distress.",
    "Watchlist Hit AND NAICS=561499 — a business in an unknown industry with a compliance hit. The absence of industry classification means industry-specific risk calibration is not possible.",
]
for p in cooc_patterns:
    add_bullet(doc, p)

add_body(doc,
    "The co-occurrence heatmap reveals which anomalies are INDEPENDENT (low co-occurrence) "
    "vs which ones CLUSTER together (high co-occurrence — systematic pipeline failures tend to cluster "
    "for the same class of micro-business or for businesses where a specific vendor integration failed).", size=10)

add_source_box(doc, [
    "All signals from stats_df + funnel_df + _home_ws_clean — zero new Redshift queries",
    "Anomaly definitions: all map to check_agent_v2.DETERMINISTIC_CHECKS and sub-sections 6.1–6.7",
    "Weighted scoring: CRITICAL=4, HIGH=3, MEDIUM=2, NOTICE/LOW=1 (consistent with check_agent_v2 severity model)",
])
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# UI ORGANISATION
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "UI Organisation & Interaction Design", level=1, color=C_BLUE)

add_heading(doc, "Overall Layout", level=2, color=C_PURPLE)
ui_layout = [
    "Section 6 opens with the 6.0 Portfolio Health Scorecard (5 KPI banner + severity distribution bar).",
    "Sub-sections 6.1–6.7 are presented in expandable sections (st.expander) or tabs, each self-contained.",
    "Section 6.8 (Anomaly Accumulation Triage) is always shown expanded at the bottom — it is the synthesis of everything above.",
    "All drilldowns use the same _drilldown_table() function and _SEG_CALC explanation format as Sections 1–4 (consistent UX).",
    "Every finding includes a detail_panel() with source SQL, JSON, and source file citations (same as all other Home tab sections).",
]
for u in ui_layout:
    add_bullet(doc, u)

add_heading(doc, "Per Sub-Section Template", level=2, color=C_PURPLE)
template = [
    "1. Header + caption — what question does this analysis answer?",
    "2. Main chart or table (matrix, heatmap, histogram, stacked bar).",
    "3. Named finding cards — one per key pattern — with severity badge, description, action, and source.",
    "4. Drilldown expanders — business IDs per finding with the exact contradicting signal columns.",
    "5. ⚙️ detail_panel — source SQL, JSON, source file citations.",
]
for t in template:
    add_bullet(doc, t)

add_heading(doc, "Data Dependencies", level=2, color=C_PURPLE)
add_body(doc,
    "Section 6 requires zero new Redshift queries. All signals are already present:", size=10)
data_deps = [
    ("stats_df", "_load_stats_for_bids()", "sos_match_boolean, sos_match_status, sos_active, sos_match_verif, sos_domestic_verif, sos_active_verif, tin_submitted, tin_match_status, tin_match, idv_passed, naics_code, watchlist_hits, num_bankruptcies, num_judgements, num_liens, adverse_media, revenue, formation_state, formation_date"),
    ("funnel_df", "_load_kyb_funnel_for_bids()", "sos_match_boolean, sos_active, formation_state, operating_state, tin_submitted, tin_match_boolean (tin_status), middesk_confidence"),
    ("_home_ws_clean", "_load_home_scores()", "weighted_score_850, risk_level, score_decision"),
]
add_matrix_table(doc,
    ["DataFrame", "Loaded by", "Signals used in Section 6"],
    data_deps)

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════
# COMPLETE ANOMALY RULE INDEX
# ═══════════════════════════════════════════════════════════════════════
add_heading(doc, "Complete Anomaly Rule Index", level=1, color=C_BLUE)
add_body(doc,
    "Every anomaly in Section 6 maps to a source-code rule. This index is the authoritative reference.", size=10)

all_rules = [
    # sub, name, severity, check_agent rule / source
    ("6.1", "Registry Active + TIN Fail",                    "MEDIUM",   "check_agent: sos_active_tin_failed · index.ts:429,482"),
    ("6.1", "Registry Inactive + TIN Pass",                  "HIGH",     "check_agent: sos_inactive_tin_ok · index.ts:1426,482"),
    ("6.1", "No Registry + TIN Pass",                        "HIGH",     "index.ts:1371-1424 (sos_match per vendor) · 1421 (sos_match_boolean)"),
    ("6.2", "Tax-Haven Formation State",                      "NOTICE",   "check_agent: tax_haven_state · verification_results.sql:36-39"),
    ("6.2", "Formation State ≠ Operating State",              "MEDIUM",   "check_agent: formation_operating_mismatch · index.ts (formation_state, primary_address)"),
    ("6.2", "Foreign-Only: sos_domestic_verif=0",             "NOTICE",   "verification_results.sql:36-39 · check_agent: tax_haven_state"),
    ("6.3", "SOS Active=True AND SOS Match=False",            "HIGH",     "check_agent: sos_active_match_contradiction · index.ts:1426 vs 1421"),
    ("6.3", "sos_match_verif=1 AND sos_active=false",         "MEDIUM",   "verification_results.sql:42-45 vs index.ts:1434"),
    ("6.3", "sos_domestic_verif=1 AND sos_match_verif=0",     "MEDIUM",   "verification_results.sql:36-45 both derivations"),
    ("6.3", "tin_match_status=success AND tin_match=false",   "CRITICAL", "check_agent: tin_bool_status_mismatch · index.ts:482-490"),
    ("6.3", "kyb_complete=true AND sos/tin not passing",      "HIGH",     "check_agent: kyb_complete_no_idv · index.ts: kyb_complete derivation"),
    ("6.4", "Non-sole-prop + TIN Fail + IDV Pass",            "MEDIUM",   "check_agent: sos_active_tin_failed · index.ts:552 (is_sole_prop)"),
    ("6.4", "Non-sole-prop + TIN Pass + IDV Fail",            "HIGH",     "check_agent: no_idv_not_sole_prop · index.ts:541 (idv_passed_boolean)"),
    ("6.4", "Non-sole-prop + TIN Fail + IDV Fail",            "CRITICAL", "check_agent: no_idv_not_sole_prop + sos_active_tin_failed (both)"),
    ("6.4", "Sole-prop + TIN Fail + IDV Pass",                "NOTICE",   "check_agent: sole_prop_tin_verification · index.ts:552"),
    ("6.5", "Watchlist Hit + No Registry",                    "HIGH",     "check_agent: watchlist_low_sos_confidence · index.ts:1503-1541"),
    ("6.5", "Watchlist Hit + SOS Inactive",                   "HIGH",     "check_agent: watchlist_low_sos_confidence + sos_active_match_contradiction"),
    ("6.5", "Multiple Public Records",                        "HIGH",     "check_agent: multiple_public_records · index.ts: num_bankruptcies, num_judgements, num_liens"),
    ("6.5", "Bankruptcy + SOS Inactive",                      "HIGH",     "check_agent: bankruptcy_active_sos (inverse case) · index.ts:1426"),
    ("6.5", "Adverse Media + No Watchlist",                   "NOTICE",   "check_agent: adverse_media_no_watchlist · rules.ts:294 (adverse media exclusion)"),
    ("6.5", "High Revenue + No Registry",                     "MEDIUM",   "check_agent: no_firmographic_data (inverse) · sources.ts ZI/EFX entity matching"),
    ("6.6", "NAICS=561499, website present (G2)",             "MEDIUM",   "check_agent: website_no_naics · aiNaicsEnrichment.ts:63,151-152"),
    ("6.6", "NAICS=561499, no website (G3)",                  "LOW",      "check_agent: naics_fallback · aiNaicsEnrichment.ts:63"),
    ("6.6", "MCC from fallback NAICS=561499",                 "LOW",      "check_agent: mcc_naics_mismatch · rel_naics_mcc lookup"),
    ("6.7", "APPROVE + KYB Critical",                         "CRITICAL", "Score vs KYB Purity override · aiscore.py L44 formula"),
    ("6.7", "DECLINE + KYB Clean",                            "HIGH",     "Score vs KYB Purity gap · worth_score_model.py feature extraction"),
    ("6.7", "Not Scored + KYB Critical",                      "CRITICAL", "data_current_scores gap · ai-score-service pipeline"),
    ("6.7", "Not Scored + KYB Clean",                         "NOTICE",   "data_current_scores gap · scoring schedule"),
]

add_matrix_table(doc,
    ["Sub-section", "Anomaly Name", "Severity", "Source Rule / File"],
    all_rules)

# ═══════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_divider(doc)
add_body(doc,
    "This specification is grounded exclusively in the Worth AI repository source files: "
    "integration-service/lib/facts/kyb/index.ts · "
    "integration-service/lib/facts/rules.ts · "
    "integration-service/lib/facts/sources.ts · "
    "integration-service/lib/aiEnrichment/aiNaicsEnrichment.ts · "
    "warehouse-service/.../verification_results.sql · "
    "check_agent_v2.py (28 DETERMINISTIC_CHECKS). "
    "No guessing, no hallucination, no invented rules.",
    italic=True, color=C_LIGHT, size=8.5)

out_path = "/workspace/Admin-Portal-KYB-App/Section6_Portfolio_Anomaly_Scanner.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
