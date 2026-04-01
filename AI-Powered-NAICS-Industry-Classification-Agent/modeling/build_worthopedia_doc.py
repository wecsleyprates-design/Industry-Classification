"""
Build a Word (.docx) document structured as Worthopedia tabs
for the INDUSTRY_FACTS_GUIDE.md content.

Each "tab" becomes a clearly separated section with a coloured heading
that mirrors the tab-based layout of the Worthopedia Google Doc.

Run:  python modeling/build_worthopedia_doc.py
Output: modeling/Industry_Classification_Worthopedia.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin  = Inches(0.8)
section.bottom_margin = Inches(0.8)

# ── Colours ───────────────────────────────────────────────────────────────────
WORTH_PURPLE = RGBColor(0x5B, 0x21, 0xB6)   # purple — main brand
WORTH_BLUE   = RGBColor(0x1D, 0x4E, 0xD8)   # blue
WORTH_GREEN  = RGBColor(0x06, 0x5F, 0x46)   # dark green
WORTH_RED    = RGBColor(0x7F, 0x1D, 0x1D)   # dark red
WORTH_GRAY   = RGBColor(0x37, 0x41, 0x51)   # dark gray
TAB_BG       = RGBColor(0xED, 0xE9, 0xFE)   # light purple bg for tab headers

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_hex: str) -> None:
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)


def add_tab_header(doc, tab_number: int, tab_name: str, description: str) -> None:
    """Add a visually distinct tab header block."""
    doc.add_paragraph()
    # Horizontal rule simulation via border paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = WORTH_PURPLE
    run.font.size = Pt(8)

    # Tab label
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    r_num = p2.add_run(f"TAB {tab_number}  ")
    r_num.font.bold = True
    r_num.font.color.rgb = WORTH_PURPLE
    r_num.font.size = Pt(9)
    r_name = p2.add_run(tab_name.upper())
    r_name.font.bold = True
    r_name.font.color.rgb = WORTH_PURPLE
    r_name.font.size = Pt(9)

    # Tab title as H1
    h = doc.add_heading(tab_name, level=1)
    h.runs[0].font.color.rgb = WORTH_PURPLE
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(4)

    # Description
    if description:
        d = doc.add_paragraph(description)
        d.runs[0].font.italic = True
        d.runs[0].font.color.rgb = WORTH_GRAY
        d.runs[0].font.size = Pt(10)
        d.paragraph_format.space_after = Pt(10)


def h2(doc, text: str, colour=WORTH_BLUE) -> None:
    p = doc.add_heading(text, level=2)
    if p.runs:
        p.runs[0].font.color.rgb = colour
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def h3(doc, text: str) -> None:
    p = doc.add_heading(text, level=3)
    if p.runs:
        p.runs[0].font.color.rgb = WORTH_GRAY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def body(doc, text: str, bold_parts=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if bold_parts:
        # Simple bold prefix: "**key:** rest"
        for segment in re.split(r'(\*\*[^*]+\*\*)', text):
            if segment.startswith('**') and segment.endswith('**'):
                r = p.add_run(segment[2:-2])
                r.bold = True
                r.font.size = Pt(10)
            else:
                r = p.add_run(segment)
                r.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)


def bullet(doc, text: str, level: int = 0, bold_prefix: str = None) -> None:
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
        r = p.add_run(text)
        r.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)


def code_block(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)


def add_table(doc, headers: list, rows: list, col_widths=None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, 'E0E7FF')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WORTH_BLUE
    # Data rows
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if val.startswith('✅'):
                r = p.add_run(val)
                r.font.color.rgb = WORTH_GREEN
                r.font.size = Pt(9)
            elif val.startswith('❌') or val.startswith('⚠️'):
                r = p.add_run(val)
                r.font.color.rgb = WORTH_RED
                r.font.size = Pt(9)
            else:
                r = p.add_run(val)
                r.font.size = Pt(9)
    doc.add_paragraph()


def callout(doc, text: str, colour=WORTH_BLUE) -> None:
    """Shaded callout box using a 1-cell table."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, 'EEF2FF')
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.italic = True
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

title = doc.add_heading('Industry Classification Facts', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
if title.runs:
    title.runs[0].font.color.rgb = WORTH_PURPLE

sub = doc.add_paragraph('Complete Reference — How Worth AI classifies businesses into industries')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.italic = True
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.color.rgb = WORTH_GRAY

doc.add_paragraph()
callout(doc,
    "This document covers: the two classification pipelines, how each confidence score is produced, "
    "the 8 industry classification facts, the winner selection rules, what customers see, and the "
    "known gaps. All content verified against the SIC-UK-Codes and Entity-Matching source code repositories.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — INDUSTRY CLASSIFICATION OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_tab_header(doc, 1, "Industry Classification",
    "What it is, why it matters, and the two pipelines that produce it.")

h2(doc, "The Core Question")
body(doc, "When a business applies through Worth AI, one of the most critical questions is: "
     "\"What industry does this business operate in?\" The answer drives risk scoring, KYB "
     "decisions, AML flags, the Worth 360 Report, and what customers see in the Worth AI UI.")

h2(doc, "Two Separate Pipelines")
callout(doc,
    "Worth AI runs two completely separate pipelines for industry classification. "
    "They share some input data but have different goals, different source coverage, "
    "and different audiences. The customer ONLY sees Pipeline A output.")

add_table(doc,
    ["", "Pipeline A — Integration Service", "Pipeline B — Warehouse Service"],
    [
        ["Goal", "Deliver best classification in real-time using ALL available sources",
         "Build analytics table for model training and data exports"],
        ["Runs", "Every time a business is submitted (real-time)", "Scheduled batch job in Redshift"],
        ["Sources", "ZoomInfo, Equifax, OC, Middesk, Trulioo, SERP, AI (6+ sources)",
         "ZoomInfo + Equifax ONLY"],
        ["Output", "rds_warehouse_public.facts (JSONB, 217 facts)",
         "datascience.customer_files (wide denormalized table)"],
        ["Customer sees?", "✅ YES — REST API, Worth 360 Report, Worth AI UI",
         "❌ NO — internal analytics only"],
        ["Can pipelines differ?", "✅ YES — if OC or Middesk won, Pipeline A shows different NAICS",
         "Always uses ZI vs EFX winner"],
    ]
)

h2(doc, "Why Two Pipelines?")
body(doc, "Pipeline B was built when only ZoomInfo and Equifax had clean pre-loaded Redshift "
     "tables with numeric NAICS fields. OC, Liberty, Middesk, and Trulioo were added later "
     "and only integrated into Pipeline A. The SQL in customer_table.sql was never extended "
     "to include them.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — PIPELINE B (BATCH)
# ══════════════════════════════════════════════════════════════════════════════

add_tab_header(doc, 2, "Pipeline B — Batch Redshift",
    "How the production winner-takes-all rule works and what it produces.")

h2(doc, "Day 0: Loading Vendor Data into Redshift")
body(doc, "Before any business is submitted, Worth AI loads bulk data from vendors into Redshift. "
     "This is offline work done by the data engineering team.")

add_table(doc,
    ["Vendor", "Redshift Table", "Key Industry Column", "Coverage"],
    [
        ["ZoomInfo", "datascience.zoominfo_standard_ml_2",
         "zi_c_naics6 (6-digit NAICS)", "Global"],
        ["Equifax", "warehouse.equifax_us_standardized",
         "efx_primnaicscode (6-digit NAICS), efx_primsic (4-digit SIC)",
         "US only ⚠️ unknown cadence"],
        ["OpenCorporates", "datascience.open_corporates_standard_ml_2",
         "industry_code_uids (pipe-delimited: us_naics-541110|gb_sic-62012)",
         "Global"],
        ["Liberty", "liberty.einmst_20260218 + einmst_15_5mn + einmst_5_3m_remaining",
         "NAICS, SIC columns", "US only"],
    ]
)

h2(doc, "The XGBoost Entity Matching Model")
body(doc, "The Worth AI XGBoost entity matching model (entity_matching_20250127 v1) answers: "
     "\"Is this submitted business the same real-world entity as this vendor record?\" "
     "It uses 33 pairwise text/address similarity features and outputs a probability 0–1.")
body(doc, "Output stored in datascience.ml_model_matches:")
bullet(doc, "zi_probability — ZoomInfo match probability (0–1)")
bullet(doc, "efx_probability — Equifax match probability (0–1)")
bullet(doc, "oc_probability — OpenCorporates match probability (0–1)")

h3(doc, "How zi_match_confidence and efx_match_confidence are built")
body(doc, "From smb_zi_oc_efx_combined.sql (verified from source code):")
code_block(doc,
    "zi_match_confidence =\n"
    "  CASE WHEN zi_probability IS NOT NULL    THEN zi_probability    -- XGBoost score\n"
    "       WHEN similarity_index/55.0 >= 0.8  THEN similarity_index/55.0  -- heuristic\n"
    "       ELSE 0\n"
    "  END\n\n"
    "MAX_CONFIDENCE_INDEX = 55  (normalises raw Levenshtein score to 0-1)\n"
    "Threshold for heuristic: similarity_index >= 45 (= 0.82 confidence)")

h2(doc, "The Winner-Takes-All Rule (customer_table.sql)")
body(doc, "The stored procedure sp_recreate_customer_files() runs this rule — it controls "
     "EVERY firmographic field in customer_files, not just NAICS:")
code_block(doc,
    "WHEN COALESCE(zi_match_confidence, 0) > COALESCE(efx_match_confidence, 0)\n"
    "  THEN zi_c_naics6          -- ZoomInfo wins ALL fields\n"
    "ELSE efx_primnaicscode       -- Equifax wins ALL fields\n\n"
    "Fields controlled: NAICS, employee count, revenue, company name,\n"
    "                   address, city, ZIP, country, website, affiliate flag")

h3(doc, "Why OC and Liberty are excluded")
bullet(doc, "OC: oc_match_confidence EXISTS in smb_zi_oc_efx_combined — but industry_code_uids "
       "is a pipe-delimited string never parsed in customer_table.sql")
bullet(doc, "Liberty: Redshift tables exist but were never joined into smb_zi_oc_efx_combined")
bullet(doc, "Middesk / Trulioo: live API sources (Pipeline A only)")

h2(doc, "Pipeline B — Scenario Table")
add_table(doc,
    ["Situation", "What gets stored in customer_files", "Impact"],
    [
        ["ZI confidence > EFX", "All fields from ZoomInfo (zi_c_naics6, name, address...)",
         "ZoomInfo data in analytics table"],
        ["EFX confidence > ZI", "All fields from Equifax (efx_primnaicscode, name, address...)",
         "Equifax data in analytics table"],
        ["Both = 0 (no match)", "COALESCE fallback to existing naics_code or NULL",
         "NULL industry in analytics"],
        ["OC had better match", "⚠️ OC IGNORED — oc_match_confidence exists but SQL never reads it",
         "Weaker ZI/EFX code used"],
        ["Liberty / Middesk / Trulioo matched",
         "⚠️ IGNORED — live API sources not in batch pipeline",
         "Same ZI vs EFX result"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — PIPELINE A (REAL-TIME / FACT ENGINE)
# ══════════════════════════════════════════════════════════════════════════════

add_tab_header(doc, 3, "Pipeline A — Real-Time Fact Engine",
    "How the Fact Engine selects the winning industry code from all 6 sources.")

h2(doc, "The 6 Sources and Their Confidence Scores")
add_table(doc,
    ["Source", "Industry Code Field", "Weight", "Confidence Method", "Platform ID"],
    [
        ["OpenCorporates", "industry_code_uids → parse us_naics-XXXXXX", "0.9",
         "XGBoost oc_probability", "23"],
        ["Middesk", "NAICS from SOS filing", "2.0 ← highest",
         "XGBoost confidenceScoreMany() OR 0.15+0.20/task", "16"],
        ["ZoomInfo", "zi_c_naics6", "0.8",
         "XGBoost zi_probability OR similarity_index/55", "24"],
        ["Equifax", "efx_primnaicscode, efx_primsic", "0.7",
         "XGBoost efx_probability OR similarity_index/55", "17"],
        ["Trulioo", "sicCode (may be 4-digit = POLLUTED)", "0.8",
         "Heuristic match.index/55 ONLY — no XGBoost", "38"],
        ["AI (GPT-5-mini)", "AI-generated NAICS + MCC", "0.1 ← lowest",
         "Self-reported HIGH/MED/LOW text", "31"],
    ]
)

callout(doc,
    "All source responses are stored in integration_data.request_response "
    "(integration-service PostgreSQL). The Fact Engine reads each source via "
    "getter() functions querying platform_id = X.")

h2(doc, "The 6 Winner Selection Rules")
h3(doc, "Rule 1 — factWithHighestConfidence()")
body(doc, "Sort all sources by confidence (0–1). Highest wins outright if the gap between "
     "top two is > WEIGHT_THRESHOLD (0.05).")

h3(doc, "Rule 2 — weightedFactSelector() — tie-break")
body(doc, "If top two are within 0.05 of each other, use source weight to break the tie.")
body(doc, "Weight ranking: Middesk (2.0) > OC (0.9) > ZoomInfo (0.8) = Trulioo (0.8) > "
     "Equifax (0.7) > AI (0.1)")

h3(doc, "Rule 3 — manualOverride() — always wins")
body(doc, "Analyst override set via PATCH /facts/business/{id}/override/{factName} "
     "always wins regardless of any model or AI result.")

h3(doc, "Rule 4 — No minimum confidence cutoff")
body(doc, "There is NO minimum confidence threshold. If only one source returned a code at "
     "confidence 0.05, it wins. Low confidence IS visible in the API response.")

h3(doc, "Rule 5 — AI safety net")
body(doc, "AI enrichment (GPT-5-mini) triggers when fewer than 3 sources returned a NAICS code. "
     "Trigger: DEPENDENT_FACTS.naics_code = { maximumSources: 3, minimumSources: 1 }")

h3(doc, "Rule 6 — Last resort: 561499")
body(doc, "If winning code not found in core_naics_code table → removeNaicsCode() replaces it "
     "with 561499 (All Other Business Support Services). Analyst override required.")

h2(doc, "AI Enrichment — What GPT-5-mini Returns")
body(doc, "Input: business name + address + website (via web_search tool) + existing codes for context.")
add_table(doc,
    ["Output field", "Example", "Notes"],
    [
        ["naics_code", "722511", "6-digit NAICS 2022 only"],
        ["naics_description", "Full-Service Restaurants", "Canonical label"],
        ["uk_sic_code", "56101", "Only if business country = GB"],
        ["mcc_code", "5812", "4-digit MCC"],
        ["confidence", "HIGH", "HIGH / MED / LOW — self-reported text"],
        ["reasoning", "Website confirms restaurant...", "Explanation"],
        ["tools_used", '["web_search"]', "Which tools GPT used"],
    ]
)
body(doc, "Post-processing: validate against core_naics_code table. If invalid → 561499.")

h2(doc, "Pipeline A — Scenario Table")
add_table(doc,
    ["Situation", "What gets stored", "What customer sees"],
    [
        ["2–6 sources returned NAICS and agree",
         "Winning code in rds_warehouse_public.facts + data_businesses.naics_id updated",
         "Strong classification, confidence, all alternatives listed"],
        ["Sources disagree (different codes)",
         "Highest-confidence source wins",
         "Winner + alternatives showing the disagreement"],
        ["Only 1 source returned a NAICS",
         "That source wins; AI also runs (<3 sources)",
         "Code with possibly low confidence"],
        ["No source returned any NAICS",
         "AI triggers; stores GPT-generated code if valid",
         "AI-generated code OR 561499 placeholder"],
        ["Winning code not in core_naics_code",
         "⚠️ Replaced with 561499 via removeNaicsCode()",
         "Placeholder 561499 — signals unknown"],
        ["Very low confidence (e.g. 0.10)",
         "Highest available wins — NO minimum cutoff",
         "Code shown with low confidence visible"],
        ["Analyst manual override",
         "Override in facts with override:{...} field",
         "Override code with override field populated"],
        ["AI also fails",
         "⚠️ 561499 stored as absolute last resort",
         "Placeholder 561499 — analyst override required"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 4 — THE 8 CLASSIFICATION FACTS
# ══════════════════════════════════════════════════════════════════════════════

add_tab_header(doc, 4, "The 8 Classification Facts",
    "Where each fact comes from, what it contains, and where it is stored.")

callout(doc,
    "All 8 facts are stored in rds_warehouse_public.facts — a JSONB key-value store. "
    "One row per fact per business. Updated via Kafka facts.v1 topic. "
    "Upsert: INSERT ON CONFLICT (business_id, name) DO UPDATE.")

h2(doc, "Fact 1: naics_code")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "The single best 6-digit NAICS 2022 code for this business"],
        ["Winner priority", "Middesk (2.0) > OC (0.9) > ZoomInfo (0.8) > Trulioo (0.8) > Equifax (0.7) > AI (0.1)"],
        ["OC source field", "industry_code_uids → parse us_naics-XXXXXX entries"],
        ["EFX source field", "efx_primnaicscode"],
        ["ZI source field", "zi_c_naics6"],
        ["Storage", "rds_warehouse_public.facts (name='naics_code') + data_businesses.naics_id → FK to core_naics_code"],
        ["JSONB example", '{"code": "722511", "source": {"confidence": 0.95, "platformId": 16}, "alternatives": [...]}'],
        ["API endpoint", "GET /facts/business/{id}/details — full fact with confidence + alternatives"],
    ]
)

h2(doc, "Fact 2: naics_description")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "Human-readable label for the winning NAICS code"],
        ["Derived from", "Lookup from core_naics_code after naics_code resolves"],
        ["JSONB example", '{"description": "Full-Service Restaurants"}'],
        ["Storage", "rds_warehouse_public.facts only"],
    ]
)

h2(doc, "Fact 3: mcc_code")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "4-digit Merchant Category Code (Visa/Mastercard standard)"],
        ["Path A (best)", "Vendor directly returns MCC — Middesk or Trulioo"],
        ["Path B (most common)", "Crosswalk: SELECT mcc_code FROM rel_naics_mcc WHERE naics_code = '722511' → 5812"],
        ["Path C (fallback)", "GPT-5-mini returns mcc_code alongside naics_code"],
        ["JSONB example", '{"code": "5812", "description": "Eating Places, Restaurants"}'],
        ["Storage", "rds_warehouse_public.facts + data_businesses.mcc_id → FK to core_mcc_code"],
    ]
)

h2(doc, "Fact 4: mcc_description")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "Human-readable label for the MCC code"],
        ["Derived from", "Lookup from core_mcc_code after mcc_code resolves"],
        ["JSONB example", '{"description": "Eating Places, Restaurants"}'],
    ]
)

h2(doc, "Fact 5: mcc_code_found")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "Boolean — was MCC found directly from a vendor API?"],
        ["true", "Middesk or Trulioo directly returned the MCC code"],
        ["false", "MCC was derived from NAICS→MCC crosswalk (rel_naics_mcc)"],
        ["Why it matters", "Direct MCC is more reliable than crosswalk-derived MCC"],
        ["JSONB example", '{"value": true}'],
    ]
)

h2(doc, "Fact 6: mcc_code_from_naics")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "MCC code derived from the NAICS→MCC crosswalk table"],
        ["SQL used", "SELECT mcc_code FROM rel_naics_mcc WHERE naics_code = '722511' → 5812"],
        ["When used", "Always computed. Equals mcc_code when mcc_code_found = false"],
        ["JSONB example", '{"code": "5812"}'],
    ]
)

h2(doc, "Fact 7: industry")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "High-level category label for UI display (less precise than NAICS)"],
        ["Derived from", "First 2 digits of NAICS → core_business_industries lookup. "
         "Also from: ZI zi_c_industry, Trulioo industry text, OC labels"],
        ["Example", "NAICS 722511 → first 2 digits '72' → 'Food Service'"],
        ["JSONB example", '{"value": "Food Service"}'],
        ["Storage", "rds_warehouse_public.facts + data_businesses.industry → FK to core_business_industries"],
    ]
)

h2(doc, "Fact 8: classification_codes ⚠️ Known Gap")
callout(doc,
    "KNOWN GAP: classification_codes correctly captures UK SIC, NACE, and Canadian NAICS "
    "codes from OC's industry_code_uids. But NO Kafka handler, API endpoint, PDF report, "
    "or database column ever reads this fact. The UK SIC code exists in the system and is "
    "silently discarded.")
add_table(doc,
    ["Field", "Detail"],
    [
        ["What it is", "ALL raw industry codes from ALL sources before any winner is chosen"],
        ["Includes", "us_naics, gb_sic, uk_sic, nace, ca_naics from OC + all vendor codes"],
        ["JSONB example", '{"codes": [{"system":"NAICS","code":"722511","source":"Middesk","confidence":0.95}, '
                           '{"system":"uk_sic","code":"56101","source":"OC","confidence":0.89}]}'],
        ["Storage", "rds_warehouse_public.facts ONLY — NOT in data_businesses"],
        ["⚠️ Gap", "No Kafka handler, API, or report reads this fact. Currently unused downstream."],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — API & STORAGE
# ══════════════════════════════════════════════════════════════════════════════

add_tab_header(doc, 5, "API & Storage",
    "Where facts are stored, which API endpoints expose naics_code, and how to query Redshift.")

h2(doc, "Storage Architecture")
add_table(doc,
    ["Table", "Database", "What it contains", "Who writes it"],
    [
        ["rds_warehouse_public.facts", "PostgreSQL + Redshift mirror",
         "All 217 facts as JSONB key-value. One row per fact per business.",
         "warehouse-service FactService (Kafka consumer)"],
        ["rds_cases_public.data_businesses", "case-service PostgreSQL",
         "naics_id (FK), mcc_id (FK), industry (FK) — integer IDs not codes",
         "case-service Kafka handler (UPDATE_NAICS_CODE event)"],
        ["core_naics_code", "case-service PostgreSQL",
         "id, code, description — NAICS lookup table",
         "DB migrations (static data)"],
        ["core_mcc_code", "case-service PostgreSQL",
         "id, code, description — MCC lookup table", "DB migrations"],
        ["rel_naics_mcc", "case-service PostgreSQL",
         "NAICS → MCC crosswalk", "DB migrations"],
        ["core_business_industries", "case-service PostgreSQL",
         "Sector category labels (e.g. '72' → 'Food Service')", "DB migrations"],
        ["datascience.customer_files", "Redshift",
         "Pipeline B wide analytics table — primary_naics_code from ZI/EFX winner",
         "warehouse-service sp_recreate_customer_files()"],
    ]
)

h2(doc, "API Endpoints That Expose naics_code")
add_table(doc,
    ["Endpoint", "naics_code?", "Who can access", "What they see"],
    [
        ["GET /facts/business/{id}/details", "✅ YES — primary",
         "Admin, Customer, Applicant",
         "Full fact: code + confidence + alternatives[] from all sources"],
        ["GET /facts/business/{id}/kyb", "✅ YES — via industry",
         "Admin, Customer, Applicant",
         "naics_code as part of KYB facts"],
        ["GET /businesses/customers/{id}", "✅ YES — simplified",
         "Admin, Customer",
         "Flat fields: naics_code + naics_title"],
        ["GET /facts/business/{id}/all", "⚠️ Admin ONLY",
         "Admin only",
         "All 217 facts — 'intentionally not cached, leaks information'"],
        ["PATCH /businesses/{id}", "❌ Stripped from body",
         "n/a",
         "naics_code removed — cannot be set via PATCH"],
        ["PATCH /facts/.../override/naics_code", "✅ Write only",
         "Admin, Customer",
         "Sets analyst override — always wins future Fact Engine runs"],
    ]
)

h2(doc, "Redshift Queries")
h3(doc, "All 8 classification facts for one business")
code_block(doc,
    "SELECT name, value, received_at\n"
    "FROM dev.rds_warehouse_public.facts\n"
    "WHERE business_id = '<uuid>'\n"
    "  AND name IN (\n"
    "    'naics_code', 'naics_description',\n"
    "    'mcc_code', 'mcc_description',\n"
    "    'mcc_code_found', 'mcc_code_from_naics',\n"
    "    'industry', 'classification_codes'\n"
    "  )\n"
    "ORDER BY name;")

h3(doc, "Denormalized business record with decoded NAICS/MCC")
code_block(doc,
    "SELECT\n"
    "  b.id,\n"
    "  n.code  AS naics_code,\n"
    "  n.label AS naics_description,\n"
    "  m.code  AS mcc_code,\n"
    "  m.label AS mcc_description,\n"
    "  i.name  AS industry\n"
    "FROM dev.rds_cases_public.data_businesses b\n"
    "LEFT JOIN dev.rds_cases_public.core_naics_code n ON n.id = b.naics_id\n"
    "LEFT JOIN dev.rds_cases_public.core_mcc_code   m ON m.id = b.mcc_id\n"
    "LEFT JOIN dev.rds_cases_public.core_business_industries i ON i.id = b.industry\n"
    "WHERE b.id = '<uuid>';")

h3(doc, "Pipeline B — customer_files for analytics")
code_block(doc,
    "SELECT\n"
    "  business_id,\n"
    "  primary_naics_code,    -- winner from ZI or EFX\n"
    "  zi_match_confidence,   -- Pipeline B XGBoost score\n"
    "  efx_match_confidence,  -- Pipeline B XGBoost score\n"
    "  match_confidence       -- max(zi_match_confidence, efx_match_confidence)\n"
    "FROM datascience.customer_files\n"
    "WHERE business_id = '<uuid>';")

h2(doc, "Kafka Events for Classification Facts")
add_table(doc,
    ["Event", "Topic", "Triggered by", "Effect"],
    [
        ["CALCULATE_BUSINESS_FACTS", "facts.v1", "Business submitted",
         "Initial fact calculation for all 217 facts"],
        ["UPDATE_NAICS_CODE", "facts.v1", "naics_code fact changes",
         "case-service updates data_businesses.naics_id"],
        ["FACT_OVERRIDE_CREATED_AUDIT", "facts.v1", "Analyst sets override",
         "Audit trail updated, override wins all future runs"],
        ["PROCESS_COMPLETION_FACTS", "facts.v1", "All verifications complete",
         "Final fact recalculation with full context"],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/workspace/AI-Powered-NAICS-Industry-Classification-Agent/modeling/Industry_Classification_Worthopedia.docx"
doc.save(out_path)
print(f"\nSaved: {out_path}")
print("File size:", round(__import__('os').path.getsize(out_path)/1024, 1), "KB")
print("\nTab names to create in Google Docs:")
print("  Tab 1: Industry Classification")
print("  Tab 2: Pipeline B — Batch Redshift")
print("  Tab 3: Pipeline A — Real-Time Fact Engine")
print("  Tab 4: The 8 Classification Facts")
print("  Tab 5: API & Storage")
