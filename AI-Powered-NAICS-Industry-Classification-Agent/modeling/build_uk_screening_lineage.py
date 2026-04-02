"""
Build: Worth AI UK Screening — Complete End-to-End Lineage
Google-Docs-compatible .docx

Covers the full UK screening process end-to-end:
  • Trulioo KYB (business verification) — UK configuration, request payload, response
  • Watchlists, Sanctions, PEPs, Adverse Media — sources, extraction, storage
  • PSC screening (UBOs/Directors) — automatic UK flow
  • SERP / Trulioo routing for UK businesses
  • Fact engine transformation → API response → 360 Report mapping
  • Score/confidence/relevancy fields — where captured, whether surfaced
  • Real UK examples (Register of Insolvencies, Employment Tribunal, Accountant in Bankruptcy)
  • Consistency guarantees across request_response, facts, and API endpoints

Verified against:
  SIC-UK-Codes/integration-service/lib/trulioo/ (all files)
  SIC-UK-Codes/integration-service/lib/facts/kyb/
  SIC-UK-Codes/integration-service/src/messaging/kafka/consumers/handlers/report.ts
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

s = doc.sections[0]
s.page_width    = Inches(11.0)
s.page_height   = Inches(8.5)
s.left_margin   = Inches(0.7)
s.right_margin  = Inches(0.7)
s.top_margin    = Inches(0.6)
s.bottom_margin = Inches(0.6)
s.orientation   = 1
PAGE_W = 9.6

PURPLE = RGBColor(0x5B, 0x21, 0xB6)
BLUE   = RGBColor(0x1E, 0x40, 0xAF)
TEAL   = RGBColor(0x04, 0x5F, 0x7E)
GREEN  = RGBColor(0x06, 0x5F, 0x46)
RED    = RGBColor(0x99, 0x1B, 0x1B)
AMBER  = RGBColor(0x78, 0x35, 0x00)
SLATE  = RGBColor(0x33, 0x41, 0x55)
DARK   = RGBColor(0x0F, 0x17, 0x2A)

def _shade(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex6)
    tcPr.append(shd)

def _left_border(cell, hex6, sz=28):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    b = OxmlElement('w:left')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
    b.set(qn('w:space'),'0');    b.set(qn('w:color'),hex6)
    tcB.append(b); tcPr.append(tcB)

def _cell_margins(cell, top=60, right=100, bottom=60, left=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('right',right),('bottom',bottom),('left',left)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa'); tcMar.append(el)
    tcPr.append(tcMar)

def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); tblB.append(b)
    tblPr.append(tblB)

def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblH = OxmlElement('w:tblHeader'); trPr.append(tblH)

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(pts)

def _heading(level, text, colour, size, sb, sa, pb=False):
    if pb: doc.add_page_break()
    p = doc.add_heading('', level=level)
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = colour
    return p

def H1(t, c=PURPLE, pb=False): return _heading(1, t, c, 20, 16, 6, pb)
def H2(t, c=BLUE):             return _heading(2, t, c, 13.5, 13, 4)
def H3(t, c=TEAL):             return _heading(3, t, c, 11.5, 10, 3)
def H4(t, c=SLATE):            return _heading(4, t, c, 10.5,  8, 2)

def body(text, size=10.5, colour=SLATE, sa=5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = colour; r.bold = bold
    p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(0)
    return p

def bullet(prefix, text, size=10.5, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
    if prefix:
        rb = p.add_run(prefix); rb.bold = True; rb.font.size = Pt(size); rb.font.color.rgb = DARK
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = SLATE

def code_block(lines, size=8.5):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell,'F1F5F9'); cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=80, right=120, bottom=80, left=120)
    first = True
    for line in lines:
        cp = cell.paragraphs[0] if first else cell.add_paragraph(); first = False
        cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(size); r.font.color.rgb = DARK
    spacer(5)

def callout(text, bg='EFF6FF', border='1D4ED8', tc=None, size=10):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, bg); _left_border(cell, border, sz=28)
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=60, right=120, bottom=60, left=160)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(0)
    r = cp.add_run(text); r.font.size = Pt(size)
    r.font.color.rgb = tc or RGBColor.from_string(border)
    spacer(7)

def warn(text): callout(text, 'FFFBEB', 'D97706', RGBColor(0x78,0x35,0x00))
def gap(text):  callout(text, 'FEE2E2', 'DC2626', RGBColor(0x7F,0x1D,0x1D))
def ok(text):   callout(text, 'F0FDF4', '059669', RGBColor(0x06,0x5F,0x46))

def lineage(lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell,'F0F9FF'); _left_border(cell,'0284C7', sz=28)
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=80, right=120, bottom=80, left=140)
    first = True
    for line in lines:
        cp = cell.paragraphs[0] if first else cell.add_paragraph(); first = False
        cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(8.5); r.font.color.rgb = DARK
    spacer(8)

def table(headers, rows, col_widths=None, hdr='DBEAFE', alt='F8FAFF', fs=9.5):
    ncols = len(headers)
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = tbl.rows[0]; _repeat_header(hrow)
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]; _shade(cell, hdr); _cell_margins(cell)
        p = cell.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(fs); r.font.color.rgb = DARK
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1: _shade(cell, alt)
            _cell_margins(cell)
            p = cell.paragraphs[0]; v = str(val)
            if   v.startswith('YES') or v.startswith('✅'): rgb, bld = GREEN, True
            elif v.startswith('NO')  or v.startswith('❌'): rgb, bld = RED, True
            elif v.startswith('WARN') or v.startswith('⚠️'): rgb, bld = AMBER, True
            elif v.startswith('GAP') or v.startswith('🔴'):  rgb, bld = RED, True
            else: rgb, bld = SLATE, False
            r = p.add_run(v); r.font.size = Pt(fs); r.font.color.rgb = rgb; r.bold = bld
    if col_widths:
        for ri2 in range(len(tbl.rows)):
            for ci2, w in enumerate(col_widths): tbl.rows[ri2].cells[ci2].width = Inches(w)
    spacer(8)


# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
spacer(16)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Worth AI — UK Screening: Complete End-to-End Lineage')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = PURPLE
p.paragraph_format.space_after = Pt(5)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Trulioo KYB · Watchlists · Sanctions · PEPs · Adverse Media · PSC · SERP')
r.font.size = Pt(13); r.font.color.rgb = BLUE
p.paragraph_format.space_after = Pt(14)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Sources: SIC-UK-Codes/integration-service/lib/trulioo/** '
    '| lib/facts/kyb/ | handlers/report.ts\n'
    'Includes real UK example payloads from test fixtures and source code.'
)
r.font.size = Pt(10); r.font.color.rgb = SLATE
p.paragraph_format.space_after = Pt(16)

H2('Document Contents')
table(
    ['Section', 'Title', 'What it covers'],
    [
        ['1', 'UK vs US Routing — When Does Trulioo Run?',
         'Country-based routing decision, canIRun() logic, GB/UK normalisation, customer config'],
        ['2', 'Trulioo KYB: Configuration & Request Payload',
         'Environment variables, OAuth flow, KYB form fields, address strategy for UK businesses'],
        ['3', 'Trulioo KYB: Full Response Structure',
         'clientData, businessData, flowData, serviceData, fullServiceDetails, StandardizedLocations'],
        ['4', 'Watchlists & Sanctions — UK Sources',
         'Advanced Business Watchlist, OFAC, BIS, UK-specific lists (Register of Insolvencies etc.)'],
        ['5', 'Watchlist Extraction — Priority 1 vs Priority 2',
         'WatchlistHitDetails, WL_results/AM_results/PEP_results, createWatchlistHit(), score field'],
        ['6', 'Score / Confidence / Relevancy Fields',
         'What score means, where it comes from, whether it is surfaced in the API today'],
        ['7', 'PEP Screening — UK-Specific',
         'PEP_results, sourceRegion mapping, PEP database sources, PSC flow'],
        ['8', 'Adverse Media — UK Flow',
         'AM_results, Trulioo adverse media, OpenAI scoring, adverse_media_articles table'],
        ['9', 'PSC (Person of Significant Control) — UK UBO/Director Screening',
         'Auto-triggered PSC flow, extractPersonsFromBusinessData(), screenPersonWithPSCFlow()'],
        ['10', 'SERP — UK Behavior',
         'Does SERP run for UK businesses? Google Profile Verified/Unverified for UK'],
        ['11', 'Data Storage — request_response, business_entity_verification, facts',
         'What is saved, where, how each table is populated for UK businesses'],
        ['12', 'Fact Engine Transformation',
         'watchlist_raw, watchlist (consolidated), adverse_media_hits, screened_people facts'],
        ['13', 'API Responses — What Is Surfaced and Where',
         'GET /facts, GET /kyb, 360 Report watchlist section — consistency across endpoints'],
        ['14', '360 Report Watchlist Section',
         '_getWatchlistData(), createWatchlistEntries(), groupWatchlistHitsByEntityName()'],
        ['15', 'Consistency Guarantees & Known Gaps',
         'Where data may appear in more than one place, dedup logic, gaps in surfacing'],
        ['16', 'Full End-to-End Flow Diagram', 'ASCII lineage from submission to 360 Report'],
    ],
    col_widths=[0.5, 2.6, 6.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — UK vs US ROUTING
# ════════════════════════════════════════════════════════════════════════════

H1('Section 1 — UK vs US Routing: When Does Trulioo Run?', c=PURPLE)
callout(
    'Worth AI uses TWO different KYB providers:\n'
    '  US businesses  → Middesk (SOS Secretary of State API)\n'
    '  UK/Canada/international businesses → Trulioo (global KYB API)\n\n'
    'The decision is made per-business at submission time based on the business\'s '
    'country code and the customer\'s enabled countries configuration.'
)

H2('canIRun() — The Routing Decision Logic')
body(
    'truliooBusiness.ts implements canIRun() which runs before every KYB task. '
    'It checks three conditions in sequence:'
)
code_block([
    '// integration-service/lib/trulioo/business/truliooBusiness.ts',
    '',
    '// STEP 1: Is Trulioo enabled for this customer?',
    "const truliooEnabled = customer_integration_status.trulioo === 'ENABLED'",
    '// Stored in: integration_data.data_integration_settings.settings.trulioo.status',
    '',
    '// STEP 2: Which countries has the customer enabled for International KYB?',
    '// case-service API: GET /customers/{id}/onboarding-setups/7/countries',
    '// INTERNATIONAL_BUSINESS_SETUP_ID = 7',
    'const enabledCountries = customerCountries',
    '  .filter(c => c.is_selected || c.is_enabled)',
    '  .map(c => c.jurisdiction_code.toUpperCase())',
    '',
    '// STEP 3: GB/UK normalisation (system may store "UK", ISO code is "GB")',
    'const normalizedEnabledCountries = enabledCountries.map(code =>',
    '  code === "UK" ? "GB" : code',
    ')',
    'const normalizedBusinessCountry = businessCountry === "UK" ? "GB" : businessCountry',
    '',
    '// STEP 4: Is this business\'s country in the enabled list?',
    'const shouldRun = normalizedEnabledCountries.includes(normalizedBusinessCountry)',
    '',
    '// For UK business: shouldRun = true if "GB" or "UK" is in enabled countries',
    '// Result: uses Trulioo, skips Middesk',
])

H2('UK Business: Trulioo Runs, Middesk Does NOT')
table(
    ['Dimension', 'US business', 'UK business (GB)'],
    [
        ['KYB provider', 'Middesk (platform_id=16)',
         'Trulioo (platform_id=38)'],
        ['SOS filing lookup', 'YES — real Secretary of State data',
         'NO — Middesk only covers US states'],
        ['Registry data', 'State SOS filings',
         'Companies House (UK) via Trulioo data sources'],
        ['Address verification', 'Middesk returns StandardizedLocations for US addresses',
         'Trulioo returns StandardizedLocations for UK addresses (postcode-based)'],
        ['Watchlist screening', 'Trulioo optional (Advanced Watchlists toggle)',
         'Trulioo AUTOMATIC — part of standard KYB flow for non-US'],
        ['PSC/UBO screening', 'Only if Advanced Watchlists enabled',
         'AUTOMATIC — shouldScreenPSCsForBusiness() returns true for GB'],
        ['Confidence method', 'XGBoost confidenceScoreMany() (task-based)',
         'Trulioo: XGBoost entity-match OR heuristic fallback (see Section 6)'],
        ['NAICS code', 'From Middesk SOS filing (if available)',
         'From Trulioo StandardizedIndustries (sicCode, naicsCode) — often empty for UK'],
        ['UK SIC code', 'N/A',
         'From OC industry_code_uids (gb_sic- prefix) — NOT from Trulioo directly'],
    ],
    col_widths=[2.2, 3.2, 4.2],
)

H2('Configuration: Where UK Is Enabled')
table(
    ['Config location', 'Table', 'Field', 'What to check'],
    [
        ['Trulioo integration enabled',
         'integration_data.data_integration_settings',
         'settings JSONB -> trulioo.status = "ENABLED"',
         'GET /customers/{id}/integration-settings'],
        ['UK (GB) country enabled',
         'onboarding_schema.rel_customer_setup_countries\n'
         'Fetched via case-service',
         'jurisdiction_code = "GB" AND is_selected = true',
         'GET /customers/{id}/onboarding-setups/7/countries'],
        ['PSC screening for UK',
         'Derived — no separate toggle',
         'shouldScreenPSCsForBusiness() returns true for GB automatically\n'
         'when Trulioo is enabled (BTTF-205 scope)',
         'No config needed — automatic for non-US countries'],
        ['Advanced Watchlists (US only)',
         'integration_data.data_integration_settings',
         'settings JSONB -> advanced_watchlist.status = "ENABLED"',
         'Only required for US businesses\' PSC screening'],
    ],
    col_widths=[2.0, 2.5, 3.1, 2.0],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TRULIOO KYB CONFIGURATION & REQUEST
# ════════════════════════════════════════════════════════════════════════════

H1('Section 2 — Trulioo KYB: Configuration & Request Payload', c=PURPLE, pb=False)

H2('Environment Variables (trulioo README.md)')
code_block([
    '# Trulioo API endpoint (Global Data Company platform):',
    'TRULIOO_API_URL=https://api.globaldatacompany.com',
    '',
    '# OAuth 2.0 credentials (client_credentials grant):',
    'TRULIOO_CLIENT_ID=your_client_id',
    'TRULIOO_CLIENT_SECRET=your_client_secret',
    '',
    '# Flow IDs (configured in Trulioo portal per customer):',
    'TRULIOO_KYB_FLOWID=kyb-flow        # Business KYB flow ID',
    'TRULIOO_PSC_FLOWID=psc-screening-flow  # Person PSC screening flow ID',
    '',
    '# HTTP timeout:',
    'TRULIOO_HTTP_TIMEOUT=30000          # 30 seconds',
])

H2('Authentication: OAuth 2.0 Client Credentials')
code_block([
    '// truliooTokenManager.ts — OAuth flow:',
    'POST https://api.globaldatacompany.com/oauth2/v1/token',
    'Content-Type: application/x-www-form-urlencoded',
    '',
    'grant_type=client_credentials',
    '&client_id={TRULIOO_CLIENT_ID}',
    '&client_secret={TRULIOO_CLIENT_SECRET}',
    '',
    '// Response:',
    '{',
    '  "access_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...",',
    '  "expires_in": 3600,',
    '  "token_type": "Bearer",',
    '  "scope": "kyb psc"',
    '}',
    '',
    '// Token is cached and auto-refreshed before expiry.',
    '// All subsequent Trulioo API calls use: Authorization: Bearer {access_token}',
])

H2('KYB Flow Initiation — Two Steps')
body(
    'The Trulioo KYB process uses a dynamic form flow. '
    'Worth AI first fetches the flow structure (element IDs and roles), '
    'then maps business data to the flow fields and submits.'
)
code_block([
    '// STEP 1: Initialize flow — get field element IDs',
    'POST https://api.globaldatacompany.com/flows/{TRULIOO_KYB_FLOWID}/init',
    'Authorization: Bearer {access_token}',
    'Content-Type: application/json',
    '',
    '// Response contains elements[] with id and role:',
    '{',
    '  "elements": [',
    '    { "id": "elem_001", "role": "company_name" },',
    '    { "id": "elem_002", "role": "company_regno" },',
    '    { "id": "elem_003", "role": "company_country_incorporation" },',
    '    { "id": "elem_004", "role": "company_state_address" },',
    '    { "id": "elem_005", "role": "company_zip" },',
    '    { "id": "elem_006", "role": "company_address_full" }',
    '  ],',
    '  "hfSession": "session-uuid-abc123"',
    '}',
    '',
    '// STEP 2: Submit business data mapped to element IDs',
    'POST https://api.globaldatacompany.com/flows/{TRULIOO_KYB_FLOWID}/submit',
    'Authorization: Bearer {access_token}',
])

H2('UK Business KYB Request Payload — TruliooKYBFormData')
body(
    'After the flow init, Worth AI maps business data to form fields. '
    'Note the critical address strategy: Worth AI sends STREET-ONLY (no suite/apartment) '
    'to Trulioo so the Comprehensive View address match shows green on Address Line 1.'
)
code_block([
    '// From truliooBusinessKYBProcessor.ts:',
    '// For a UK business: "Acme Ltd, 123 Baker Street, London, W1U 6RG"',
    '',
    '// businessPayload (stored internally — full address with suite):',
    'const businessPayload: TruliooKYBFormData = {',
    '  companyName:                "Acme Limited",',
    '  companyCountryIncorporation: "GB",',
    '  companyStateAddress:         "England",',
    '  companyCity:                 "London",',
    '  companyZip:                  "W1U 6RG",',
    '  companyregno:                "12345678",   // UK Companies House number',
    '  companyAddressFull:          "123 Baker Street, Suite 2, London"  // full addr',
    '}',
    '',
    '// truliooPayload (sent to Trulioo — street only, no suite):',
    'const truliooPayload: TruliooKYBFormData = {',
    '  ...businessPayload,',
    '  companyAddressFull: "123 Baker Street"  // street only -> all-green Address Line 1',
    '}',
    '',
    '// IMPORTANT: The address split is intentional.',
    '// Trulioo\'s Comprehensive View compares street against Companies House data.',
    '// Including the suite number often causes nomatch on Address Line 1.',
    '// Worth AI strips it before sending so the registry street name matches.',
])

H2('TruliooKYBFormData Field Reference')
table(
    ['Field', 'UK value example', 'Required?', 'Source in Worth AI'],
    [
        ['companyName', '"Acme Limited"', 'YES', 'data_businesses.name'],
        ['companyCountryIncorporation', '"GB"', 'YES',
         'data_businesses.address_country (or inferred from address)'],
        ['companyStateAddress', '"England"', 'YES (may be empty string)',
         'data_businesses.address_state'],
        ['companyCity', '"London"', 'NO', 'data_businesses.address_city'],
        ['companyZip', '"W1U 6RG"', 'YES', 'data_businesses.address_postal_code'],
        ['companyState', '"England"', 'NO (duplicate of StateAddress)', 'Same as StateAddress'],
        ['companyregno', '"12345678"', 'NO but improves match',
         'data_businesses.tin (decrypted)\nOR registration_number field'],
        ['companyAddressFull', '"123 Baker Street"', 'NO',
         'address_line_1 only (no apartment/suite) — intentionally truncated'],
        ['companyEmail', '"info@acme.co.uk"', 'NO', 'Not sent currently'],
        ['companyVat', '"GB123456789"', 'NO', 'Not sent currently'],
    ],
    col_widths=[2.2, 1.8, 1.0, 4.6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TRULIOO KYB RESPONSE STRUCTURE
# ════════════════════════════════════════════════════════════════════════════

H1('Section 3 — Trulioo KYB: Full Response Structure', c=PURPLE, pb=False)

H2('Top-Level Response Shape (TruliooFlowResult)')
code_block([
    '// What Trulioo returns after flow submission:',
    '{',
    '  "hfSession": "session-uuid-abc123",   // session ID used for storage',
    '  "clientData": {',
    '    "status": "completed",              // "completed" | "pending" | "failed" | "REJECTED"',
    '    "external_id": "session-uuid-abc123",',
    '    "businessData": { ... },            // extracted business entity data',
    '    "watchlistResults": [ ... ],        // pre-processed watchlist hits (if any)',
    '    "flowData": {                        // raw form submission data',
    '      "kyb-flow-id": {',
    '        "serviceData": [',
    '          {',
    '            "nodeType": "trulioo_business_wl",  // or "trulioo_kyb", "trulioo_person_wl"',
    '            "fullServiceDetails": {',
    '              "Record": {',
    '                "DatasourceResults": [  // per-datasource verification results',
    '                  {',
    '                    "DatasourceName": "Comprehensive View",',
    '                    "DatasourceFields": [',
    '                      { "FieldName": "BusinessName",  "Status": "match" },',
    '                      { "FieldName": "AddressLine1",  "Status": "match" },',
    '                      { "FieldName": "PostalCode",    "Status": "match" }',
    '                    ]',
    '                  },',
    '                  {',
    '                    "DatasourceName": "Advanced Business Watchlist",',
    '                    "AppendedFields": [',
    '                      {',
    '                        "FieldName": "WatchlistHitDetails",',
    '                        "Data": {',
    '                          "WL_results": [ ... ],   // Sanctions hits',
    '                          "AM_results": [ ... ],   // Adverse Media hits',
    '                          "PEP_results": [ ... ]   // PEP hits',
    '                        }',
    '                      }',
    '                    ]',
    '                  }',
    '                ]',
    '              }',
    '            },',
    '            "watchlistResults": { ... }  // summary format (fallback)',
    '          }',
    '        ],',
    '        "fieldData": { ... }',
    '      }',
    '    }',
    '  }',
    '}',
])

H2('businessData — UK Entity Information Extracted')
code_block([
    '// TruliooBusinessData for a UK company:',
    '{',
    '  "name": "Acme Limited",',
    '  "country": "GB",',
    '  "state": "England",',
    '  "city": "London",',
    '  "postalCode": "W1U 6RG",',
    '  "business_addresses": [',
    '    {',
    '      "is_primary": true,',
    '      "addressLine1": "123 Baker Street",',
    '      "city": "London",',
    '      "state": "England",',
    '      "postalCode": "W1U 6RG",',
    '      "country": "GB"',
    '    }',
    '  ],',
    '  "ubos": [            // Persons with Significant Control found in Companies House',
    '    {',
    '      "fullName": "John Smith",',
    '      "firstName": "John",',
    '      "lastName": "Smith",',
    '      "dateOfBirth": "1980-05-15",',
    '      "ownershipPercentage": 75,',
    '      "title": "Director",',
    '      "nationality": "British"',
    '    }',
    '  ],',
    '  "directors": [       // Directors from Companies House filing',
    '    { "fullName": "Jane Doe", "title": "Secretary" }',
    '  ]',
    '}',
])

H2('StandardizedLocations — Reported Addresses for UK')
code_block([
    '// StandardizedLocations in Trulioo response (Companies House registry addresses):',
    '// Extracted by: extractTruliooAddressesAsStrings(businessData)',
    '{',
    '  "StandardizedLocations": [',
    '    {',
    '      "Address1": "123 Baker Street",',
    '      "BuildingNumber": "123",',
    '      "StreetName": "Baker Street",',
    '      "City": "London",',
    '      "StateProvinceCode": "ENG",',
    '      "PostalCode": "W1U 6RG",',
    '      "CountryCode": "GB",',
    '      "LocationType": "Registered"      // "Registered" | "Trading" | "Previous"',
    '    }',
    '  ]',
    '}',
    '',
    '// These become "Reported Addresses" in the Contact Information tab',
    '// submitted=false in integration_data.business_entity_address_source',
])

H2('Comprehensive View — DatasourceFields (Address Verification for UK)')
callout(
    'Trulioo\'s "Comprehensive View" datasource is the definitive address match result '
    'for UK businesses. It aggregates data from Companies House, Dun & Bradstreet UK, '
    'and other UK-specific data providers.\n\n'
    'Worth AI ONLY uses the Comprehensive View for match/nomatch decisions. '
    'Individual datasources (Business Insights, Business Essentials) are ignored '
    'because they may have incomplete UK data and report false "nomatch" results.'
)
code_block([
    '// DatasourceResults from Comprehensive View for a UK business:',
    '{',
    '  "DatasourceName": "Comprehensive View",',
    '  "DatasourceFields": [',
    '    { "FieldName": "BusinessName",       "Status": "match" },',
    '    { "FieldName": "AddressLine1",        "Status": "match" },',
    '    { "FieldName": "PostalCode",          "Status": "match" },',
    '    { "FieldName": "City",                "Status": "match" },',
    '    { "FieldName": "BusinessRegistrationNumber", "Status": "match" },',
    '    { "FieldName": "DayOfBirth",          "Status": "missing" },',
    '    { "FieldName": "MonthOfBirth",        "Status": "missing" }',
    '  ]',
    '}',
    '',
    '// Status values: "match" | "nomatch" | "missing"',
    '// "missing" = field not in Trulioo\'s data for this entity (NOT a mismatch)',
    '// ADDRESS RULE (extractAddressMatchStatusFromDatasourceFields):',
    '//   ANY "nomatch" on AddressLine1/PostalCode/City -> overall "nomatch" -> red No Match badge',
    '//   All "match" or "missing"                      -> overall "match"  -> green Verified badge',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WATCHLISTS & SANCTIONS UK SOURCES
# ════════════════════════════════════════════════════════════════════════════

H1('Section 4 — Watchlists & Sanctions: UK-Specific Sources', c=PURPLE, pb=False)

callout(
    'For UK businesses, Trulioo screens against the "Advanced Business Watchlist" — '
    'a single consolidated datasource that checks hundreds of global lists simultaneously. '
    'Results come back in three buckets:\n'
    '  WL_results   = Watchlist / Sanctions hits\n'
    '  AM_results   = Adverse Media hits\n'
    '  PEP_results  = Politically Exposed Persons hits'
)

H2('UK-Specific Watchlists (sourceListType examples from test fixtures)')
table(
    ['sourceListType (UK-specific)', 'listType bucket', 'sourceAgencyName', 'Description'],
    [
        ['Register of Insolvencies', 'SANCTIONS (WL_results)',
         'Accountant in Bankruptcy', 'Scottish insolvency register (Accountant in Bankruptcy Scotland)'],
        ['Employment Tribunal Decisions', 'SANCTIONS (WL_results)',
         'Employment Tribunal', 'UK Employment Tribunal decisions — adverse rulings'],
        ['Companies House Disqualified Directors', 'SANCTIONS (WL_results)',
         'Companies House UK', 'Directors barred from company directorships in UK'],
        ['Insolvency Register', 'SANCTIONS (WL_results)',
         'Insolvency Service UK', 'England & Wales insolvency register (bankruptcies, IVAs, DROs)'],
        ['Financial Sanctions List', 'SANCTIONS (WL_results)',
         'OFSI (Office of Financial Sanctions Implementation)',
         'UK HM Treasury sanctions — UK equivalent of OFAC SDN for post-Brexit UK'],
        ['PEP Database', 'PEP (PEP_results)',
         'Various (Politically Exposed Persons)',
         'UK-specific PEP data including MPs, Lords, civil servants, local officials'],
        ['Adverse Media', 'ADVERSE_MEDIA (AM_results)',
         'Trulioo (aggregated news sources)',
         'UK news sources: BBC, The Guardian, Financial Times, Reuters, etc.'],
        ['Specially Designated Nationals (SDN)', 'SANCTIONS (WL_results)',
         'Office of Foreign Assets Control (OFAC)',
         'US OFAC list — applies globally including UK businesses'],
        ['Entity List', 'SANCTIONS (WL_results)',
         'Bureau of Industry and Security (BIS)',
         'US BIS export restrictions — applies to UK entities too'],
        ['EU Consolidated Sanctions List', 'SANCTIONS (WL_results)',
         'European External Action Service',
         'Post-Brexit: EU list is separate from UK OFSI list but both checked'],
        ['UN Consolidated List', 'SANCTIONS (WL_results)',
         'United Nations Security Council',
         'Global UN sanctions — included for all jurisdictions'],
    ],
    col_widths=[2.8, 1.8, 2.2, 3.0],
)

H2('Global Lists Also Applied to UK Businesses')
table(
    ['Agency', 'Lists', 'Region'],
    [
        ['OFAC', 'SDN, SSI, Non-SDN CMIC, Non-SDN Palestinian LC,\nCAPTA Foreign Sanctions Evaders, Non-SDN Iran',
         'North America / Global'],
        ['BIS', 'Entity List, Denied Persons List, Unverified List', 'North America / Global'],
        ['EU EEAS', 'EU Consolidated Sanctions', 'European Union'],
        ['UN Security Council', 'UN Consolidated List', 'Global'],
        ['HM Treasury (UK)', 'UK Financial Sanctions (OFSI)', 'United Kingdom'],
        ['Insolvency Service', 'England & Wales Insolvency Register', 'United Kingdom'],
        ['Accountant in Bankruptcy', 'Scottish Register of Insolvencies', 'Scotland / United Kingdom'],
        ['Companies House', 'Disqualified Directors Register', 'United Kingdom'],
        ['Various news sources', 'Adverse Media (aggregated)', 'Global (UK-relevant)'],
        ['Various PEP databases', 'Politically Exposed Persons', 'Global (UK MPs, Lords, etc.)'],
    ],
    col_widths=[2.5, 4.2, 2.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — WATCHLIST EXTRACTION PRIORITY 1 vs 2
# ════════════════════════════════════════════════════════════════════════════

H1('Section 5 — Watchlist Extraction: Priority 1 vs Priority 2', c=PURPLE, pb=False)
callout(
    'extractWatchlistResultsFromTruliooResponse() implements a two-priority extraction strategy.\n\n'
    'PRIORITY 1 (preferred): fullServiceDetails.Record.DatasourceResults.AppendedFields.WatchlistHitDetails\n'
    '  → Contains individual hit-level detail: entity name, score, sourceAgencyName, URL, sourceListType\n\n'
    'PRIORITY 2 (fallback): watchlistResults summary format\n'
    '  → Contains only aggregate counts: wlHitsNumber, amHitsNumber, pepHitsNumber\n'
    '  → Creates placeholder entries — less detail, no entity-level match information'
)

H2('Priority 1 — Detailed Extraction from fullServiceDetails')
H3('Real UK example from test fixtures')
code_block([
    '// Real Trulioo response structure for a UK person/business watchlist hit:',
    '// (from extractWatchlistResults.test.ts)',
    '{',
    '  "flowData": {',
    '    "flow1": {',
    '      "serviceData": [',
    '        {',
    '          "nodeType": "trulioo_person_wl",  // or "trulioo_business_wl"',
    '          "fullServiceDetails": {',
    '            "Record": {',
    '              "DatasourceResults": [',
    '                {',
    '                  "DatasourceName": "Advanced Watchlist",',
    '                  "AppendedFields": [',
    '                    {',
    '                      "FieldName": "WatchlistHitDetails",',
    '                      "Data": {',
    '                        "WL_results": [',
    '                          {',
    '                            "score": 1,                           // relevancy: 0.0 to 1.0',
    '                            "subjectMatched": "Jacqueline Martin",// matched entity name',
    '                            "sourceListType": "Register of Insolvencies",',
    '                            "sourceRegion": "Europe",',
    '                            "sourceAgencyName": "Accountant in Bankruptcy",',
    '                            "URL": "https://example.com/hit1"',
    '                          },',
    '                          {',
    '                            "score": 0.98,',
    '                            "subjectMatched": "Martin Jacqueline",// name permutation',
    '                            "sourceListType": "Employment Tribunal Decisions",',
    '                            "sourceRegion": "Europe",',
    '                            "sourceAgencyName": "Employment Tribunal",',
    '                            "URL": "https://example.com/hit2"',
    '                          }',
    '                        ],',
    '                        "AM_results": [',
    '                          {',
    '                            "score": 0.95,',
    '                            "subjectMatched": "Jacqueline Martin",',
    '                            "sourceListType": "Press Releases",',
    '                            "sourceRegion": "North America",',
    '                            "sourceAgencyName": "FBI"',
    '                          }',
    '                        ],',
    '                        "PEP_results": [',
    '                          {',
    '                            "score": 1,',
    '                            "subjectMatched": "Jane Smith",',
    '                            "sourceListType": "PEP Database",',
    '                            "sourceRegion": "Asia"',
    '                          }',
    '                        ]',
    '                      }',
    '                    }',
    '                  ]',
    '                }',
    '              ]',
    '            }',
    '          }',
    '        }',
    '      ]',
    '    }',
    '  }',
    '}',
])

H3('createWatchlistHit() — What gets stored per hit')
code_block([
    '// integration-service/lib/trulioo/common/utils.ts: createWatchlistHit()',
    '// Called for each WL_results / AM_results / PEP_results entry:',
    '',
    'function createWatchlistHit(hit, listType, defaultListName): TruliooWatchlistHit {',
    '  return {',
    '    listType,                                 // "SANCTIONS" | "ADVERSE_MEDIA" | "PEP"',
    '    listName: hit.sourceListType || defaultListName,   // "Register of Insolvencies"',
    '    confidence: Number(hit.score) || 0,       // THE SCORE FIELD: 0.0 to 1.0',
    '    matchDetails: hit.subjectMatched           // matched entity name',
    '                || hit.entityName',
    '                || hit.remarks',
    '                || `${listType} hit`,',
    '    url: hit.URL || undefined,                // source document URL',
    '    sourceAgencyName: hit.sourceAgencyName,   // "Accountant in Bankruptcy"',
    '    sourceRegion: hit.sourceRegion,           // "Europe"',
    '    sourceListType: hit.sourceListType,       // "Register of Insolvencies"',
    '    listCountry: mapRegionToCountry(hit.sourceRegion)  // "Europe" -> "Europe"',
    '  }',
    '}',
    '',
    '// For the UK example above, WL_results[0] produces:',
    '{',
    '  listType: "SANCTIONS",',
    '  listName: "Register of Insolvencies",',
    '  confidence: 1.0,                  // score = 1.0 (exact match)',
    '  matchDetails: "Jacqueline Martin",',
    '  url: "https://example.com/hit1",',
    '  sourceAgencyName: "Accountant in Bankruptcy",',
    '  sourceRegion: "Europe",',
    '  sourceListType: "Register of Insolvencies",',
    '  listCountry: "Europe"',
    '}',
])

H2('Priority 2 — Summary Format Fallback')
code_block([
    '// When fullServiceDetails is absent or WatchlistHitDetails is empty,',
    '// extraction falls back to the summary watchlistResults format:',
    '{',
    '  "watchlistResults": {',
    '    "Advanced Watchlist": {',
    '      "watchlistStatus": "Potential Hit",',
    '      "watchlistHitDetails": {',
    '        "wlHitsNumber": 138,    // count of WL/Sanctions hits',
    '        "amHitsNumber": 12,     // count of Adverse Media hits',
    '        "pepHitsNumber": 3,     // count of PEP hits',
    '        "confidence": 0.85     // aggregate confidence (less useful)',
    '      }',
    '    }',
    '  }',
    '}',
    '',
    '// This creates placeholder TruliooWatchlistHit entries:',
    '// 138 × { listType: "SANCTIONS", listName: "Advanced Watchlist", confidence: 0.85,',
    '//         matchDetails: "Watchlist hit 1 of 138" }  <- no entity name, no URL',
    '',
    '// ⚠️ Priority 2 hits lack entity names, URLs, and sourceAgencyName.',
    '// This is a known limitation when Trulioo returns summary-only responses.',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — SCORE / CONFIDENCE / RELEVANCY
# ════════════════════════════════════════════════════════════════════════════

H1('Section 6 — Score / Confidence / Relevancy Fields', c=PURPLE, pb=False)

callout(
    'There are TWO different score/confidence concepts in the Trulioo UK pipeline:\n\n'
    '1. Watchlist match score (hit.score / TruliooWatchlistHit.confidence)\n'
    '   The per-hit relevancy score returned by Trulioo (0.0 to 1.0).\n'
    '   1.0 = exact name match. 0.85 = close fuzzy match.\n\n'
    '2. Source confidence (Trulioo\'s overall match quality as a KYB source)\n'
    '   Used by the Fact Engine to decide whether Trulioo wins industry/firmographic facts.\n'
    '   Computed differently — see below.'
)

H2('1. Watchlist Hit Score (hit.score / confidence field)')
table(
    ['Aspect', 'Detail'],
    [
        ['What it is', 'Trulioo\'s relevancy score for how well the entity name matches the watchlist entry.\n'
         '1.0 = exact match. 0.95 = very close. 0.7 = fuzzy match.'],
        ['Where it comes from', 'Trulioo fullServiceDetails.WatchlistHitDetails.WL_results[].score\n'
         '(or AM_results[].score, PEP_results[].score)'],
        ['How it is stored', 'TruliooWatchlistHit.confidence field.\n'
         'Stored in: integration_data.business_entity_review_task.metadata JSONB\n'
         'as WatchlistValueMetadatum[] — but the "score" field is NOT included in WatchlistValueMetadatum.'],
        ['Is it surfaced in the API?',
         'PARTIALLY. The watchlist fact schema includes an optional "score" field:\n'
         '  score: z.number().optional()  (kyb/index.ts line 1527)\n'
         'BUT: transformTruliooBusinessWatchlistResults() and mapWatchlistHits() do NOT\n'
         'currently copy the confidence/score into WatchlistValueMetadatum.\n'
         'Result: score IS captured in request_response JSONB but NOT propagated to the\n'
         'watchlist fact or the API response /facts/business/{id}/details.'],
        ['GAP: score not surfaced',
         'The confidence (0.0-1.0) from Trulioo is LOST between:\n'
         '  TruliooWatchlistHit.confidence (has it)\n'
         '  → mapWatchlistHits() / storeBusinessWatchlistResults() (does NOT copy it to DB)\n'
         '  → WatchlistValueMetadatum (no confidence field)\n'
         '  → watchlist fact (schema has "score" but it is never populated from Trulioo)'],
    ],
    col_widths=[2.2, 7.4],
)

H2('2. Source Confidence — Trulioo as a KYB Fact Source')
body(
    'This is different from the watchlist hit score. This is how confident Worth AI is '
    'that Trulioo found the right business entity (for NAICS, firmographics, address matching). '
    'Used by the Fact Engine to decide if Trulioo\'s data wins over ZoomInfo/Equifax/OC.'
)
code_block([
    '// integration-service/lib/facts/sources.ts:',
    '// Trulioo source confidence computation:',
    '',
    '// PRIMARY: XGBoost entity matching (if both customerBusiness and truliooBusiness available)',
    'const confidenceResult = await confidenceScore({',
    '  business: customerBusiness,         // submitted business (normalised)',
    '  integration_business: truliooBusiness  // Trulioo business data (normalised)',
    '})',
    'this.confidence = confidenceResult.prediction  // 0.0 to 1.0 from XGBoost',
    '',
    '// FALLBACK: Heuristic based on verification status',
    'if (!this.confidence) {',
    '  if (status === "completed" || status === "success")  calculatedConfidence = 0.7',
    '  if (status === "pending"  || status === "in_progress") calculatedConfidence = 0.4',
    '  if (status === "failed"   || status === "REJECTED")    calculatedConfidence = 0.2',
    '  else calculatedConfidence = 0.3  // unknown status',
    '',
    '  // Bonus for data completeness:',
    '  if (businessData.name)                     calculatedConfidence += 0.1',
    '  if (businessData.address || businessData.business_addresses?.length > 0)',
    '                                             calculatedConfidence += 0.1',
    '  if (businessData.ubos?.length > 0 || businessData.directors?.length > 0)',
    '                                             calculatedConfidence += 0.05',
    '}',
    '',
    '// This confidence is stored in:',
    '//   integration_data.request_response.confidence (column)',
    '// And used by Fact Engine: factWithHighestConfidence() to pick winning source',
    '// Surfaced in API: GET /facts/business/{id}/details -> source.confidence',
])

H2('Summary: Score/Confidence Field Status')
table(
    ['Field', 'Captured?', 'Stored where?', 'Surfaced in API?'],
    [
        ['Watchlist hit score\n(hit.score from WL_results/AM_results/PEP_results)',
         'YES — in TruliooWatchlistHit.confidence in memory',
         'Stored in request_response.response JSONB (full Trulioo response)\n'
         'NOT in business_entity_review_task.metadata (WatchlistValueMetadatum has no confidence field)',
         'NO — NOT exposed in /facts/ or 360 report. watchlist fact schema has\n'
         '"score: z.number().optional()" but it is never populated.'],
        ['Trulioo source confidence\n(entity match quality for KYB facts)',
         'YES — computed by XGBoost or heuristic',
         'integration_data.request_response.confidence (dedicated column)',
         'YES — GET /facts/business/{id}/details returns source.confidence'],
        ['Watchlist aggregate confidence\n(Priority 2 summary format)',
         'YES — hits.confidence from watchlistHitDetails',
         'Used during extraction but not separately stored',
         'NO — only used to create placeholder hit entries'],
    ],
    col_widths=[2.8, 1.5, 3.3, 2.0],
)

gap(
    'GAP: The per-hit Trulioo watchlist score (0.0-1.0) indicating match relevancy '
    'is captured in the raw response but never surfaced to API consumers or the 360 Report. '
    'The WatchlistValueMetadatum schema lacks a confidence/score field. '
    'To surface it: add "score?: number" to WatchlistValueMetadatum, '
    'copy TruliooWatchlistHit.confidence into it in mapWatchlistHits(), '
    'and include it in the watchlist fact calculation.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — PEP SCREENING UK
# ════════════════════════════════════════════════════════════════════════════

H1('Section 7 — PEP Screening: UK-Specific', c=PURPLE, pb=False)

H2('What is a PEP in the UK Context?')
body(
    'PEP = Politically Exposed Person. Under UK AML regulations (JMLSG guidance and '
    'the Money Laundering Regulations 2017), a PEP is defined as:\n'
)
bullet('Current or former', ' senior politicians (MPs, Lords, Devolved Assembly Members)')
bullet('Government ministers', ' and senior civil servants')
bullet('Senior judicial officers', ' (High Court judges and above)')
bullet('Ambassadors and', ' high commissioners')
bullet('Senior executives', ' of state-owned enterprises')
bullet('Members of governing bodies', ' of international organisations (UN, NATO, etc.)')
bullet('Immediate family members', ' and known close associates of any of the above')
spacer(6)
callout(
    'Worth AI checks PEPs for BOTH the business (via KYB flow) AND for each UBO/Director '
    '(via PSC flow). PEP hits appear in PEP_results[] within WatchlistHitDetails. '
    'For UK businesses, Trulioo uses UK-specific PEP databases alongside global PEP databases.'
)

H2('PEP_results — UK Example')
code_block([
    '// Real test fixture format for a PEP hit:',
    '"PEP_results": [',
    '  {',
    '    "score": 1,                        // exact match',
    '    "subjectMatched": "Jane Smith",    // matched name',
    '    "sourceListType": "PEP Database",  // "UK PEP Database" | "Global PEP" | etc.',
    '    "sourceRegion": "Europe"           // maps to listCountry = "Europe"',
    '  }',
    ']',
    '',
    '// createWatchlistHit() produces:',
    '{',
    '  listType: "PEP",',
    '  listName: "PEP Database",',
    '  confidence: 1.0,',
    '  matchDetails: "Jane Smith",',
    '  sourceRegion: "Europe",',
    '  listCountry: "Europe"',
    '}',
    '',
    '// mapAgencyFromListType() for PEP:',
    '// if listType === "PEP" OR n.includes("pep") OR n.includes("politically exposed")',
    '// -> { agency: "Politically Exposed Persons", agencyAbbr: "PEP" }',
])

H2('mapAgencyFromListType() — UK Agency Mapping')
code_block([
    '// integration-service/lib/trulioo/business/truliooWatchlist.ts',
    '// Maps Trulioo listType + listName to human-readable agency:',
    '',
    'const patterns = {',
    '  PEP:          { agency: "Politically Exposed Persons", agencyAbbr: "PEP" },',
    '  ADVERSE_MEDIA:{ agency: "Adverse Media", agencyAbbr: "ADVERSE" },',
    '  SANCTIONS:    { agency: "Sanctions List", agencyAbbr: "SANCTIONS" },',
    '  ofac: { agency: "Office of Foreign Assets Control", agencyAbbr: "OFAC" },',
    '  bis:  { agency: "Bureau of Industry and Security", agencyAbbr: "BIS" },',
    '  dos:  { agency: "State Department", agencyAbbr: "DOS" }',
    '}',
    '',
    '// Matching logic (case-insensitive on listName):',
    '// "register of insolvencies"   -> patterns.SANCTIONS (no OFAC/BIS match -> SANCTIONS)',
    '// "employment tribunal"         -> patterns.SANCTIONS',
    '// "financial sanctions"         -> patterns.SANCTIONS (matches "sanction")',
    '// "specially designated nationals" -> patterns.ofac   (matches "sdn")',
    '// "entity list"                 -> patterns.bis       (matches "entity list")',
    '// "pep database"                -> patterns.PEP       (matches "pep")',
    '',
    '// NOTE: UK-specific lists (Register of Insolvencies, Employment Tribunal)',
    '// fall into the generic SANCTIONS bucket because there is no UK-specific pattern.',
    '// The sourceAgencyName field preserves the real agency: "Accountant in Bankruptcy".',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — ADVERSE MEDIA UK FLOW
# ════════════════════════════════════════════════════════════════════════════

H1('Section 8 — Adverse Media: UK Flow', c=PURPLE, pb=False)

H2('Two Sources of Adverse Media for UK Businesses')
table(
    ['Source', 'Platform', 'Trigger', 'What it produces'],
    [
        ['Trulioo KYB AM_results',
         'Trulioo (platform_id=38)',
         'Automatic — part of KYB flow when watchlist screening runs',
         'News article URLs from Trulioo\'s adverse media data sources.\n'
         'listType="ADVERSE_MEDIA" in WatchlistHitDetails.AM_results.'],
        ['SERP Adverse Media',
         'SERP (platform_id=22)',
         'Separate integration — if enabled and configured for UK',
         'Google Search results for "[business name] fraud/scam/lawsuit".\n'
         'Stored in adverse_media + adverse_media_articles tables.'],
    ],
    col_widths=[1.8, 1.4, 2.4, 4.0],
)

H2('Trulioo Adverse Media Processing Pipeline')
lineage([
    'Trulioo KYB response contains:',
    '  fullServiceDetails.WatchlistHitDetails.AM_results[]:',
    '    { score, subjectMatched, sourceListType, URL, sourceAgencyName }',
    '       |',
    '  extractWatchlistResultsFromTruliooResponse() extracts AM_results as:',
    '    TruliooWatchlistHit { listType: "ADVERSE_MEDIA", listName, confidence, url, ... }',
    '       |',
    '  truliooBusinessResultsStorage.storeBusinessVerificationResults():',
    '    storeBusinessWatchlistResults() -> stores in business_entity_review_task',
    '    processAndPersistTruliooAdverseMedia() -> post-processing for scoring',
    '       |',
    '  processAndPersistTruliooAdverseMedia():',
    '    For each AM hit:',
    '      1. Extract title from hit.listName (or URL slug if generic "Adverse Media")',
    '         extractTitleFromUrl(hit.url) -> derives title from URL slug',
    '         e.g. "https://bbc.co.uk/news/business-acme-ltd-fraud" -> "Acme ltd fraud"',
    '      2. scoreAdverseMedia(title, entityNames, individuals) via OpenAI API:',
    '         -> keywordsScore (0-1): fraud/crime/sanction keywords',
    '         -> negativeSentimentScore (0-1): sentiment analysis',
    '         -> entityFocusScore (0-1): is the article about this entity?',
    '         -> finalScore (0-1): weighted combination',
    '         -> riskLevel: "LOW" | "MEDIUM" | "HIGH"',
    '         -> description: human-readable risk explanation',
    '         -> mediaType: "business" | individual name',
    '      3. Deduplicate by (link, mediaType)',
    '      4. insertAdverseMedia() ->',
    '           adverse_media table (business_id, total/high/medium/low risk counts)',
    '           adverse_media_articles table (per article: title, link, scores, riskLevel)',
])

H2('adverse_media_articles Table — What is Stored')
code_block([
    '// adverse_media_articles table (case-service PostgreSQL):',
    '{',
    '  business_id: "<uuid>",',
    '  title: "Acme Ltd fraud investigation",  // from URL slug or listName',
    '  link: "https://bbc.co.uk/news/...",      // Trulioo hit URL',
    '  date: "2026-03-27T10:00:00Z",            // processing timestamp',
    '  source: "bbc.co.uk",                     // extracted from URL domain',
    '  keywords_score: 0.92,                    // OpenAI: fraud/crime keyword score',
    '  negative_sentiment_score: 0.88,          // OpenAI: negative sentiment',
    '  entity_focus_score: 0.95,               // OpenAI: entity relevance',
    '  final_score: 0.91,                      // OpenAI: combined risk score',
    '  risk_level: "HIGH",                     // "LOW" | "MEDIUM" | "HIGH"',
    '  risk_description: "Article mentions specific fraud allegation against entity",',
    '  media_type: "business"                   // "business" | owner name',
    '}',
])

H2('When Trulioo Returns Generic "Adverse Media" as listName')
callout(
    'Trulioo often returns listName = "Adverse Media" (generic) without a meaningful title. '
    'Worth AI extracts the real article title from the URL slug:\n\n'
    'URL: https://www.bbc.co.uk/news/business-50236936-acme-ltd-fraud-case\n'
    'slug: "business-50236936-acme-ltd-fraud-case"\n'
    'after cleaning: "Acme ltd fraud case"\n\n'
    'extractTitleFromUrl() strips: file extensions, numeric IDs (5+ digits), hyphens → spaces.\n'
    'Returns null if slug < 10 chars (too short to be meaningful).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — PSC SCREENING (UK UBOs/DIRECTORS)
# ════════════════════════════════════════════════════════════════════════════

H1('Section 9 — PSC Screening: UK UBO/Director Automatic Flow', c=PURPLE, pb=False)
callout(
    'For UK businesses, PSC (Person of Significant Control) screening is AUTOMATIC. '
    'It is triggered immediately after KYB verification completes with status="approved". '
    'No separate API call is needed — it follows the Middesk pattern: '
    '"one API call → business verification + automatic person screening".'
)

H2('When PSC Screening Triggers for UK')
code_block([
    '// pscScreeningHelpers.ts: shouldScreenPSCsForBusiness()',
    '',
    '// For UK (GB) business:',
    'const isUS = false  // GB is not US',
    '',
    '// Prerequisite: Trulioo (International KYB) must be enabled',
    'const internationalKYBEnabled = await checkInternationalKYBEnabled(customerId)',
    '// -> checks integration_data.data_integration_settings: trulioo.status = "ENABLED"',
    '',
    '// If enabled AND not US -> PSC screening is AUTOMATIC:',
    'return { shouldScreen: true, reason: "Non-US business with International KYB enabled" }',
    '',
    '// UK businesses do NOT need the "Advanced Watchlists" toggle.',
    '// Advanced Watchlists is a separate toggle ONLY for US businesses.',
    '',
    '// PSC flow is skipped if:',
    '// - Trulioo is not enabled for this customer',
    '// - Business country is not in enabled countries list (canIRun() check)',
    '// - KYB verification did not succeed (status != "approved")',
])

H2('Person Sources: Who Gets Screened for a UK Business?')
body(
    'extractPersonsFromBusinessData() gathers persons from MULTIPLE sources and deduplicates:'
)
table(
    ['Source', 'How persons are found', 'deduplication'],
    [
        ['Trulioo KYB response — UBOs',
         'businessData.ubos[] — Persons with Significant Control from Companies House\n'
         'TruliooUBO: { fullName, firstName, lastName, dateOfBirth, ownershipPercentage, title }',
         'Deduped by fullName (case-insensitive)'],
        ['Trulioo KYB response — Directors',
         'businessData.directors[] — Directors from Companies House filing\n'
         'TruliooDirector: { fullName, firstName, lastName, title, nationality }',
         'Deduped by fullName (case-insensitive)'],
        ['Applicant-submitted owners',
         'getOwnersUnencrypted(businessId) from case-service\n'
         'data_owners table: owners submitted via onboarding form',
         'convertOwnersToTruliooPersons() normalises to TruliooUBOPersonData'],
        ['Middesk-discovered officers (US only)',
         'integration_data.business_entity_people WHERE submitted=false\n'
         'Officers discovered by Middesk from US SOS filings',
         'NOT included for UK businesses (this path only runs for US businesses\n'
         'with Advanced Watchlists enabled)'],
    ],
    col_widths=[2.2, 4.3, 3.1],
)

H2('PSC Screening Request — TruliooUBOPersonData')
code_block([
    '// For each UK UBO/Director, the PSC flow sends:',
    '{',
    '  fullName:             "John Smith",',
    '  firstName:            "John",',
    '  lastName:             "Smith",',
    '  dateOfBirth:          "1980-05-15",  // YYYY-MM-DD',
    '  controlType:          "UBO",          // "UBO" | "DIRECTOR" | "BENEFICIARY" | "CONTROL"',
    '  ownershipPercentage:  75,',
    '  title:                "Director",',
    '  nationality:          "British",',
    '  // Address from businessData (fallback to business address if no personal address)',
    '  addressLine1:         "123 Baker Street",',
    '  city:                 "London",',
    '  postalCode:           "W1U 6RG",',
    '  country:              "GB"',
    '}',
])

H2('PSC Screening Results Storage')
code_block([
    '// truliooPersonDataStorage.ts: storePersonInBusinessEntityPeople()',
    'INSERT INTO integration_data.business_entity_people',
    '  (business_entity_verification_id, name, submitted, source, titles, screening_results)',
    'VALUES',
    '  (',
    '    "<kyb_verification_uuid>",',
    '    "John Smith",',
    '    true,',
    '    \'[{"type": "trulioo_psc", "provider": "trulioo", "id": "John Smith", "controlType": "UBO"}]\',',
    '    \'["Director"]\',',
    '    \'{"watchlistHits": [...], "provider": "trulioo", "screenedAt": "2026-03-27T10:00:00Z"}\'',
    '  )',
    'ON CONFLICT (business_entity_verification_id, name) DO IGNORE  -- atomic lock',
    '',
    '// The ON CONFLICT ignore is a critical atomic lock:',
    '// Prevents duplicate PSC screenings if two processes race',
    '// Only ONE process successfully inserts -> it triggers PSC',
    '// Others see the existing row and skip',
])

H2('PSC Watchlist Hits — Stored in business_entity_review_task')
code_block([
    '// PSC screening watchlist hits are stored with entity_type = PERSON:',
    'WatchlistValueMetadatum {',
    '  id: "<uuid>",',
    '  type: "pep",            // "sanctions" | "pep" | "adverse_media" | "other"',
    '  entity_type: "person",  // ← distinguishes PSC hits from KYB business hits',
    '  metadata: {',
    '    abbr: "PEP",',
    '    title: "PEP Database",',
    '    agency: "Politically Exposed Persons",',
    '    agency_abbr: "PEP",',
    '    entity_name: "John Smith"',
    '  },',
    '  url: null,',
    '  list_country: "Europe",',
    '  list_region: "Europe"',
    '}',
    '',
    '// Business hits have entity_type = "business"',
    '// Person (PSC) hits have entity_type = "person"',
    '// Both stored in: integration_data.business_entity_review_task',
    '//   (key="watchlist", business_entity_verification_id=<kyb_verification_id>)',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — SERP UK BEHAVIOR
# ════════════════════════════════════════════════════════════════════════════

H1('Section 10 — SERP: UK Behavior', c=PURPLE, pb=False)

H2('Does SERP Run for UK Businesses?')
callout(
    'YES — SERP (platform_id=22) runs for UK businesses when configured for the customer. '
    'SERP is NOT country-specific — it runs based on customer integration settings, not business country. '
    'However, its results are more limited for UK businesses because some data sources '
    '(e.g., Google Business Profile, Google Maps) provide different data density for UK vs US.'
)

table(
    ['SERP feature', 'Runs for UK?', 'Notes'],
    [
        ['Google SERP scraping', 'YES — if SERP enabled for customer',
         'Searches Google for "[business name] [city]" to find contact info, website, description'],
        ['Google Business Profile', 'YES — platform_id=39 (SERP_GOOGLE_PROFILE)',
         'Fetches Google Maps/Business listing for UK businesses.\n'
         'UK businesses may or may not have Google Business profiles.'],
        ['Google Business Reviews', 'YES — platform_id=19/20',
         'Google reviews and ratings — works globally including UK'],
        ['Contact info extraction', 'YES', 'Phone, email, website from Google SERP results'],
        ['Address Verified (Google Profile)', 'YES — for UK addresses',
         'caseTabValuesManager: if Google address matches submitted UK address\n'
         '-> googleProfileStatus = "passed" -> "Verified" badge\n'
         'Uses UK postcode matching'],
        ['Worth Website Scanning', 'YES — platform_id=30',
         'Domain age, SSL, malware check — works for .co.uk and other UK TLDs'],
    ],
    col_widths=[2.5, 1.5, 5.6],
)

H2('Google Profile Verified / Unverified for UK Addresses')
code_block([
    '// caseTabValuesManager.ts: Google Profile verification for UK',
    '',
    '// Checks if Google Maps address matches submitted address:',
    'const googleAddr = google_profile?.address  // from SERP_GOOGLE_PROFILE response',
    'const addrStr = submitted_address',
    '',
    '// Verification logic:',
    'if (verified) {',
    '  googleProfileStatus = "passed"',
    '  googleProfileDesc = "Business address (Google) verified."',
    '} else if (hasAddress) {',
    '  googleProfileStatus = "missing"',
    '  googleProfileDesc = "Address present but not verified."',
    '} else {',
    '  googleProfileStatus = "failed"',
    '  googleProfileDesc = MISSING',
    '}',
    '',
    '// For UK businesses:',
    '// "Verified" = Google Maps shows "123 Baker Street, London W1U 6RG" matching submitted',
    '// "Unverified" = Google found the business but address format differs',
    '//   (common for UK: "W1U 6RG" vs "W1U6RG" or "Baker St" vs "Baker Street")',
    '// No badge = business not found on Google Maps at all',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 11 — DATA STORAGE
# ════════════════════════════════════════════════════════════════════════════

H1('Section 11 — Data Storage: What is Saved and Where', c=PURPLE, pb=False)

H2('UK Business Screening — All Storage Locations')
table(
    ['Data type', 'Table', 'Schema', 'Key fields', 'Written by'],
    [
        ['Raw Trulioo KYB response',
         'request_response', 'integration_data',
         'business_id, platform_id=38, response JSONB (FULL Trulioo response),\n'
         'confidence (source quality 0-1), external_id (hfSession)',
         'truliooBusinessResultsStorage.saveRequestResponse()'],
        ['Business entity verification result',
         'business_entity_verification', 'integration_data',
         'business_id, status ("approved"/"rejected"/"in_review"),\n'
         'name, incorporation_date, entity_type, sos_status,\n'
         'external_id (hfSession), unique_external_id',
         'truliooBusinessResultsStorage.storeInitialVerificationRecord()'],
        ['Submitted + Reported addresses',
         'business_entity_address_source', 'integration_data',
         'business_entity_verification_id, address_line_1, city, state,\n'
         'postal_code, submitted (bool), deliverable (bool)',
         'truliooBusinessResultsStorage.storeAddressData()'],
        ['Business names (submitted + reported)',
         'business_entity_names', 'integration_data',
         'business_entity_verification_id, name, submitted (bool), type ("legal"/"dba"),\n'
         'source JSONB, metadata JSONB',
         'truliooBusinessResultsStorage.upsertBusinessEntityNames()'],
        ['Watchlist + PEP + Adverse Media hits',
         'business_entity_review_task', 'integration_data',
         'business_entity_verification_id, key ("watchlist"),\n'
         'status ("success"/"failure"), metadata JSON (WatchlistValueMetadatum[]),\n'
         'entity_type ("business"/"person")',
         'storeBusinessWatchlistResults() (business)\n'
         'truliooWatchlist.storePersonWatchlistResults() (PSC)'],
        ['Adverse media articles (scored)',
         'adverse_media + adverse_media_articles', 'case-service PostgreSQL',
         'business_id, title, link, source, keywords_score,\n'
         'negative_sentiment_score, entity_focus_score, final_score,\n'
         'risk_level, risk_description, media_type',
         'adverseMedia.insertAdverseMedia() (after OpenAI scoring)'],
        ['UBOs/Directors (UK PSC results)',
         'business_entity_people', 'integration_data',
         'business_entity_verification_id, name, submitted (true for PSC),\n'
         'source JSONB, titles[], screening_results JSONB\n'
         '(watchlistHits, provider, screenedAt)',
         'truliooPersonDataStorage.storePersonInBusinessEntityPeople()'],
        ['Address verification review task',
         'business_entity_review_task', 'integration_data',
         'key="address_verification", status ("success"/"failure"),\n'
         'sublabel ("Verified"/"Unverified")',
         'truliooBusinessResultsStorage.createOrUpdateAddressVerificationTask()'],
        ['TIN/EIN verification review task',
         'business_entity_review_task', 'integration_data',
         'key="tin", status ("success"/"failure"),\n'
         'message, sublabel',
         'Middesk (US) / Trulioo (UK via extractRegistrationNumberFromTruliooResponse)'],
        ['Computed facts (217 total)',
         'facts', 'rds_warehouse_public',
         'business_id, name (e.g. "watchlist"), value JSONB',
         'warehouse-service FactService after Kafka facts.v1 event'],
    ],
    col_widths=[2.2, 2.2, 1.4, 2.5, 1.3],
    fs=8.5,
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 12 — FACT ENGINE TRANSFORMATION
# ════════════════════════════════════════════════════════════════════════════

H1('Section 12 — Fact Engine Transformation: From DB to Computed Facts', c=PURPLE, pb=False)

H2('watchlist_raw Fact — Source Selection Priority')
code_block([
    '// integration-service/lib/facts/kyb/index.ts: watchlist_raw fact',
    '// Sources checked in order:',
    '',
    '// SOURCE 1: Middesk (US businesses only)',
    '// Middesk returns reviewTasks[] with key="watchlist"',
    '// -> metadata: WatchlistValueMetadatum[]',
    '',
    '// SOURCE 2: Trulioo (UK/Canada businesses)',
    '// a) First try clientData.watchlistResults (pre-processed array)',
    '// b) If empty, try extractWatchlistResultsFromTruliooResponse(clientData)',
    '//    -> fullServiceDetails Priority 1 extraction',
    '//    -> then transformTruliooBusinessWatchlistResults()',
    '// c) If still empty, try reviewTasks[key="watchlist"].metadata',
    '',
    '// transformTruliooBusinessWatchlistResults():',
    '//   Takes TruliooWatchlistHit[] -> WatchlistValueMetadatum[]',
    '//   entity_type = WATCHLIST_ENTITY_TYPE.BUSINESS',
    '',
    '// Output:',
    '{',
    '  metadata: WatchlistValueMetadatum[],  // all business-level hits',
    '  message: "Found 3 watchlist hit(s)"   // or empty string',
    '}',
])

H2('watchlist (Consolidated) Fact — Business + PSC Combined')
code_block([
    '// calculateConsolidatedWatchlist() merges:',
    '//   watchlist_raw (business KYB hits, entity_type="business")',
    '//   screened_people (PSC hits per person, entity_type="person")',
    '',
    '// Deduplication: createWatchlistDedupKey(hit):',
    '//   key = title + "::" + agency + "::" + entity_name',
    '//   Prevents same hit appearing twice (e.g., from KYB and PSC)',
    '',
    '// Output fact: rds_warehouse_public.facts name="watchlist"',
    '{',
    '  "metadata": [',
    '    {',
    '      "id": "<uuid>",',
    '      "type": "sanctions",        // "sanctions" | "pep" | "adverse_media" | "other"',
    '      "entity_type": "business",  // "business" or "person"',
    '      "metadata": {',
    '        "abbr": "OFSI",',
    '        "title": "UK Financial Sanctions",',
    '        "agency": "OFSI",',
    '        "agency_abbr": "OFSI",',
    '        "entity_name": "Acme Limited"',
    '      },',
    '      "url": "https://example.com/...",',
    '      "list_country": "United Kingdom",',
    '      "list_region": "Europe",',
    '      "score": null               // ⚠️ NOT POPULATED — see Section 6 GAP',
    '    },',
    '    {',
    '      "id": "<uuid>",',
    '      "type": "pep",',
    '      "entity_type": "person",   // ← PSC person hit',
    '      "metadata": {',
    '        "entity_name": "John Smith"',
    '      }',
    '    }',
    '  ],',
    '  "message": "Found 2 watchlist hit(s)"',
    '}',
])

H2('watchlist_hits and adverse_media_hits Facts')
code_block([
    '// watchlist_hits: total count of hits (business + person)',
    'watchlist_hits = watchlist.metadata.length  // simple count',
    '',
    '// adverse_media_hits: count from Trulioo AM_results + adverse_media table',
    'adverse_media_hits = Math.max(',
    '  truliooClientData.watchlistResults.filter(h => h.listType === "ADVERSE_MEDIA").length,',
    '  adverseMediaData.records[0]?.total_risk_count ?? 0',
    ')',
    '// Note: Trulioo AM_results and the scored adverse_media table may have different counts.',
    '// Trulioo AM_results = raw Trulioo hits (unscored)',
    '// adverse_media table = OpenAI-scored articles that passed scoring quality threshold',
    '// adverse_media_hits fact uses whichever is larger (or the scored count if available)',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 13 — API RESPONSES
# ════════════════════════════════════════════════════════════════════════════

H1('Section 13 — API Responses: What Is Surfaced and Where', c=PURPLE, pb=False)

H2('Where UK Watchlist/Screening Data Appears in the API')
table(
    ['API endpoint', 'What watchlist data is included', 'Source table'],
    [
        ['GET /facts/business/{id}/details\n(or /kyb)',
         'watchlist fact: { metadata: WatchlistValueMetadatum[], message }\n'
         'watchlist_hits: total count\n'
         'adverse_media_hits: count\n'
         'source.confidence: Trulioo source quality (0.0-1.0)',
         'rds_warehouse_public.facts (computed fact)\n'
         'confidence from request_response.confidence'],
        ['GET /facts/business/{id}/all\n(admin only)',
         'All 217 facts including:\n'
         'watchlist, watchlist_raw, watchlist_hits,\n'
         'adverse_media_hits, screened_people',
         'rds_warehouse_public.facts'],
        ['GET /businesses/{id}/kyb\n(KYB facts endpoint)',
         'watchlist and related facts included as part of KYB fact set',
         'rds_warehouse_public.facts via Fact Engine'],
        ['360 Report (Kafka: FETCH_REPORT_DATA)',
         'watchlist_hits: hits\nwatchlist_hits_count: total\n'
         'people_watchlist: per-person PSC hits grouped by entity name\n'
         'watchlist_entries: grouped by business/person entity',
         'Fact Engine: watchlist, watchlist_raw, screened_people facts\n'
         'business_entity_review_task (via createWatchlistEntries)'],
        ['Admin UI — KYB Watchlists tab',
         'Reads from the watchlist fact\n'
         'Groups hits by: "Hits for {businessName}" (business type)\n'
         '"Hits for {ownerName}" (person type)',
         'watchlist fact -> metadata filtered by entity_type'],
    ],
    col_widths=[2.8, 3.8, 3.0],
)

H2('Can the Same Data Appear in More Than One Place?')
callout(
    'YES — watchlist data intentionally appears in multiple places for different consumers:\n\n'
    '1. request_response.response JSONB: raw full Trulioo response (ALL data, unprocessed)\n'
    '2. business_entity_review_task.metadata: processed WatchlistValueMetadatum[]\n'
    '3. rds_warehouse_public.facts: computed watchlist fact (Fact Engine output)\n'
    '4. adverse_media_articles: OpenAI-scored adverse media (subset of AM_results)\n\n'
    'These are NOT duplicates — each serves a different purpose:\n'
    '  request_response = audit log / raw data recovery\n'
    '  business_entity_review_task = structured match data for case management\n'
    '  facts = customer-facing API data\n'
    '  adverse_media_articles = risk-scored articles for detailed analysis'
)

H2('Consistency: How Worth AI Ensures Consistency Across Endpoints')
table(
    ['Consistency mechanism', 'What it protects', 'How it works'],
    [
        ['Single source of truth: watchlist fact',
         'All API endpoints reading watchlist data',
         'All endpoints read from rds_warehouse_public.facts (name="watchlist").\n'
         'Computed once by the Fact Engine, cached, re-used.'],
        ['Deduplication in calculateConsolidatedWatchlist()',
         'Business + PSC hits not double-counted',
         'createWatchlistDedupKey(hit): title + agency + entity_name composite key.\n'
         'Same hit from KYB and PSC flows appears only once.'],
        ['Upsert pattern in facts table',
         'Score recalculations do not create duplicate facts',
         'INSERT ON CONFLICT (business_id, name) DO UPDATE SET value = EXCLUDED.value.\n'
         'Always the latest value.'],
        ['atomic lock in business_entity_people',
         'PSC screening not triggered twice for same person',
         'INSERT ... ON CONFLICT (business_entity_verification_id, name) IGNORE.\n'
         'First process wins; subsequent processes skip.'],
        ['Comprehensive View as single arbiter (address)',
         'Address match result consistent between Contact Info tab and Facts API',
         'extractAddressMatchStatusFromDatasourceFields() always uses Comprehensive View.\n'
         'Individual datasource results are ignored.'],
    ],
    col_widths=[2.5, 2.5, 4.6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 14 — 360 REPORT WATCHLIST SECTION
# ════════════════════════════════════════════════════════════════════════════

H1('Section 14 — 360 Report: Watchlist Section Detail', c=PURPLE, pb=False)

H2('_getWatchlistData() — 360 Report Watchlist Assembly')
code_block([
    '// integration-service/src/messaging/kafka/consumers/handlers/report.ts',
    '// Called as part of fetchReportData() parallel Promise.all()',
    '',
    'async _getWatchlistData(body: I360Report) {',
    '  const factsToFetch = [',
    '    "watchlist_raw",    // raw hits from Trulioo/Middesk',
    '    "watchlist",        // consolidated (business + PSC)',
    '    "names_submitted",  // all names the business submitted',
    '    "people",           // all persons on file (owners + discovered)',
    '    "screened_people",  // persons who went through PSC screening',
    '    "legal_name"        // winner legal name fact',
    '  ]',
    '',
    '  const { watchlist, names_submitted, people, legal_name } = await factEngine.getResults()',
    '',
    '  const hits = watchlist?.value?.metadata ?? []   // all consolidated hits',
    '',
    '  // createWatchlistEntries(): groups hits by business name / person name',
    '  // Returns: [{ entity: "Acme Ltd", hits: [...] }, { entity: "John Smith", hits: [...] }]',
    '  const watchlistEntries = createWatchlistEntries(watchlist, names_submitted, people, legal_name)',
    '',
    '  // groupWatchlistHitsByEntityName(): creates { "John Smith": [hit1, hit2], ... }',
    '  const groupedHitsByEntity = groupWatchlistHitsByEntityName(hits)',
    '',
    '  // peopleWatchlist: [{ id, name: "John Smith", titles: [], watchlist_results: [...] }]',
    '  const peopleWatchlist = Object.keys(groupedHitsByEntity).map(entityName => ({',
    '    id: groupedHitsByEntity[entityName][0]?.id || "",',
    '    name: entityName,',
    '    titles: [],',
    '    watchlist_results: groupedHitsByEntity[entityName]',
    '  }))',
    '',
    '  return {',
    '    watchlist_hits: hits.length > 0 ? hits : null,   // all consolidated hits',
    '    watchlist_hits_count: watchlistHitsCount,          // total across all entities',
    '    people_watchlist: peopleWatchlist,                 // per-person grouping',
    '    watchlist_entries: watchlistEntries                // grouped by entity name',
    '  }',
    '}',
])

H2('360 Report Watchlist Output — Real UK Example')
code_block([
    '// _getWatchlistData() output for UK business with hits:',
    '{',
    '  "watchlist_hits": [',
    '    {',
    '      "id": "<uuid>",',
    '      "type": "sanctions",',
    '      "entity_type": "business",',
    '      "metadata": {',
    '        "abbr": "AiB",',
    '        "title": "Register of Insolvencies",',
    '        "agency": "Accountant in Bankruptcy",',
    '        "agency_abbr": "AiB",',
    '        "entity_name": "Acme Limited"',
    '      },',
    '      "url": "https://www.aib.gov.uk/insolvency/acme-ltd"',
    '    },',
    '    {',
    '      "id": "<uuid>",',
    '      "type": "pep",',
    '      "entity_type": "person",',
    '      "metadata": {',
    '        "abbr": "PEP",',
    '        "title": "PEP Database",',
    '        "agency": "Politically Exposed Persons",',
    '        "entity_name": "John Smith"',
    '      }',
    '    }',
    '  ],',
    '  "watchlist_hits_count": 2,',
    '  "people_watchlist": [',
    '    { "name": "John Smith", "watchlist_results": [{ "type": "pep", ... }] }',
    '  ],',
    '  "watchlist_entries": [',
    '    { "entity": "Acme Limited", "hits": [{ "type": "sanctions", ... }] },',
    '    { "entity": "John Smith",   "hits": [{ "type": "pep", ... }] }',
    '  ]',
    '}',
    '',
    '// NOTE: Watchlist hit score/confidence is NOT in this output.',
    '// It is available in request_response.response JSONB only.',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 15 — CONSISTENCY & KNOWN GAPS
# ════════════════════════════════════════════════════════════════════════════

H1('Section 15 — Consistency Guarantees & Known Gaps', c=PURPLE, pb=False)

H2('Where Data May Appear in More Than One Place')
table(
    ['Data element', 'Appears in', 'Are they consistent?'],
    [
        ['Watchlist hits (business)',
         '1. request_response.response JSONB (raw, full Trulioo response)\n'
         '2. business_entity_review_task.metadata (processed)\n'
         '3. rds_warehouse_public.facts name="watchlist" (computed)\n'
         '4. 360 Report watchlist section (assembled from facts)',
         'YES — all derive from the same Trulioo response.\n'
         'Request_response is source of truth for raw data.\n'
         'Facts are computed from business_entity_review_task.\n'
         '360 Report reads the computed fact.'],
        ['Watchlist hits (PSC/person)',
         '1. request_response.response JSONB (raw PSC Trulioo response)\n'
         '2. business_entity_people.screening_results JSONB\n'
         '3. business_entity_review_task.metadata (entity_type="person")\n'
         '4. screened_people fact\n'
         '5. watchlist consolidated fact (entity_type="person" entries)',
         'YES — all derive from PSC flow.\n'
         'Deduplication in consolidated fact prevents double-counting.'],
        ['Adverse media',
         '1. request_response.response JSONB (raw AM_results from Trulioo)\n'
         '2. business_entity_review_task.metadata (type="adverse_media")\n'
         '3. adverse_media + adverse_media_articles (OpenAI-scored subset)\n'
         '4. adverse_media_hits fact (count)',
         'PARTIAL — counts may differ.\n'
         'Trulioo AM_results count = all raw Trulioo hits.\n'
         'adverse_media_articles count = only hits that passed OpenAI scoring.\n'
         'adverse_media_hits fact uses scored count if available.'],
        ['Address verification',
         '1. business_entity_address_source (raw submitted + reported)\n'
         '2. business_entity_review_task key="address_verification" (match/nomatch)\n'
         '3. address_verification_boolean fact\n'
         '4. Contact Information tab (reads address_source + review task)',
         'YES — all from same Trulioo Comprehensive View result.'],
        ['Business name',
         '1. data_businesses.name (submitted)\n'
         '2. business_entity_names (submitted=true + reported from Trulioo)\n'
         '3. legal_name fact (winner from Fact Engine)\n'
         '4. names_submitted fact',
         'YES but may differ:\n'
         'submitted name = what applicant typed\n'
         'reported name = what Trulioo/Middesk found in registry\n'
         'legal_name fact = Fact Engine winner (usually reported name)'],
    ],
    col_widths=[2.2, 3.8, 3.6],
)

H2('Known Gaps in UK Screening Pipeline')
gap(
    'GAP 1: Watchlist hit score/confidence NOT surfaced in API or 360 Report.\n'
    'Trulioo returns score (0.0-1.0) per hit (e.g., "Register of Insolvencies" hit with score=0.98).\n'
    'This is captured in TruliooWatchlistHit.confidence but NOT copied to WatchlistValueMetadatum.\n'
    'Fix: Add score?: number to WatchlistValueMetadatum; copy in mapWatchlistHits().\n'
    'Impact: Analysts cannot see match relevancy — a score=0.3 hit looks the same as score=1.0.'
)
gap(
    'GAP 2: UK SIC codes NOT surfaced to customers despite being captured.\n'
    'OC\'s industry_code_uids contains gb_sic-56101 for UK businesses.\n'
    'Trulioo may also return StandardizedIndustries with UK SIC codes.\n'
    'These are stored in classification_codes fact but NO downstream consumer reads them.\n'
    'Fix: Add Kafka handler for UPDATE_CLASSIFICATION_CODES + new API endpoint.'
)
gap(
    'GAP 3: Adverse media article count inconsistency between Trulioo raw hits and OpenAI-scored articles.\n'
    'Trulioo may return 12 AM_results but OpenAI scoring may reject 3 as low-quality.\n'
    'The adverse_media_hits fact returns the scored count (9) not the raw Trulioo count (12).\n'
    'The watchlist_hits count includes all 12 (from consolidated watchlist fact).\n'
    'This can confuse analysts: "12 hits" in one place, "9 articles" in another.'
)
warn(
    'LIMITATION: Priority 2 (summary format) watchlist results lack entity-level detail.\n'
    'If Trulioo returns only the summary format (wlHitsNumber=138, amHitsNumber=12),\n'
    'Worth AI creates placeholder hits: "Watchlist hit 1 of 138" with no entity name or URL.\n'
    'This occurs when Trulioo does not return fullServiceDetails.WatchlistHitDetails.\n'
    'Workaround: Use the raw request_response.response JSONB to inspect full Trulioo output.'
)
ok(
    'CONFIRMED WORKING: PSC screening deduplication.\n'
    'The atomic ON CONFLICT lock in business_entity_people prevents duplicate PSC screenings\n'
    'even if two processes race. Idempotent for UK UBO/Director screening.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 16 — FULL END-TO-END FLOW DIAGRAM
# ════════════════════════════════════════════════════════════════════════════

H1('Section 16 — Full End-to-End Flow: UK Business Submission to 360 Report', c=PURPLE, pb=False)

lineage([
    'T+0:00  UK Business submits onboarding application',
    '  data_businesses: name="Acme Limited", address_country="GB", tin="12345678" (encrypted)',
    '  data_cases: created, status=SUBMITTED',
    '  Kafka: BUSINESS_SUBMITTED',
    '       |',
    'T+0:01  integration-service receives Kafka event',
    '  truliooBusiness.canIRun():',
    '    1. data_integration_settings.trulioo.status = "ENABLED" -> OK',
    '    2. rel_customer_setup_countries: "GB" is_selected=true -> OK',
    '    3. normalizedBusinessCountry = "GB" (UK->GB normalised)',
    '    -> shouldRun = true',
    '       |',
    'T+0:02  Trulioo OAuth token acquired (cached)',
    '  POST https://api.globaldatacompany.com/oauth2/v1/token',
    '  -> access_token cached in truliooTokenManager',
    '       |',
    'T+0:03  KYB Flow Initiation',
    '  POST /flows/{TRULIOO_KYB_FLOWID}/init',
    '  -> elements[] with field roles returned',
    '  -> hfSession = "session-uuid-abc123"',
    '       |',
    'T+0:04  KYB Payload Built and Submitted',
    '  truliooBusinessKYBProcessor.processKYBFlow():',
    '    businessPayload: { companyName, companyCountryIncorporation="GB", companyZip,',
    '                       companyregno, companyAddressFull (full address, stored internally) }',
    '    truliooPayload:  { ...businessPayload, companyAddressFull (STREET ONLY) }',
    '  POST /flows/{TRULIOO_KYB_FLOWID}/submit with truliooPayload',
    '       |',
    'T+0:05 to T+0:15  Trulioo processes KYB request',
    '  - Companies House lookup for UK registration number',
    '  - Address verification against UK Comprehensive View (Royal Mail / Companies House)',
    '  - Advanced Business Watchlist screening:',
    '      OFAC SDN/SSI/Non-SDN, BIS Entity List, UK OFSI Financial Sanctions,',
    '      UK Register of Insolvencies, UK Companies House Disqualified Directors,',
    '      Employment Tribunal, EU Consolidated, UN Consolidated, PEP Database,',
    '      Adverse Media (UK + global sources)',
    '       |',
    'T+0:15  Trulioo returns KYB response (TruliooFlowResult)',
    '  clientData.status = "completed"',
    '  clientData.businessData = { name, ubos[], directors[], business_addresses[] }',
    '  clientData.watchlistResults[] OR flowData.serviceData[].fullServiceDetails',
    '  DatasourceResults[Comprehensive View].DatasourceFields: name=match, addr=match, etc.',
    '  DatasourceResults[Advanced Business Watchlist].AppendedFields[WatchlistHitDetails]:',
    '    WL_results: [{ score:1, subjectMatched:"Acme Ltd", sourceListType:"Register of Insolvencies" }]',
    '    AM_results: [{ score:0.95, subjectMatched:"Acme Ltd", URL:"https://bbc.co.uk/..." }]',
    '    PEP_results: [] (empty for this business)',
    '       |',
    'T+0:16  storeBusinessVerificationResults():',
    '  1. saveRequestResponse() -> integration_data.request_response (FULL raw response)',
    '  2. storeInitialVerificationRecord() -> integration_data.business_entity_verification',
    '  3. storeRegistrationData() -> business_entity_verification (name, status, entity_type)',
    '  4. storeAddressData():',
    '     - business_entity_address_source: submitted=true (street only, no suite)',
    '     - business_entity_address_source: submitted=false (StandardizedLocations from Trulioo)',
    '     - createOrUpdateAddressVerificationTask() -> business_entity_review_task key="address_verification"',
    '       status = "success" (Comprehensive View all-match)',
    '  5. upsertBusinessEntityNames():',
    '     - business_entity_names: "Acme Limited" (submitted=true)',
    '     - business_entity_names: "ACME LTD" (from Trulioo, submitted=false)',
    '  6. storeBusinessWatchlistResults():',
    '     - extractWatchlistResultsFromTruliooResponse() -> Priority 1 extraction',
    '       -> WL_results[0]: { listType:SANCTIONS, listName:"Register of Insolvencies",',
    '                           confidence:1.0, matchDetails:"Acme Ltd",',
    '                           sourceAgencyName:"Accountant in Bankruptcy" }',
    '     - business_entity_review_task: key="watchlist", status="failure",',
    '       metadata=[WatchlistValueMetadatum(entity_type=BUSINESS)]',
    '  7. processAndPersistTruliooAdverseMedia():',
    '     - AM_results[0]: score=0.95, URL="https://bbc.co.uk/..."',
    '     - extractTitleFromUrl() -> "Acme ltd fraud investigation"',
    '     - scoreAdverseMedia() via OpenAI -> finalScore=0.88, riskLevel="HIGH"',
    '     - insertAdverseMedia() -> adverse_media + adverse_media_articles tables',
    '       |',
    'T+0:17  PSC Screening (AUTOMATIC for UK):',
    '  shouldScreenPSCsForBusiness(businessId, "GB"):',
    '    internationalKYBEnabled=true, isUS=false -> shouldScreen=true',
    '  extractPersonsFromBusinessData():',
    '    - ubos: [{ fullName:"John Smith", ownershipPercentage:75 }] from Trulioo',
    '    - directors: [{ fullName:"Jane Doe", title:"Secretary" }] from Trulioo',
    '    - applicant owners from data_owners (case-service)',
    '    - deduplicatePersons(): 2 unique persons',
    '  For each person (concurrent):',
    '    ATOMIC LOCK: INSERT business_entity_people ON CONFLICT IGNORE',
    '    screenPersonWithPSCFlow():',
    '      POST /flows/{TRULIOO_PSC_FLOWID}/init -> hfSession',
    '      POST /flows/{TRULIOO_PSC_FLOWID}/submit { firstName, lastName, dob, country }',
    '      PSC result: PEP_results = [] (John Smith is not a PEP)',
    '      storePersonInBusinessEntityPeople(): screening_results = { watchlistHits: [] }',
    '       |',
    'T+0:25  Kafka facts.v1 CALCULATE_BUSINESS_FACTS event fired',
    '  Fact Engine runs for this business:',
    '  1. watchlist_raw fact:',
    '     business source (Trulioo) -> reads business_entity_review_task key="watchlist"',
    '     -> WatchlistValueMetadatum[{type:sanctions, entity_type:business, ...}]',
    '  2. screened_people fact:',
    '     -> reads business_entity_people.screening_results',
    '     -> extractWatchlistHitsFromScreenedPeople(): no PEP hits for John Smith',
    '  3. watchlist (consolidated):',
    '     calculateConsolidatedWatchlist() merges business + person hits',
    '     dedup: no overlaps',
    '     -> { metadata: [SANCTIONS hit for "Acme Ltd"], message: "Found 1 hit" }',
    '  4. adverse_media_hits: 1 (from adverse_media_articles table)',
    '  5. address_verification_boolean: true (Comprehensive View matched)',
    '  6. naics_code: 561499 (no Trulioo NAICS for UK, AI safety net triggered)',
    '  7. industry: "Other Services" (from 561499 sector)',
    '       |',
    '  Kafka facts.v1 UPDATE_NAICS_CODE, UPDATE_INDUSTRY etc.',
    '  warehouse-service: UPSERT rds_warehouse_public.facts (all facts)',
    '       |',
    'T+0:30  Customer API now available:',
    '  GET /facts/business/{id}/details:',
    '    watchlist: { metadata: [...], message: "Found 1 hit" }',
    '    watchlist_hits: 1',
    '    adverse_media_hits: 1',
    '    address_verification_boolean: true',
    '    source.confidence: 0.85 (Trulioo entity match quality)',
    '  GET /facts/business/{id}/kyb:',
    '    Same facts via KYB endpoint',
    '       |',
    'T+varies  360 Report generated (Kafka: FETCH_REPORT_DATA):',
    '  _getWatchlistData():',
    '    Facts: watchlist, names_submitted, people, screened_people, legal_name',
    '    hits = watchlist.metadata (1 hit: Register of Insolvencies)',
    '    watchlistEntries = [{ entity:"Acme Limited", hits:[{type:sanctions,...}] }]',
    '    peopleWatchlist = [] (John Smith had no PSC hits)',
    '    Return: { watchlist_hits:[...], watchlist_hits_count:1, people_watchlist:[], ... }',
    '',
    '  360 Report section: "Watchlists" shows:',
    '    "Acme Limited": 1 SANCTIONS hit (Register of Insolvencies, Accountant in Bankruptcy)',
    '    score/confidence: NOT SHOWN (GAP — see Section 6)',
])

# ── Save ──────────────────────────────────────────────────────────────────────
out = ('/workspace/AI-Powered-NAICS-Industry-Classification-Agent/'
       'modeling/Worth_AI_UK_Screening_Lineage.docx')
doc.save(out)
size_kb = round(os.path.getsize(out) / 1024, 1)
print(f'Saved  : {out}')
print(f'Size   : {size_kb} KB')
print('Ready  : Import into Google Docs via File > Import')
