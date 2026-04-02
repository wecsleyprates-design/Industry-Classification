"""
Build: Worth AI Admin UI — Complete Lineage Reference
Google-Docs-compatible .docx

Maps every field, badge, status, tab, and panel visible in
admin.joinworth.com/customers to its exact origin:
source code file, database table, API call, Kafka event, ML model.

Verified against:
  - SIC-UK-Codes (case-service, integration-service, manual-score-service, warehouse-service)
  - Entity-Matching-Ref
  - Screenshots of the admin UI (Lisa's Nail Salon case, customers list)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page (landscape for wide tables) ──────────────────────────────────────────
s = doc.sections[0]
s.page_width    = Inches(11.0)
s.page_height   = Inches(8.5)
s.left_margin   = Inches(0.7)
s.right_margin  = Inches(0.7)
s.top_margin    = Inches(0.6)
s.bottom_margin = Inches(0.6)
s.orientation   = 1
PAGE_W = 9.6

# ── Colours ────────────────────────────────────────────────────────────────────
PURPLE = RGBColor(0x5B, 0x21, 0xB6)
BLUE   = RGBColor(0x1E, 0x40, 0xAF)
TEAL   = RGBColor(0x04, 0x5F, 0x7E)
GREEN  = RGBColor(0x06, 0x5F, 0x46)
RED    = RGBColor(0x99, 0x1B, 0x1B)
AMBER  = RGBColor(0x78, 0x35, 0x00)
SLATE  = RGBColor(0x33, 0x41, 0x55)
DARK   = RGBColor(0x0F, 0x17, 0x2A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Low-level XML ──────────────────────────────────────────────────────────────
def _shade(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)

def _left_border(cell, hex6, sz=28):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    b = OxmlElement('w:left')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '0');   b.set(qn('w:color'), hex6)
    tcB.append(b); tcPr.append(tcB)

def _cell_margins(cell, top=60, right=100, bottom=60, left=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('right',right),('bottom',bottom),('left',left)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'), 'none'); tblB.append(b)
    tblPr.append(tblB)

def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblH = OxmlElement('w:tblHeader'); trPr.append(tblH)

# ── High-level helpers ─────────────────────────────────────────────────────────
def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

def _heading(level, text, colour, size, sb, sa, pb=False):
    if pb: doc.add_page_break()
    p = doc.add_heading('', level=level)
    p.paragraph_format.space_before   = Pt(sb)
    p.paragraph_format.space_after    = Pt(sa)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = colour
    return p

def H1(t, c=PURPLE, pb=False): return _heading(1, t, c, 20, 16, 6, pb)
def H2(t, c=BLUE):             return _heading(2, t, c, 13.5, 13, 4)
def H3(t, c=TEAL):             return _heading(3, t, c, 11.5, 10, 3)
def H4(t, c=SLATE):            return _heading(4, t, c, 10.5,  8, 2)

def body(text, size=10.5, colour=SLATE, sa=5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = colour; r.bold = bold
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(0)
    return p

def body_parts(parts, sa=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, colour in parts:
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(10.5)
        r.font.color.rgb = colour or SLATE
    return p

def bullet(prefix, text, size=10.5, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if prefix:
        rb = p.add_run(prefix); rb.bold = True
        rb.font.size = Pt(size); rb.font.color.rgb = DARK
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.color.rgb = SLATE

def code_block(lines, size=8.5):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, 'F1F5F9'); cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=80, right=120, bottom=80, left=120)
    first = True
    for line in lines:
        cp = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        r = cp.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(size); r.font.color.rgb = DARK
    spacer(5)

def callout(text, bg='EFF6FF', border='1D4ED8', tc=None, size=10):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, bg); _left_border(cell, border, sz=28)
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=60, right=120, bottom=60, left=160)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(0)
    r = cp.add_run(text); r.font.size = Pt(size)
    r.font.color.rgb = tc or RGBColor.from_string(border)
    spacer(7)

def warn(text):  callout(text, 'FFFBEB', 'D97706', RGBColor(0x78,0x35,0x00))
def gap(text):   callout(text, 'FEE2E2', 'DC2626', RGBColor(0x7F,0x1D,0x1D))
def ok(text):    callout(text, 'F0FDF4', '059669', RGBColor(0x06,0x5F,0x46))

def lineage(lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, 'F0F9FF'); _left_border(cell, '0284C7', sz=28)
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=80, right=120, bottom=80, left=140)
    first = True
    for line in lines:
        cp = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        r = cp.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(8.5); r.font.color.rgb = DARK
    spacer(8)

def table(headers, rows, col_widths=None, hdr='DBEAFE', alt='F8FAFF', fs=9.5):
    ncols = len(headers)
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = tbl.rows[0]; _repeat_header(hrow)
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]; _shade(cell, hdr); _cell_margins(cell)
        p = cell.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(fs); r.font.color.rgb = DARK
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1: _shade(cell, alt)
            _cell_margins(cell)
            p = cell.paragraphs[0]; v = str(val)
            if v.startswith('YES') or v.startswith('✅'): rgb, bld = GREEN, True
            elif v.startswith('NO') or v.startswith('❌'): rgb, bld = RED, True
            elif v.startswith('WARN') or v.startswith('⚠️'): rgb, bld = AMBER, True
            elif v.startswith('GAP') or v.startswith('🔴'): rgb, bld = RED, True
            else: rgb, bld = SLATE, False
            r = p.add_run(v); r.font.size = Pt(fs)
            r.font.color.rgb = rgb; r.bold = bld
    if col_widths:
        for ri2 in range(len(tbl.rows)):
            for ci2, w in enumerate(col_widths):
                tbl.rows[ri2].cells[ci2].width = Inches(w)
    spacer(8)

# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT START
# ════════════════════════════════════════════════════════════════════════════

spacer(18)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Worth AI Admin Portal — Complete Data Lineage Reference')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = PURPLE
p.paragraph_format.space_after = Pt(5)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Every Field, Badge, Status, Tab & Score: Where It Comes From')
r.font.size = Pt(13); r.font.color.rgb = BLUE
p.paragraph_format.space_after = Pt(14)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Verified against case-service, integration-service, manual-score-service source code.\n'
    'Based on screenshots of admin.joinworth.com/customers — Lisa\'s Nail Salon case.'
)
r.font.size = Pt(10); r.font.color.rgb = SLATE
p.paragraph_format.space_after = Pt(18)

H2('Document Sections')
table(
    ['Section', 'Title', 'What it covers'],
    [
        ['1', 'Customers List Page',
         'Customer ID, Business Name, Account Type (Sandbox/Production), Status, Owner'],
        ['2', 'Case Overview Tab',
         'Case ID, Business ID, Applicant, Worth Score, Risk Level, Case Activity feed'],
        ['3', 'Worth Score Deep Dive',
         'How 819/850 is computed, Model 3.0, score categories, auto-approve vs manual review'],
        ['4', 'KYB Tab — All Sub-tabs',
         'Background, Business Registration, Contact Information, Website, Watchlists'],
        ['5', 'KYB — Business Details Fields',
         'Every field lineage: name, address, age, revenue, employees, industry, MCC, NPI'],
        ['6', 'KYC Tab',
         'Owner identity, IDV, field-level Match badges, Verified/Unverified status'],
        ['7', 'Case Statuses — Full Decision Tree',
         'All 20 statuses, who sets each, auto vs manual, which pipeline, which tables'],
        ['8', 'Sandbox vs Production Deep Dive',
         'What it means, how it is set, which vendors change behavior, what data is real vs fake'],
        ['9', 'Public Records Tab',      'Middesk SOS data, NPI lookups, adverse media'],
        ['10', 'Banking & Accounting Tabs', 'Plaid, Rutter connections, data tables'],
        ['11', 'Full Integration & Pipeline Map',
         'Every vendor → platform_id → API call → Kafka event → table → UI field'],
    ],
    col_widths=[0.5, 2.2, 6.9],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CUSTOMERS LIST
# ════════════════════════════════════════════════════════════════════════════

H1('Section 1 — Customers List Page', c=PURPLE)
callout(
    'The Customers list page (admin.joinworth.com/customers) shows one row per customer '
    '(a "customer" in Worth AI is the bank/lender/fintech that has integrated Worth AI — '
    'NOT the end-business being verified). Each row represents one onboarding application '
    'submitted by a business to that customer\'s Worth AI-powered portal.'
)

H2('What is a "Customer" vs a "Business"?')
table(
    ['Entity', 'Worth AI term', 'Example', 'Stored in table'],
    [
        ['The bank / lender / fintech using Worth AI',
         'Customer (or Tenant)', 'Worldpay, Branch, Wellfit',
         'data_customers (case-service PostgreSQL)'],
        ['The merchant / SMB being verified',
         'Business', 'Lisa\'s Nail Salon',
         'data_businesses (case-service PostgreSQL)'],
        ['The onboarding application',
         'Case', 'Lisa\'s Nail Salon → Worldpay application',
         'data_cases (case-service PostgreSQL)'],
        ['The owner / controller',
         'Applicant / Owner',
         'Thanh Nguyen (25% owner)',
         'data_applicants + data_owners (case-service PostgreSQL)'],
    ],
    col_widths=[2.4, 1.8, 2.1, 3.3],
)

H2('Columns in the Customers List — Field Lineage')
table(
    ['Column shown', 'UI label', 'Source table / field', 'How populated'],
    [
        ['Customer ID',
         'e.g. 3ed63df1-... (truncated UUID)',
         'data_cases.id (UUID primary key)',
         'Auto-generated by case-service on case creation via uuid() call'],
        ['Business Name',
         'e.g. "Lisa\'s Nail Salon"',
         'data_businesses.name',
         'Submitted by the business applicant during onboarding form'],
        ['Account Type',
         'Sandbox (yellow badge) / Production (green badge)',
         'data_customers.customer_type  OR  derived from Kafka payload field customer_type',
         'Set when the customer (tenant) is provisioned. '
         'Valid values: "PRODUCTION" | "SANDBOX" — from integration-service '
         'src/messaging/kafka/consumers/handlers/schema.ts line 323'],
        ['Onboarding Date',
         'e.g. 3/27/2026',
         'data_cases.created_at',
         'Timestamp of case creation (UTC)'],
        ['Customer Owner',
         'e.g. "Joannah Sesebo"',
         'data_users.first_name + last_name WHERE assigned to this customer',
         'Set during customer provisioning or manually assigned in admin'],
        ['Status',
         'Active (green) / Invited (yellow)',
         'data_cases.status_id -> CASE_STATUS constant',
         'Updated by case-service state machine on every status transition. '
         'Invited = status_id 1, Active/Auto Approved = 6 (see Section 7)'],
        ['Actions (eye icon)',
         'View case detail',
         'Navigates to /customers/{case_id}',
         'Frontend routing only — no additional API call'],
    ],
    col_widths=[1.8, 2.0, 2.7, 3.1],
)

H2('Sandbox vs Production — What It Actually Means')
callout(
    'SANDBOX: The bank/lender is testing Worth AI. Real vendor APIs may be replaced '
    'with mock/sandbox versions. Business data entered is test data. Worth Scores are '
    'generated but should not be used for real credit decisions.\n\n'
    'PRODUCTION: The bank/lender is live. Real vendor APIs are called. '
    'Data is real. Worth Scores are used for real underwriting decisions.'
)
body(
    'Sandbox/Production is a PER-CUSTOMER (per-tenant) setting, not per-case. '
    'Every case submitted by a Sandbox customer is automatically a sandbox case. '
    'The setting is stored in data_customers.customer_type and propagated via Kafka '
    'to integration-service when the customer is provisioned. '
    'Integration-service then stores it in data_integration_settings for each customer.'
)
H3('Which integrations change behavior in Sandbox mode?')
table(
    ['Integration', 'Platform ID', 'Sandbox behavior', 'Production behavior'],
    [
        ['Equifax', '17',
         'EquifaxSandboxStrategy — calls Equifax sandbox API endpoint\n'
         '(envConfig.EQUIFAX_SANDBOX_*)',
         'EquifaxProductionStrategy — calls live Equifax API'],
        ['Plaid IDV', '18',
         'PlaidIdvSandboxStrategy — uses Plaid sandbox environment\n'
         'Returns simulated identity verification flows',
         'PlaidIdvProductionStrategy — real identity checks'],
        ['GIACT gVerify / gAuthenticate', '26',
         'GiactSandboxStrategy — calls GIACT sandbox endpoint\n'
         'Can also fall back to GiactMockStrategy (fake data)',
         'GiactProductionStrategy — real bank account verification'],
        ['KYX', '40',
         'KyxSandboxStrategy — uses KYX_SANDBOX_* env vars',
         'KyxProductionStrategy — real KYX calls'],
        ['Middesk', '16',
         'No explicit sandbox strategy — uses same API\n'
         'but may use Middesk\'s own test business records',
         'Same API — live SOS registry lookups'],
        ['OpenCorporates', '23',
         'No sandbox strategy — always queries Redshift pre-loaded data',
         'Same — Redshift query (no live API call)'],
        ['ZoomInfo', '24',
         'No sandbox strategy — always queries Redshift pre-loaded data',
         'Same — Redshift query (no live API call)'],
        ['Trulioo', '38',
         'No explicit sandbox — test data from Trulioo test environment if configured',
         'Live Trulioo KYB API'],
    ],
    col_widths=[1.5, 1.0, 3.5, 3.6],
)
H3('Source code reference for Sandbox strategy pattern')
code_block([
    '// integration-service/src/helpers/strategyPlatformFactory.ts',
    '// Maps platform ID to PRODUCTION / SANDBOX class path:',
    'const platformMap = {',
    '  [INTEGRATION_ID.EQUIFAX]: {',
    '    PRODUCTION: "../../lib/equifax/strategies/EquifaxProductionStrategy",',
    '    SANDBOX:    "../../lib/equifax/strategies/EquifaxSandboxStrategy"',
    '  },',
    '  [INTEGRATION_ID.PLAID_IDV]: {',
    '    PRODUCTION: "../../lib/plaid/strategies/PlaidIdvProductionStrategy",',
    '    SANDBOX:    "../../lib/plaid/strategies/PlaidIdvSandboxStrategy"',
    '  }',
    '}',
    '',
    '// Per-customer strategy is stored in:',
    '// data_integration_settings.settings JSONB -> { equifax: { mode: "SANDBOX" } }',
    '// Table: integration_data.data_integration_settings (integration-service PostgreSQL)',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CASE OVERVIEW
# ════════════════════════════════════════════════════════════════════════════

H1('Section 2 — Case Overview Tab', c=PURPLE, pb=False)

H2('Header: "Case Details — Lisa\'s Nail Salon • 03/27/2026 4:17 PM • Integrations Complete"')
table(
    ['UI element', 'Source', 'Table / field'],
    [
        ['"Lisa\'s Nail Salon"', 'Business legal name',
         'data_businesses.name'],
        ['"03/27/2026 4:17 PM"', 'Case creation timestamp (UTC)',
         'data_cases.created_at'],
        ['"Integrations Complete"', 'All integration tasks finished',
         'Derived: all rows in integration_data.business_integration_tasks '
         'WHERE business_id = X have status = "completed" or "failed".\n'
         'Computed by integrationsCompletionTracker.ts in integration-service.'],
    ],
    col_widths=[2.4, 2.2, 5.0],
)

H2('Left Panel — Case Identity Fields')
table(
    ['UI label', 'Value shown', 'Source table / field', 'Notes'],
    [
        ['CASE ID', 'UUID (truncated)', 'data_cases.id',
         'Primary key of the case. Unique per (business, customer) pair.'],
        ['BUSINESS ID', 'UUID (truncated)', 'data_businesses.id',
         'Primary key of the business. One business can have multiple cases '
         'across different customers.'],
        ['Applicant(s)', '"Joannah Sesebo"', 'data_applicants.first_name + last_name',
         'The person who submitted the application via the onboarding flow. '
         'Linked to the case via rel_case_applicants.'],
        ['Assigned to', '"Unassigned"',
         'data_cases.assigned_to_user_id -> data_users',
         'Analyst assignment. NULL shown as "Unassigned". '
         'Set manually via admin UI PATCH /cases/{id}'],
    ],
    col_widths=[1.6, 1.8, 2.8, 3.4],
)

H2('Right Panel — Case Results (Worth Score)')
table(
    ['UI element', 'Value', 'Source table', 'How computed'],
    [
        ['Risk Level badge',
         '"Low Risk ✓"',
         'business_scores.risk_level',
         'Derived from Worth Score via score_decision_matrix table.\n'
         'SQL: SELECT risk_level FROM score_decision_matrix\n'
         'WHERE score >= range_start AND score <= range_end AND customer_id = X'],
        ['Worth Score',
         '"819/850"',
         'business_scores.weighted_score_850',
         'Manual scoring model (manual-score-service). '
         'Scale: 0–850 (analogous to FICO). See Section 3 for full computation.'],
        ['Score bar (red/yellow/green)',
         'Visual representation',
         'business_scores.weighted_score_850',
         'Three coloured segments: RED (0–499 High Risk), '
         'YELLOW (500–699 Medium Risk), GREEN (700–850 Low Risk). '
         'Cutoffs defined in score_decision_matrix per customer.'],
        ['* Generated on 03/27/26 using Model 3.0',
         'Model version label',
         'score_config_history.version (or similar)',
         '"Model 3.0" refers to the score weightage configuration version stored in '
         'score_config_history table. business_scores.score_weightage_config FK '
         'references the active config row at the time of scoring.'],
    ],
    col_widths=[1.8, 1.4, 2.2, 4.2],
)

H2('Case Overview — Right Panel (Key Insights, Worth Score, Company Profile...)')
body(
    'The right panel items (Key Insights, Worth Score, Company Profile, Financial Trends, '
    'Public Records) are section links in the Worth AI PDF report, NOT separate API calls. '
    'They navigate to sections of the same case detail view.'
)
table(
    ['Panel link', 'What it shows', 'Data source'],
    [
        ['Key Insights', 'AI-generated narrative summary of the case',
         'Generated by AI scoring pipeline from score_inputs table.\n'
         'Also pulled from business_scores.score_decision field (text).'],
        ['Worth Score',
         'Breakdown of score by category (Banking, Public Records, Financial Strength)',
         'business_score_factors table — one row per score factor per score.\n'
         'Each row: category_id, factor_id, weighted_score_850, status.'],
        ['Company Profile',
         'KYB data: name, address, age, industry, employees, revenue',
         'data_businesses + rds_warehouse_public.facts (JSONB) via Pipeline A.'],
        ['Financial Trends',
         'Bank statement analysis, revenue trends, cash flow',
         'Banking integration data (Plaid): plaid_transactions, bank_statements table.'],
        ['Public Records',
         'SOS registration, watchlists, adverse media',
         'Middesk response (platform_id=16), Trulioo watchlist results (platform_id=38).'],
    ],
    col_widths=[1.8, 3.0, 4.8],
)

H2('Case Activity Feed')
table(
    ['Activity type', 'Example shown', 'Source table', 'Who writes it'],
    [
        ['System message',
         '"A new Worth Score of 819 has been generated for Lisa\'s Nail Salon"',
         'case_audit_logs (or similar audit table)',
         'Written automatically by manual-score-service after score computation completes. '
         'Kafka event: GENERATE_AI_SCORE -> score stored -> audit log written.'],
        ['User action',
         '"Case for Lisa\'s Nail Salon has been updated by Joannah Sesebo"',
         'case_audit_logs.user_id + action_type',
         'Written by case-service on any PATCH /cases/{id} operation.'],
        ['System message',
         '"Joannah Sesebo added Thanh to first_name"',
         'case_audit_logs.field_changes JSONB',
         'Diff of before/after values on data_businesses or data_applicants.'],
        ['Add a Comment box',
         'Free-text analyst comment',
         'case_comments table: (case_id, user_id, text, created_at)',
         'Written by POST /cases/{id}/comments — case-service API.'],
    ],
    col_widths=[1.8, 2.8, 2.2, 2.8],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — WORTH SCORE DEEP DIVE
# ════════════════════════════════════════════════════════════════════════════

H1('Section 3 — Worth Score: How 819/850 Is Computed', c=PURPLE, pb=False)

callout(
    'The Worth Score is NOT produced by the XGBoost entity-matching model '
    '(that model only produces match probabilities for vendor record matching). '
    'The Worth Score is a weighted multi-factor scoring model built in manual-score-service. '
    'It aggregates data from Banking (Plaid), Public Records (Middesk/Trulioo), '
    'and Financial Strength (accounting/tax) integrations into a single 0–850 score.'
)

H2('Score Architecture — How It Flows')
lineage([
    'TRIGGER: Business onboarding completes',
    '  -> Kafka: ONBOARDING_INVITE / APPLICATION_EDIT / MONITORING_REFRESH event',
    '  -> integration-service fires all integrations (Middesk, Plaid, Trulioo, Equifax...)',
    '  -> Each integration completes a "task" -> Kafka: UPDATE_INTEGRATION_DATA_FOR_SCORE',
    '       |',
    '       v',
    'manual-score-service/src/workers/score/taskHandler.ts',
    '  taskQueue.process("UPDATE_INTEGRATION_DATA_FOR_SCORE", async (job) => {',
    '    saveIntegrationData(job.data)           -> Redis + business_score_triggers table',
    '    checkScoreDataStatus(...)               -> are all required integrations done?',
    '    if (isScoreReadyToGenerate.status) {',
    '      producer.send({ topic: kafkaTopics.AI_SCORES,',
    '                      event: kafkaEvents.GENERATE_AI_SCORE, ... })',
    '    }',
    '  })',
    '       |',
    '       v',
    'Score calculation worker receives GENERATE_AI_SCORE:',
    '  calculateBankingScore()        -> banking.js  (Plaid data)',
    '  calculatePublicRecordsScore()  -> public_records.js  (Middesk, Trulioo)',
    '  calculateFinancialScore()      -> financial_strength.js  (Rutter/QuickBooks)',
    '       |',
    '  Each factor: score_100 (0-100) * weightage -> weighted_score_100',
    '  Sum all weighted_score_100 -> total_score_100',
    '  Scale: weighted_score_850 = total_score_100 * 8.5',
    '       |',
    '       v',
    'getScoreDecision(weighted_score_850, customerID)',
    '  -> SELECT * FROM score_decision_matrix WHERE customer_id = X',
    '     -> score 700-850: risk_level="LOW",    decision="AUTO_APPROVED"',
    '     -> score 500-699: risk_level="MEDIUM", decision="UNDER_MANUAL_REVIEW"',
    '     -> score 0-499:   risk_level="HIGH",   decision="UNDER_MANUAL_REVIEW" or "AUTO_REJECTED"',
    '       |',
    '       v',
    'updateScore() -> UPDATE business_scores SET weighted_score_850=819, risk_level="LOW",',
    '                  score_decision="AUTO_APPROVED", status="COMPLETED"',
    '       |',
    '       v',
    'Kafka: UPDATE_CASE_STATUS -> case-service',
    '  -> UPDATE data_cases SET status_id = CASE_STATUS.AUTO_APPROVED (6)',
    '     OR status_id = CASE_STATUS.UNDER_MANUAL_REVIEW (4)',
])

H2('Score Computation Tables')
table(
    ['Table', 'Database', 'Purpose'],
    [
        ['score_categories', 'case-service PostgreSQL',
         'Defines the 3 scoring categories: banking, public_records, financial_strength'],
        ['score_category_factors', 'case-service PostgreSQL',
         'Individual factors within each category (e.g. average_daily_balance, '
         'public_records_count, monthly_revenue_trend)'],
        ['rel_score_factor_evaluation_config', 'case-service PostgreSQL',
         'Links factors to evaluation configs — defines weightage per factor'],
        ['score_evaluation_config', 'case-service PostgreSQL',
         'Range-based scoring rules: IF value >= X AND value < Y THEN score = Z'],
        ['score_decision_matrix', 'case-service PostgreSQL',
         'Per-customer thresholds: score range -> risk_level + decision '
         '(AUTO_APPROVED / UNDER_MANUAL_REVIEW / AUTO_REJECTED)'],
        ['score_decision_history', 'case-service PostgreSQL',
         'History of decisioning config changes per customer'],
        ['score_config_history', 'case-service PostgreSQL',
         'History of score weightage config changes — referenced as "Model 3.0"'],
        ['business_score_triggers', 'case-service PostgreSQL',
         'One row per scoring run trigger: (business_id, customer_id, trigger_type)'],
        ['business_scores', 'case-service PostgreSQL',
         'Final score per trigger: weighted_score_850, risk_level, score_decision, status'],
        ['business_score_factors', 'case-service PostgreSQL',
         'Per-factor results: category_id, factor_id, value, score_100, weighted_score_100, score_850'],
        ['business_score_history', 'case-service PostgreSQL',
         'Audit trail of score status changes (PROCESSING -> COMPLETED / FAILED)'],
        ['data_current_scores', 'case-service PostgreSQL',
         'Latest score per (business_id, customer_id) — what the UI reads'],
        ['score_inputs', 'case-service PostgreSQL',
         'Snapshot of all raw data fed into the score model (JSON): inputs, inputs_raw'],
    ],
    col_widths=[2.6, 2.2, 4.8],
)

H2('Score Categories and Their Data Sources')
table(
    ['Category', 'Code', 'Integration source', 'Platform IDs', 'What is measured'],
    [
        ['Banking', 'banking',
         'Plaid bank account connection',
         '1 (Plaid), 34 (Manual Banking)',
         'Average daily balance, NSF count, overdraft frequency, '
         'monthly deposit volume, revenue consistency'],
        ['Public Records', 'public_records',
         'Middesk (SOS), Trulioo KYB, NPI, Adverse Media',
         '16 (Middesk), 38 (Trulioo), 27 (Adverse Media), 28 (NPI)',
         'SOS registration status, watchlist hits, adverse media hits, '
         'NPI validity, business age'],
        ['Financial Strength', 'financial_strength',
         'Rutter (accounting integrations), Tax Status',
         '5-13 (Rutter platforms), 15 (Tax Status)',
         'Revenue trends, gross profit, accounts receivable, '
         'tax compliance, IRS balance'],
    ],
    col_widths=[1.5, 1.0, 2.3, 1.6, 3.6],
)

H2('Auto-Approve vs Manual Review vs Archived — Who Decides What')
table(
    ['Case status', 'ID', 'Trigger', 'Who sets it', 'Source'],
    [
        ['AUTO_APPROVED', '6',
         'Worth Score >= customer\'s auto-approve threshold\n'
         '(score_decision_matrix.decision = "AUTO_APPROVED")',
         'Automatically by manual-score-service after score calculation',
         'manual-score-service -> Kafka AI_SCORES -> case-service '
         'UPDATE data_cases SET status_id = 6'],
        ['AUTO_REJECTED', '13',
         'Worth Score below customer\'s auto-reject threshold\n'
         '(score_decision_matrix.decision = "AUTO_REJECTED")',
         'Automatically by manual-score-service',
         'Same as above — status_id = 13'],
        ['UNDER_MANUAL_REVIEW', '4',
         'Score in the medium-risk range (no automatic decision)\n'
         'OR analyst manually triggers review',
         'Automatically by scoring OR manually by analyst',
         'score_decision_matrix.decision = "UNDER_MANUAL_REVIEW"\n'
         'OR PATCH /cases/{id} { status: "UNDER_MANUAL_REVIEW" }'],
        ['MANUALLY_APPROVED', '5',
         'Analyst approves the case after reviewing',
         'Human analyst clicking Approve in admin UI',
         'PATCH /cases/{id} { status: "MANUALLY_APPROVED" }\n'
         'Written to data_cases.status_id = 5'],
        ['MANUALLY_REJECTED', '8',
         'Analyst rejects the case',
         'Human analyst clicking Reject in admin UI',
         'PATCH /cases/{id} { status: "MANUALLY_REJECTED" }'],
        ['ARCHIVED', '9',
         'Case is archived (no longer active)',
         'Manual analyst action or automated archiving policy',
         'PATCH /cases/{id} { status: "ARCHIVED" }'],
        ['PENDING_DECISION', '10',
         'Score calculated but no auto decision yet',
         'Automatic interim state during scoring pipeline',
         'Set by manual-score-service before final decision'],
        ['INFORMATION_REQUESTED', '11',
         'Analyst requested more info from applicant',
         'Analyst action via admin UI',
         'PATCH /cases/{id} + creates case_info_request record'],
        ['RISK_ALERT', '14',
         'Risk monitoring triggered a new alert',
         'Automatic — risk monitoring rules engine',
         'integration-service risk-alerts module -> Kafka -> data_cases'],
        ['INVESTIGATING', '15',
         'Analyst is actively investigating a risk alert',
         'Manual analyst action',
         'PATCH /cases/{id} { status: "INVESTIGATING" }'],
        ['ESCALATED', '17',
         'Case escalated to senior analyst or compliance',
         'Manual analyst action',
         'PATCH /cases/{id} { status: "ESCALATED" }'],
    ],
    col_widths=[2.1, 0.5, 2.4, 2.1, 2.9],
)

H3('SQL queries for score data')
code_block([
    '-- Current score for a business:',
    'SELECT bs.weighted_score_850, bs.risk_level, bs.score_decision, bs.status',
    'FROM data_current_scores dcs',
    'JOIN business_scores bs ON bs.id = dcs.score_id',
    "WHERE dcs.business_id = '<uuid>' AND dcs.customer_id = '<uuid>';",
    '',
    '-- Per-factor breakdown:',
    'SELECT sc.code AS category, scf.code AS factor,',
    '       bsf.value, bsf.score_100, bsf.weighted_score_100, bsf.score_850',
    'FROM business_score_factors bsf',
    'JOIN score_category_factors scf ON scf.id = bsf.factor_id',
    'JOIN score_categories sc ON sc.id = bsf.category_id',
    "WHERE bsf.score_id = '<score_uuid>'",
    'ORDER BY sc.code, scf.code;',
    '',
    '-- Decision matrix for a customer:',
    'SELECT range_start, range_end, risk_level, decision',
    'FROM score_decision_matrix',
    "WHERE customer_id = '<uuid>' ORDER BY range_start;",
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — KYB TAB
# ════════════════════════════════════════════════════════════════════════════

H1('Section 4 — KYB Tab: What It Is and All Sub-tabs', c=PURPLE, pb=False)

callout(
    'KYB = Know Your Business. This tab aggregates data about the business entity itself '
    '(not the owners — that is KYC). It shows data from ALL vendor integrations that '
    'verify the business: Middesk (SOS filings), OpenCorporates (global registry), '
    'ZoomInfo (commercial data), Equifax (bureau), Trulioo (KYB API), and SERP/AI scraping.'
)

H2('KYB Sub-tabs and Their Sources')
table(
    ['Sub-tab', 'What is shown', 'Primary source', 'Platform ID(s)'],
    [
        ['Background',
         'Map with business location, Business Details panel (name, DBA, address, age, '
         'revenue, employees, corporation type, phone, email, MBE/WBE/VBE flags), '
         'Industry section (NAICS, MCC), NPI section',
         'Pipeline A Fact Engine winner from all sources.\n'
         'Map: Google Maps embed using business address.\n'
         'Business details: rds_warehouse_public.facts JSONB.',
         '16 (Middesk), 23 (OC), 24 (ZI), 17 (EFX), 38 (Trulioo), 31 (AI)'],
        ['Business Registration',
         'SOS filing: registered name, status (Active/Inactive), registration date, '
         'registered agent, state of formation, filing type',
         'Middesk SOS lookup — live API call per business submission.\n'
         'Stored in integration_data.business_entity_verification.',
         '16 (Middesk)'],
        ['Contact Information',
         'Phone numbers, email addresses, website URL, LinkedIn, social media',
         'SERP scraping + ZoomInfo + Middesk.\n'
         'rds_warehouse_public.facts: phone_number, email, official_website facts.',
         '22 (SERP), 24 (ZI), 16 (Middesk)'],
        ['Website',
         'Website screenshot, Worth website analysis score, domain age, '
         'SSL status, malware flag, Google Business rating',
         'WORTH_WEBSITE_SCANNING (platform 30) + Google Business Reviews (19/20).\n'
         'Stored in integration_data.request_response.',
         '30 (WorthWebsiteScanning), 19 (Google Places), 20 (Google Business)'],
        ['Watchlists',
         'OFAC SDN, EU sanctions, UN consolidated, PEP lists, adverse media hits',
         'Trulioo KYB watchlist screening.\n'
         'Stored in integration_data.business_entity_verification (watchlist field).\n'
         'Also: Adverse Media integration (platform 27).',
         '38 (Trulioo), 42 (Trulioo PSC), 27 (Adverse Media)'],
    ],
    col_widths=[1.8, 3.0, 2.9, 1.9],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — KYB BUSINESS DETAILS FIELDS
# ════════════════════════════════════════════════════════════════════════════

H1('Section 5 — KYB Background Tab: Every Field Lineage', c=PURPLE, pb=False)

H2('Business Details Section — Field-by-Field Lineage')
table(
    ['Field label', 'Example value', 'Fact name in facts table', 'Winner priority',
     'Source table (winning vendor)'],
    [
        ['Provided Business Name', 'Lisa\'s Nail Salon',
         'Data from application — NOT a fact',
         'n/a — submitted by applicant',
         'data_businesses.name (case-service PostgreSQL)'],
        ['Legal Business Name', 'Lisa\'s Nail Salon',
         'legal_business_name (rds_warehouse_public.facts)',
         'Middesk SOS filing (highest confidence)\nFallback: OC, ZI, EFX',
         'integration_data.business_entity_verification.name (Middesk)\n'
         'OR zoominfo_standard_ml_2.zi_c_name (ZoomInfo)'],
        ['DBA', 'N/A',
         'dba_names (facts)',
         'Middesk SOS -> OC -> ZI',
         'integration_data.business_entity_verification.dba_names (Middesk)'],
        ['Business Address', '7554 Louisiana 182, Morgan City, LA 70380',
         'address_line_1, address_city, address_state, address_postal_code (facts)',
         'Middesk SOS -> OC -> ZI -> EFX\n'
         'Rule: factWithHighestConfidence() then weightedFactSelector()',
         'integration_data.business_entity_address_source\n'
         'OR equifax_us_standardized.efx_address*'],
        ['Mailing Address', 'N/A',
         'mailing_address (facts)',
         'Middesk SOS only (when different from registered)',
         'integration_data.business_entity_verification.mailing_address'],
        ['Business Age', '08/18/2004 (21 years)',
         'year_of_incorporation / date_of_incorporation (facts)',
         'Middesk SOS (formation date from state filing)\n'
         'Fallback: OC incorporation_date, ZI zi_c_founded',
         'integration_data.business_entity_verification.incorporation_date\n'
         'OR open_corporates_standard_ml_2.oc_incorporation_date'],
        ['Annual Revenue', 'N/A (no banking/accounting connected)',
         'annual_revenue (facts)',
         'Equifax (efx_annual_sales) -> ZoomInfo (zi_c_annual_sales)\n'
         'Also from Rutter accounting integration if connected',
         'equifax_us_standardized.efx_annual_sales\n'
         'OR zoominfo_standard_ml_2.zi_c_annual_sales'],
        ['Avg. Annual Revenue', 'N/A',
         'Computed from Plaid bank statements if banking connected',
         'Plaid (platform_id=1)',
         'plaid_transactions or bank_statements table'],
        ['Net Income', 'N/A',
         'net_income (facts)',
         'Rutter accounting -> Equifax',
         'accounting_statements table (Rutter)'],
        ['Corporation Type', 'Private',
         'corporation_type (facts)',
         'Middesk SOS -> OC',
         'integration_data.business_entity_verification.entity_type'],
        ['Number of Employees', '1',
         'employee_count (facts)',
         'ZoomInfo (zi_c_employee_count) -> Equifax (efx_employees)',
         'zoominfo_standard_ml_2.zi_c_employee_count\n'
         'OR equifax_us_standardized.efx_employees'],
        ['Business Phone Number', 'N/A',
         'phone_number (facts)',
         'Middesk -> SERP -> ZoomInfo',
         'integration_data.business_entity_verification.phone_numbers\n'
         'OR integration_data.request_response (SERP, platform=22)'],
        ['Business Email', 'N/A',
         'email (facts)',
         'Middesk -> SERP -> ZoomInfo',
         'integration_data.request_response (SERP, platform=22)'],
        ['MBE / WBE / VBE flags', 'N/A',
         'minority_business_enterprise, woman_owned_business, veteran_owned_business (facts)',
         'Equifax -> ZoomInfo (these are vendor-provided flags)',
         'equifax_us_standardized.efx_minority_business_enterprise\n'
         'OR zoominfo_standard_ml_2.zi_minority_business_enterprise'],
    ],
    col_widths=[1.8, 1.6, 2.0, 2.0, 2.2],
    fs=8.5,
)

H2('Industry Section — NAICS, MCC, and NPI')
table(
    ['Field', 'Example', 'Fact name', 'Winner priority (see Industry Facts doc for full detail)'],
    [
        ['Industry Name', 'Other Services (except Public Administration)',
         'industry (facts)',
         'Derived from NAICS code sector (first 2 digits)\n'
         '812xxx -> sector 81 -> "Other Services" (core_business_industries lookup)'],
        ['NAICS Code', '812113',
         'naics_code (facts)',
         'Middesk (2.0) > OC (0.9) > ZoomInfo (0.8) > Trulioo (0.8) > Equifax (0.7) > AI (0.1)\n'
         'Validation: must exist in core_naics_code table; else replaced with 561499'],
        ['NAICS Description', 'Nail Salons',
         'naics_description (facts)',
         'Lookup from core_naics_code WHERE code = "812113"'],
        ['MCC Code', '7230',
         'mcc_code (facts)',
         'Path A: vendor returns MCC directly (Middesk/Trulioo)\n'
         'Path B: NAICS->MCC crosswalk: rel_naics_mcc WHERE naics_code="812113" -> 7230\n'
         'Path C: AI enrichment'],
        ['MCC Description', 'Beauty and Barber Shops',
         'mcc_description (facts)',
         'Lookup: core_mcc_code WHERE code = "7230"'],
    ],
    col_widths=[1.8, 2.2, 1.8, 3.8],
)

H2('NPI Section')
body(
    'NPI = National Provider Identifier (for healthcare businesses only). '
    'Integrated via NPI lookup (platform_id = 28). '
    'For Lisa\'s Nail Salon: all N/A + "No records found" = NPI registry returned '
    'no matching healthcare provider for this business. '
    'NPI data is fetched from NPPES (CMS public registry) via integration-service/lib/npi/npi.ts.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — KYC TAB
# ════════════════════════════════════════════════════════════════════════════

H1('Section 6 — KYC Tab: Owner Identity Verification', c=PURPLE, pb=False)

callout(
    'KYC = Know Your Customer. This tab shows identity verification for each beneficial owner / '
    'control person. Data comes from two main sources:\n\n'
    '1. SELF-REPORTED: Fields submitted by the applicant via the onboarding form '
    '(yellow highlight = data entered by the business owner).\n\n'
    '2. VERIFIED: The green "Match" badge means a third-party verification service '
    'confirmed the submitted data matches official records. '
    '"Unverified" means IDV did not complete or was disabled.'
)

H2('KYC Owner Panel — "Thanh Nguyen" Example (Unverified — IDV Disabled)')
table(
    ['Field', 'Value shown', 'Yellow highlight?', 'Source', 'Table'],
    [
        ['Legal First Name', 'Thanh', 'YES — applicant-submitted',
         'Applicant submitted during onboarding form',
         'data_applicants.first_name OR data_owners.first_name'],
        ['Legal Last Name', 'Nguyen', 'YES',
         'Applicant-submitted', 'data_applicants.last_name'],
        ['Date of Birth', '1971-07-21', 'YES',
         'Applicant-submitted', 'data_applicants.date_of_birth OR data_owners.dob'],
        ['Social Security Number', 'N/A', 'N/A',
         'Not provided / masked for display. If provided, stored encrypted.',
         'data_owners.ssn (AES-256 encrypted in PostgreSQL)'],
        ['Home Address', '700 Cypress St, Morgan City, LA 70380', 'YES',
         'Applicant-submitted home address',
         'data_owner_addresses.line_1, city, state, postal_code'],
        ['Phone Number', '+1 (985) 519-5173', 'YES',
         'Applicant-submitted', 'data_applicants.phone OR data_owners.phone'],
        ['Email Address', 'gino5@bellsouth.net', 'YES',
         'Applicant-submitted', 'data_applicants.email OR data_owners.email'],
        ['Job Title', 'Partner', 'YES',
         'Applicant-submitted', 'data_owners.title'],
        ['Ownership', '25.00%', 'YES',
         'Applicant-submitted', 'data_owners.ownership_percentage'],
    ],
    col_widths=[1.8, 2.0, 1.3, 2.2, 2.7],
)

body(
    '"IDV Disabled" next to "Unverified" in the KYC header means the customer (Worldpay/sandbox) '
    'has Identity Verification (IDV) disabled for this case. '
    'IDV status is stored in data_integration_settings.settings JSONB '
    'under identity_verification.status = "INACTIVE".'
)

H2('KYC Owner Panel — "Leslie Knope" Example (Verified — Green "Match" Badges)')
callout(
    'The green "Match" badge on each field means the submitted value was verified '
    'against a third-party data source and MATCHED. This is NOT the XGBoost entity-matching model. '
    'Match badges come from identity verification services — primarily Trulioo (person screening) '
    'or Plaid IDV (document + facial recognition). The matching is done field-by-field.'
)

H2('What Produces the "Match" Badge — The Full Process')
table(
    ['Step', 'What happens', 'Service', 'Source code'],
    [
        ['1', 'Applicant submits owner info via onboarding form',
         'case-service',
         'POST /businesses/{id}/owners -> data_owners table'],
        ['2', 'integration-service triggers identity verification\n'
         'based on customer config (IDV enabled = ACTIVE)',
         'integration-service',
         'src/core/businessOnboarding/businessOnboardingManager.ts\n'
         'platform_id = 18 (Plaid IDV) or 38 (Trulioo) or 3 (Persona)'],
        ['3A — Plaid IDV path',
         'Plaid IDV: applicant completes photo ID + selfie flow\n'
         'Plaid compares submitted name/DOB/address to government ID\n'
         'Returns per-field match status',
         'Plaid (platform_id=18)',
         'integration-service/lib/plaid/plaid.ts\n'
         'Stores result in: data_identity_verifications table\n'
         'IDV_STATUS: SUCCESS=1, PENDING=2, FAILED=99'],
        ['3B — Trulioo PSC path',
         'Trulioo person screening: submits owner details to\n'
         'Trulioo API, which checks against credit bureau / '
         'government databases per country\n'
         'Returns DatasourceFields with Status="match"|"nomatch"|"missing"',
         'Trulioo (platform_id=38 / 42 for PSC)',
         'integration-service/lib/trulioo/person/truliooPersonVerificationProcessor.ts\n'
         'Stored in: integration_data.business_entity_verification\n'
         'Field match: DatasourceFields[].Status per FieldName'],
        ['4', 'case-service reads IDV / Trulioo results and determines\n'
         'per-field Match status for UI display',
         'case-service Fact Engine',
         'integration-service/lib/facts/kyb/index.ts\n'
         'idv_status fact: counts SUCCESS/PENDING/FAILED IDV records\n'
         'idv_passed_boolean fact: true if at least 1 SUCCESS'],
        ['5', '"Match" badge shown in UI when field is verified',
         'Frontend React app',
         'Field-level badge: green "Match" = field confirmed by IDV/Trulioo\n'
         'No badge = field not verified\n'
         '"Verified" header badge = ALL required fields passed'],
    ],
    col_widths=[0.4, 3.2, 1.8, 4.2],
)

H2('Per-Field Match Logic — Trulioo DatasourceFields')
code_block([
    '// integration-service/lib/trulioo/common/utils.ts',
    '// Trulioo returns per-field match status in DatasourceResults:',
    '//   DatasourceResults[].DatasourceFields[{ FieldName, Status }]',
    '// Status values: "match" | "nomatch" | "missing"',
    '',
    '// Address matching example (Rule: aligned with product 2026-02-13):',
    'extractAddressMatchStatusFromDatasourceFields(clientData):',
    '  - Only use "Comprehensive View" datasource (Trulioo\'s consolidated result)',
    '  - If ANY address field has status "nomatch" -> overall "nomatch"',
    '  - If all address fields are "match" or "missing" -> overall "match"',
    '  - Falls back to all datasources if Comprehensive View not found',
    '',
    '// Stored in: integration_data.business_entity_address_source.status',
    '//   ("match" -> address_verification review task status = "success")',
    '//   ("nomatch" -> status = "failure")',
])

H2('KYC Fields: Self-Reported (Yellow) vs Verified (Green Match)')
table(
    ['Field', 'Yellow highlight meaning', 'Green Match meaning', 'Match source'],
    [
        ['Legal First Name', 'Submitted by applicant in form',
         'First name matches government ID or bureau record',
         'Trulioo DatasourceFields.FirstGivenName.Status = "match"\n'
         'OR Plaid IDV name comparison'],
        ['Legal Last Name', 'Submitted by applicant',
         'Last name matches', 'Same as above'],
        ['Date of Birth', 'Submitted by applicant',
         'DOB matches government ID',
         'Trulioo DatasourceFields.DayOfBirth + MonthOfBirth + YearOfBirth\n'
         'OR Plaid IDV DOB comparison'],
        ['Social Security Number', 'Applicant-submitted (masked XXX-XX-XXXX)',
         'SSN/TIN verified against bureau',
         'Trulioo DatasourceFields.NationalId.Status = "match"\n'
         'OR Equifax TIN verification'],
        ['Home Address', 'Applicant-submitted',
         'Address matches registered address in bureau/registry',
         'Trulioo extractAddressMatchStatusFromDatasourceFields()\n'
         '"Comprehensive View" datasource address fields'],
        ['Phone Number', 'Applicant-submitted',
         'Phone number verified as belonging to this person',
         'Trulioo DatasourceFields (varies by country)'],
        ['Email Address', 'Applicant-submitted',
         'Email verified (typically not a Match field — shown as submitted only)',
         'No verification service for email — shown as submitted'],
    ],
    col_widths=[1.7, 2.2, 2.2, 3.5],
)

H2('"Verified" vs "Unverified" Header Badge')
table(
    ['Badge', 'Meaning', 'Condition', 'Source'],
    [
        ['✅ Verified',
         'Owner identity fully verified',
         'IDV or Trulioo completed successfully\n'
         'idv_passed_boolean = true (integration-service facts/kyb/index.ts)\n'
         'OR Trulioo screening status = "COMPLETED" with match results',
         'data_identity_verifications.status = "SUCCESS" (IDV_STATUS.SUCCESS = 1)\n'
         'OR integration_data.business_entity_verification.status = "completed"'],
        ['⚠️ Unverified — IDV Disabled',
         'Identity verification is not active for this customer/case',
         'integration_data.data_integration_settings.settings\n'
         '-> identity_verification.status = "INACTIVE"\n'
         'OR IDV never triggered',
         'data_integration_settings table in integration-service PostgreSQL'],
        ['🔴 Unverified',
         'IDV was attempted but failed or not completed',
         'IDV_STATUS = FAILED (99) or EXPIRED (4)\n'
         'OR Trulioo screening returned nomatch on critical fields',
         'data_identity_verifications.status_id = 99\n'
         'OR integration_data.business_entity_verification.status = "failed"'],
    ],
    col_widths=[1.8, 2.2, 2.8, 2.8],
)

H2('"Control Person" Badge')
body(
    '"Control Person" next to the owner name means this owner has been flagged as the '
    'primary controlling person of the business (e.g., CEO, Managing Member). '
    'This designation comes from the applicant\'s submitted ownership data: '
    'data_owners.is_control_person = true, set during the onboarding form '
    'when the applicant marks themselves as the control person.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ALL CASE STATUSES
# ════════════════════════════════════════════════════════════════════════════

H1('Section 7 — Case Statuses: All 20 States', c=PURPLE, pb=False)

callout(
    'Case status is stored in data_cases.status_id (integer FK to CASE_STATUS constants). '
    'Every status transition is written by case-service and logged in case_audit_logs. '
    'Source: case-service/src/constants/case-status.constant.ts'
)

table(
    ['Status', 'ID', 'Set by', 'Automatic?', 'Trigger', 'Pipeline step'],
    [
        ['CREATED',          '20', 'case-service API',          'YES',
         'Case object created (POST /businesses/customers/{id})',
         'Immediately on submission, before any integrations fire'],
        ['INVITED',          '1',  'case-service (deprecated)',  'YES',
         'Onboarding invitation sent to applicant',
         'Legacy — being phased out'],
        ['INVITE_EXPIRED',   '2',  'Scheduled job (deprecated)', 'YES',
         'Invite not accepted within TTL',
         'Legacy — being phased out'],
        ['ONBOARDING',       '3',  'case-service API',          'YES',
         'Applicant started filling the onboarding form',
         'First form field submitted'],
        ['SUBMITTED',        '12', 'case-service API',          'YES',
         'Applicant submitted the completed onboarding form',
         'POST /businesses/{id}/submit -> triggers all integrations via Kafka'],
        ['SCORE_CALCULATED', '7',  'manual-score-service',      'YES',
         'Worth Score computation completed',
         'After all integration tasks finish + GENERATE_AI_SCORE Kafka event processed'],
        ['PENDING_DECISION', '10', 'manual-score-service',      'YES',
         'Score ready but decision rule not yet applied',
         'Interim state during scoring pipeline'],
        ['AUTO_APPROVED',    '6',  'manual-score-service',      'YES',
         'Score >= auto-approve threshold (score_decision_matrix)',
         'AUTO — no human needed. Most common final state for Sandbox test data.'],
        ['AUTO_REJECTED',    '13', 'manual-score-service',      'YES',
         'Score <= auto-reject threshold',
         'AUTO — no human needed'],
        ['UNDER_MANUAL_REVIEW', '4', 'manual-score-service OR analyst', 'BOTH',
         'Score in manual review range OR analyst triggered review',
         'Requires human analyst decision'],
        ['INFORMATION_REQUESTED', '11', 'Analyst', 'NO',
         'Analyst needs more info from applicant',
         'Sends notification to applicant; creates case_info_request record'],
        ['INFORMATION_UPDATED', '19', 'case-service API',       'YES',
         'Applicant submitted requested information',
         'Re-triggers scoring pipeline'],
        ['MANUALLY_APPROVED', '5', 'Analyst',                   'NO',
         'Analyst clicks Approve after manual review',
         'PATCH /cases/{id} { status: "MANUALLY_APPROVED" }'],
        ['MANUALLY_REJECTED', '8', 'Analyst',                   'NO',
         'Analyst clicks Reject',
         'PATCH /cases/{id} { status: "MANUALLY_REJECTED" }'],
        ['ARCHIVED',          '9', 'Analyst OR automated policy','BOTH',
         'Case is closed/archived',
         'PATCH /cases/{id} { status: "ARCHIVED" }'],
        ['RISK_ALERT',        '14', 'integration-service (risk monitoring)', 'YES',
         'Risk monitoring rule triggered a new alert (ongoing monitoring)',
         'Kafka event from risk-alerts module -> data_cases.status_id = 14'],
        ['INVESTIGATING',     '15', 'Analyst',                  'NO',
         'Analyst is actively investigating a risk alert',
         'PATCH /cases/{id} { status: "INVESTIGATING" }'],
        ['ESCALATED',         '17', 'Analyst',                  'NO',
         'Escalated to compliance/senior review',
         'PATCH /cases/{id} { status: "ESCALATED" }'],
        ['PAUSED',            '18', 'Analyst',                  'NO',
         'Investigation paused',
         'PATCH /cases/{id} { status: "PAUSED" }'],
        ['DISMISSED',         '16', 'Analyst',                  'NO',
         'Risk alert dismissed after investigation',
         'PATCH /cases/{id} { status: "DISMISSED" }'],
    ],
    col_widths=[1.9, 0.4, 1.7, 1.0, 2.2, 2.4],
    fs=8.5,
)

H3('Which pipeline controls each status transition?')
table(
    ['Phase', 'Pipeline / service', 'Statuses it writes'],
    [
        ['Submission', 'case-service (real-time API)',
         'CREATED -> ONBOARDING -> SUBMITTED'],
        ['Integration execution', 'integration-service (real-time)',
         'No status change — fires vendor integrations in background'],
        ['Score computation', 'manual-score-service (Kafka worker)',
         'SCORE_CALCULATED -> PENDING_DECISION -> AUTO_APPROVED / AUTO_REJECTED / UNDER_MANUAL_REVIEW'],
        ['Analyst decisions', 'case-service (API, manual)',
         'MANUALLY_APPROVED / MANUALLY_REJECTED / ARCHIVED / INFORMATION_REQUESTED'],
        ['Ongoing risk monitoring', 'integration-service risk-alerts module',
         'RISK_ALERT -> INVESTIGATING -> ESCALATED / PAUSED / DISMISSED'],
        ['Pipeline B (batch Redshift)', 'warehouse-service',
         'NO status changes — Pipeline B is analytics only, never changes case_status'],
    ],
    col_widths=[1.8, 3.0, 4.8],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SANDBOX vs PRODUCTION DEEP DIVE
# ════════════════════════════════════════════════════════════════════════════

H1('Section 8 — Sandbox vs Production: Complete Deep Dive', c=PURPLE, pb=False)

H2('Where is Sandbox/Production Set?')
lineage([
    'STEP 1: Worth AI admin provisions a new customer (tenant)',
    '  -> Admin sets customer_type = "PRODUCTION" or "SANDBOX"',
    '  -> Stored in: data_customers.customer_type (case-service PostgreSQL)',
    '',
    'STEP 2: Kafka event fired to integration-service',
    '  -> Topic: integration_settings (or similar)',
    '  -> Payload: { customerID, customer_type: "SANDBOX", settings: { ... } }',
    '  -> Schema validated: Joi.string().valid("PRODUCTION", "SANDBOX")',
    '  -> Source: integration-service/src/messaging/kafka/consumers/handlers/schema.ts line 323',
    '',
    'STEP 3: integration-service stores per-customer settings',
    '  -> Table: integration_data.data_integration_settings',
    '     (customer_id, settings JSONB)',
    '  -> JSONB example: {',
    '       "equifax":              { "status": "ACTIVE",   "mode": "SANDBOX" },',
    '       "identity_verification":{ "status": "INACTIVE", "mode": "PRODUCTION" },',
    '       "bjl":                  { "status": "INACTIVE", "mode": "PRODUCTION" }',
    '     }',
    '',
    'STEP 4: When a business is submitted under this customer',
    '  -> integration-service reads data_integration_settings for customer',
    '  -> Uses mode from settings to select PRODUCTION or SANDBOX strategy',
    '  -> strategyPlatformFactory.ts loads the correct strategy class',
    '',
    'STEP 5: The UI reads customer_type from the API',
    '  -> GET /customers/{id} -> returns customer_type',
    '  -> Frontend shows yellow "Sandbox" badge or green "Production" badge',
])

H2('Per-Integration Mode Settings (Default Values)')
table(
    ['Integration', 'Default status', 'Default mode', 'Available modes'],
    [
        ['Equifax (personal credit)', 'ACTIVE', 'PRODUCTION', 'PRODUCTION | SANDBOX | DISABLE'],
        ['GIACT gVerify', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | SANDBOX | DISABLE'],
        ['GIACT gAuthenticate', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | SANDBOX | DISABLE'],
        ['Identity Verification (Plaid IDV)', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | SANDBOX | DISABLE'],
        ['BJL (Bankruptcies, Judgements, Liens)', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | DISABLE'],
        ['Website analysis', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | DISABLE'],
        ['NPI', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | DISABLE'],
        ['Adverse Media', 'INACTIVE', 'PRODUCTION', 'PRODUCTION | DISABLE'],
        ['Middesk (KYB)', 'Always runs', 'n/a — no sandbox mode', 'Same API for all customers'],
        ['OpenCorporates', 'Always runs', 'n/a — Redshift query', 'Same Redshift data always'],
        ['ZoomInfo', 'Always runs', 'n/a — Redshift query', 'Same Redshift data always'],
        ['Trulioo (KYB)', 'Varies by customer', 'PRODUCTION typically', 'Depends on customer config'],
    ],
    col_widths=[2.8, 1.5, 1.5, 4.2],
)
callout(
    'Source: integration-service/src/constants/customer-integration-settings.constants.ts\n'
    'DEFAULT_CUSTOMER_INTEGRATION_SETTINGS defines the default mode for every integration.\n'
    'These can be overridden per customer via the admin portal integration settings page.'
)

H2('What Data Is Real vs Simulated in Sandbox?')
table(
    ['Data type', 'Sandbox behavior', 'Production behavior'],
    [
        ['Equifax credit data', 'Equifax sandbox API — returns test credit profiles\n'
         'Not real consumer credit data', 'Real Equifax bureau data'],
        ['Plaid bank connections', 'Plaid sandbox — fake institutions, test transactions\n'
         'No real bank accounts', 'Real Plaid bank connections'],
        ['GIACT bank verification', 'GIACT sandbox API — returns test pass/fail results',
         'Real GIACT bank account verification'],
        ['Middesk SOS filings', 'Same API — but test businesses may use\n'
         'Middesk\'s own test merchant records',
         'Real Secretary of State filings'],
        ['ZoomInfo company data', 'Same Redshift pre-loaded data\n'
         '(ZI data loaded once, used for all)',
         'Same — no distinction'],
        ['OpenCorporates data', 'Same Redshift pre-loaded data',
         'Same — no distinction'],
        ['Trulioo KYB watchlists', 'Depends on customer config\n'
         'May use Trulioo test environment',
         'Live Trulioo watchlist screening'],
        ['Worth Score', 'Computed identically to production\n'
         'but based on simulated/sandbox integration data',
         'Computed from real vendor data'],
        ['Case decisions', 'AUTO_APPROVED/UNDER_MANUAL_REVIEW applied same\n'
         'but should not drive real credit decisions',
         'Real underwriting decisions'],
    ],
    col_widths=[2.2, 3.7, 3.7],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — PUBLIC RECORDS, BANKING, ACCOUNTING
# ════════════════════════════════════════════════════════════════════════════

H1('Section 9 — Other Tabs: Public Records, Banking, Accounting, Taxes, Documents',
   c=PURPLE, pb=False)

H2('Public Records Tab')
table(
    ['Panel', 'What is shown', 'Source vendor', 'Platform ID', 'Tables'],
    [
        ['SOS Registration',
         'Registered business name, state, status (Active/Inactive),\n'
         'registration date, registered agent, file number, entity type',
         'Middesk — live SOS registry API call per business',
         '16',
         'integration_data.business_entity_verification\n'
         '(registration_number, incorporation_date, entity_type, sos_status)'],
        ['Watchlists',
         'OFAC SDN, EU Consolidated, UN Consolidated, PEP lists,\n'
         'Adverse Media — hit/no-hit per list',
         'Trulioo KYB screening\n'
         'Also: ADVERSE_MEDIA integration (platform 27)',
         '38, 27',
         'integration_data.business_entity_verification.watchlist_hits JSONB\n'
         'OR data_adverse_media table'],
        ['NPI (Healthcare)',
         'NPI provider details, taxonomy, state license',
         'NPPES public API (CMS) via integration-service/lib/npi/npi.ts',
         '28',
         'integration_data.request_response (platform_id=28)\n'
         'Fact: npi_data in rds_warehouse_public.facts'],
        ['Adverse Media',
         'News articles mentioning the business negatively',
         'Adverse Media integration (dedicated search)',
         '27',
         'integration_data.request_response (platform_id=27)'],
    ],
    col_widths=[1.8, 2.4, 2.0, 0.9, 2.5],
)

H2('Banking Tab')
table(
    ['Element', 'What is shown', 'Source', 'Platform ID', 'Tables'],
    [
        ['Bank accounts list', 'Institution, account type, balance, connection status',
         'Plaid bank connection',
         '1 (Plaid)',
         'data_connections (status), plaid_accounts, plaid_balances'],
        ['Bank statements', 'PDF statements uploaded or pulled from Plaid',
         'Plaid OR manual upload (MANUAL_BANKING=34)',
         '1, 34',
         'bank_statements table, s3_documents (S3 storage)'],
        ['Transaction analysis', 'NSF count, overdraft frequency, deposit volume',
         'Plaid transactions feed',
         '1',
         'plaid_transactions -> aggregated into score_inputs for Banking category'],
        ['GIACT verification',
         'Bank account open/active/good standing status',
         'GIACT gVerify API',
         '26',
         'integration_data.request_response (platform_id=26)'],
    ],
    col_widths=[1.8, 2.5, 1.8, 1.0, 2.5],
)

H2('Accounting Tab')
table(
    ['Element', 'What is shown', 'Source', 'Platform ID', 'Tables'],
    [
        ['QuickBooks / Xero connection',
         'Revenue, expenses, profit/loss, accounts receivable',
         'Rutter (unified accounting API)\n'
         'Supports: QBO(5), Xero(6), Zoho(7), FreshBooks(8), Wave(10), Stripe(12), Square(13)',
         '5-13',
         'accounting_statements, data_connections'],
        ['Financial statements', 'P&L, balance sheet, cash flow',
         'Rutter -> normalized financial data',
         '5-13',
         'financial_statements table -> fed into Financial Strength score category'],
    ],
    col_widths=[1.8, 2.5, 2.7, 0.9, 1.7],
)

H2('Taxes Tab')
table(
    ['Element', 'What is shown', 'Source', 'Platform ID'],
    [
        ['IRS tax transcripts', 'Filed date, total sales, IRS balance, liens',
         'Tax Status integration (IRS e-services)',
         '15 (TAX_STATUS)'],
        ['Tax compliance',
         'Whether taxes are current, outstanding liabilities',
         'Tax Status API -> stored in integration_data.request_response',
         '15'],
    ],
    col_widths=[1.8, 3.0, 3.2, 1.6],
)

H2('Documents Tab')
body(
    'Shows all documents uploaded by the applicant or generated by the system '
    '(e.g. bank statements, tax forms, identity documents, e-signature agreements). '
    'Stored in S3 via integration-service/lib/aws. '
    'Document references stored in: s3_documents table (integration-service PostgreSQL) '
    'with metadata: (business_id, document_type, s3_key, created_at).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — WORTH SCORE vs PENDING
# ════════════════════════════════════════════════════════════════════════════

H1('Section 10 — Why Some Cases Show Worth Score as "Pending"', c=PURPLE, pb=False)

callout(
    'The second KYC screenshot shows Case Results = "Pending" with Worth Score = "-" '
    'and a message "Integrations are currently processing. Data on this case is still '
    'populating. Please wait until this process is done before reviewing this case."\n\n'
    'This is NOT an error. It is the normal state between submission and score completion.'
)

H2('The "Pending" State — Why It Happens')
lineage([
    'T+0:00  Applicant submits onboarding form',
    '  case-service: data_cases.status_id = SUBMITTED (12)',
    '  Kafka: BUSINESS_SUBMITTED event fired',
    '       |',
    'T+0:01  integration-service receives Kafka event',
    '  Creates business_integration_tasks for each integration:',
    '    - Middesk SOS lookup (platform_id=16)         -> task created, status="pending"',
    '    - OpenCorporates query (platform_id=23)       -> task created',
    '    - ZoomInfo query (platform_id=24)             -> task created',
    '    - Equifax lookup (platform_id=17)             -> task created',
    '    - Trulioo KYB (platform_id=38)                -> task created',
    '    - NPI lookup (platform_id=28)                 -> task created (if enabled)',
    '    - Identity Verification (platform_id=18)      -> task created (if IDV enabled)',
    '       |',
    'T+0:02  to T+2:00  Integrations run in parallel',
    '  Each integration completes -> task.status = "completed" or "failed"',
    '  Each completion fires: Kafka UPDATE_INTEGRATION_DATA_FOR_SCORE',
    '       |',
    '  During this window: Case Results shows "Pending"',
    '  business_scores.status = PROCESSING',
    '  data_current_scores exists but score values are NULL',
    '       |',
    'T+2:00  Last integration completes',
    '  integrationsCompletionTracker detects all tasks done',
    '  Fires: Kafka GENERATE_AI_SCORE',
    '       |',
    'T+2:05  manual-score-service generates Worth Score',
    '  calculateBankingScore() + calculatePublicRecordsScore() + calculateFinancialScore()',
    '  -> weighted_score_850 = 819',
    '  -> UPDATE business_scores SET status = "COMPLETED"',
    '  -> UPDATE data_cases SET status_id = AUTO_APPROVED (6)',
    '       |',
    'T+2:06  UI refresh: Case Results shows "Low Risk — 819/850"',
])

H2('Tables Involved During Pending State')
table(
    ['Table', 'State during Pending', 'State after score completes'],
    [
        ['data_cases.status_id', 'SUBMITTED (12)', 'AUTO_APPROVED (6) or UNDER_MANUAL_REVIEW (4)'],
        ['business_scores.status', 'PROCESSING', 'COMPLETED'],
        ['business_scores.weighted_score_850', 'NULL', '819'],
        ['business_scores.risk_level', 'NULL', '"LOW"'],
        ['business_integration_tasks.status', 'pending / in_progress', 'completed / failed'],
        ['data_current_scores.score_id', 'Points to PROCESSING score', 'Points to COMPLETED score'],
    ],
    col_widths=[2.8, 2.4, 4.4],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 11 — COMPLETE INTEGRATION MAP
# ════════════════════════════════════════════════════════════════════════════

H1('Section 11 — Complete Integration & Pipeline Map', c=PURPLE, pb=False)

H2('Every Integration: Vendor → Platform ID → API → Kafka → Table → UI')
table(
    ['Vendor', 'Platform ID', 'Call type', 'Triggered by', 'Kafka event',
     'Primary storage table', 'UI tabs showing it'],
    [
        ['Middesk (KYB)', '16',
         'Live REST API call\n(SOS registry lookup)',
         'Business submission',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE\n+ facts.v1 naics_code/industry',
         'integration_data.business_entity_verification\n'
         'rds_warehouse_public.facts',
         'KYB Background, Business Registration,\nPublic Records, Industry, Worth Score (PR category)'],
        ['Equifax', '17',
         'PRODUCTION: Live Equifax API\nSANDBOX: Equifax sandbox API\n'
         '(via EquifaxProductionStrategy / EquifaxSandboxStrategy)',
         'Business submission',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE',
         'integration_data.request_response\n'
         'warehouse.equifax_us_standardized (Redshift, pre-loaded)\n'
         'rds_warehouse_public.facts (via Pipeline A)',
         'KYB Background (firmographic facts)\nKYC (owner fields if EFX personal)\n'
         'Accounting/Worth Score'],
        ['Plaid (banking)', '1',
         'OAuth bank connection\n(applicant completes flow)',
         'Applicant connects bank in onboarding',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE (banking category)',
         'data_connections, plaid_accounts,\nplaid_transactions, bank_statements',
         'Banking tab, Worth Score (banking category)'],
        ['Plaid IDV', '18',
         'PRODUCTION: Real Plaid IDV flow\n(photo ID + selfie)\n'
         'SANDBOX: Simulated IDV',
         'Owner submits identity info\n(if IDV enabled for customer)',
         'IDV_COMPLETED / IDV_FAILED',
         'data_identity_verifications\n(IDV_STATUS: SUCCESS=1, FAILED=99)',
         'KYC tab (Verified badge, Match per field)'],
        ['OpenCorporates', '23',
         'Redshift query (pre-loaded bulk data)\nNot a live API call',
         'Business submission',
         'facts.v1 naics_code',
         'datascience.open_corporates_standard_ml_2 (Redshift)\n'
         'rds_warehouse_public.facts (Pipeline A)',
         'KYB Background (industry, registration)\nKYB Business Registration'],
        ['ZoomInfo', '24',
         'Redshift query (pre-loaded bulk data)\nNot a live API call',
         'Business submission',
         'facts.v1 naics_code / firmographic facts',
         'datascience.zoominfo_standard_ml_2 (Redshift)\n'
         'rds_warehouse_public.facts (Pipeline A)',
         'KYB Background (employees, revenue, industry)'],
        ['Trulioo (KYB)', '38',
         'Live Trulioo KYB API\n(business + watchlist screening)',
         'Business submission',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE\n(public_records category)',
         'integration_data.business_entity_verification\n'
         '(watchlist, address match, entity data)',
         'KYB Background, KYB Watchlists,\nPublic Records, Worth Score (PR category)'],
        ['Trulioo PSC', '42',
         'Live Trulioo PSC API\n(person / UBO screening)',
         'Owner data submitted + PSC enabled',
         'PSC_COMPLETED',
         'integration_data.business_entity_verification\n'
         '(person screening results)',
         'KYC tab (owner watchlist hits)'],
        ['SERP (web scraping)', '22',
         'Google SERP + website scraping\n(no external API key — crawls web)',
         'Business submission',
         'facts.v1 (website, phone, email facts)',
         'integration_data.request_response (platform_id=22)\n'
         'rds_warehouse_public.facts',
         'KYB Contact Information, KYB Website'],
        ['Worth Website Scanning', '30',
         'Internal website analysis\n(domain age, SSL, malware)',
         'Business submission (if website enabled)',
         'facts.v1 (website_score facts)',
         'integration_data.request_response (platform_id=30)',
         'KYB Website tab'],
        ['AI Enrichment (GPT-5-mini)', '31',
         'OpenAI API call\n(conditional: fires if <3 vendors have NAICS)',
         'After other integrations (safety net)',
         'facts.v1 naics_code (if AI wins)',
         'integration_data.request_response (platform_id=31)\n'
         'rds_warehouse_public.facts (if AI code wins)',
         'KYB Background Industry section\n(low confidence AI result)'],
        ['Rutter (accounting)', '5-13',
         'OAuth accounting connection\n(QuickBooks, Xero, etc.)',
         'Applicant connects accounting in onboarding',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE (financial category)',
         'accounting_statements, financial_statements',
         'Accounting tab, Worth Score (financial category)'],
        ['Tax Status', '15',
         'IRS e-services API\n(tax transcript retrieval)',
         'Business submission (if Tax Status enabled)',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE (financial category)',
         'integration_data.request_response (platform_id=15)',
         'Taxes tab, Worth Score (financial category)'],
        ['GIACT gVerify', '26',
         'PRODUCTION: Live GIACT API\nSANDOX: GIACT sandbox\nMOCK: GiactMockStrategy (fake data)',
         'Business submission (if gVerify enabled)',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE (banking category)',
         'integration_data.request_response (platform_id=26)',
         'Banking tab (account verification status)'],
        ['Adverse Media', '27',
         'Adverse media search API',
         'Business submission (if enabled)',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE (public_records category)',
         'integration_data.request_response (platform_id=27)',
         'Public Records tab (adverse media section)'],
        ['NPI', '28',
         'NPPES public API (CMS healthcare registry)',
         'Business submission (if NPI enabled)',
         'facts.v1 npi_data fact',
         'integration_data.request_response (platform_id=28)\n'
         'rds_warehouse_public.facts',
         'KYB Background NPI section'],
        ['Match (Mastercard)', '37',
         'Mastercard Match termination inquiry API\n(merchant terminated for fraud check)',
         'Business submission (if Match enabled for customer)',
         'MATCH_PRO_BULK Kafka event',
         'integration_data.request_response (platform_id=37)',
         'Public Records / Watchlists'],
        ['KYX', '40',
         'PRODUCTION: KYX_BASE_URL API\nSANDBOX: KYX_SANDBOX_BASE_URL',
         'Business submission (if KYX enabled)',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE',
         'integration_data.request_response (platform_id=40)',
         'Depends on KYX integration purpose'],
        ['Canada Open', '32',
         'Canadian business registry\n(for Canadian businesses)',
         'Business submission (if Canadian address detected)',
         'facts.v1',
         'integration_data.request_response (platform_id=32)',
         'KYB Business Registration (Canada)'],
    ],
    col_widths=[1.5, 1.0, 1.7, 1.4, 1.6, 1.9, 1.5],
    fs=8,
)

H2('XGBoost Model — What It Does and Does NOT Do in the UI')
callout(
    'IMPORTANT CLARIFICATION: The XGBoost entity-matching model (entity_matching_20250127 v1) '
    'does NOT produce the Worth Score, the Match badges, the Risk Level, or any directly '
    'visible UI element on its own. Its outputs are INTERMEDIATE — they feed into other systems:\n\n'
    '1. zi_probability / efx_probability / oc_probability -> smb_zi_oc_efx_combined -> '
    'customer_files (Pipeline B internal analytics only)\n\n'
    '2. These XGBoost match probabilities are ALSO used as the confidence score in '
    'integration_data.request_response (column: confidence) for Pipeline A source selection\n\n'
    '3. The Fact Engine uses these confidence scores to pick the winning vendor for '
    'NAICS code, industry, firmographic fields — which THEN appears in the KYB tab\n\n'
    'So: XGBoost -> confidence -> Fact Engine winner -> rds_warehouse_public.facts -> KYB tab'
)

table(
    ['XGBoost output', 'Where it goes', 'What it powers in the UI'],
    [
        ['zi_probability (0-1)\nP(ZoomInfo record = same business)',
         'datascience.ml_model_matches (Redshift)\n'
         '-> smb_zi_oc_efx_combined.zi_match_confidence\n'
         '-> datascience.customer_files.zi_match_confidence',
         'INTERNAL (Pipeline B) — not directly visible in UI.\n'
         'Indirectly: determines ZI vs EFX winner in customer_files analytics table.'],
        ['efx_probability (0-1)\nP(Equifax record = same business)',
         'Same path as above\n'
         '-> customer_files.efx_match_confidence',
         'INTERNAL (Pipeline B) — not directly visible.\n'
         'Also used as source.confidence in Pipeline A for EFX winner selection.'],
        ['oc_probability (0-1)\nP(OC record = same business)',
         'datascience.ml_model_matches (Redshift)\n'
         '-> integration_data.request_response.confidence (Pipeline A)',
         'Determines OC source confidence in Fact Engine.\n'
         'If OC wins: industry_code_uids parsed -> naics_code fact -> KYB Industry section.'],
        ['Source confidence scores\n(via similarity_index/55 fallback)',
         'integration_data.request_response.confidence\n'
         '(per source per business)',
         'GET /facts/business/{id}/details returns source.confidence.\n'
         'Shown to admin in source breakdown panel (if exposed).'],
    ],
    col_widths=[2.4, 3.2, 4.0],
)

H2('Final Summary: Where Each UI Screen Gets Its Data')
table(
    ['UI Screen / Field', 'Pipeline', 'Storage table', 'Who wrote it'],
    [
        ['Customers list — Business Name, Status, Date',
         'A (real-time)', 'data_businesses, data_cases (case-service PostgreSQL)',
         'case-service API on form submission'],
        ['Customers list — Account Type badge',
         'Config (not a pipeline)',
         'data_customers.customer_type (case-service PostgreSQL)',
         'Set at tenant provisioning by Worth AI admin'],
        ['Case Overview — Worth Score 819/850',
         'Scoring service (manual-score-service)',
         'business_scores.weighted_score_850 (case-service PostgreSQL)',
         'manual-score-service Kafka worker after all integrations complete'],
        ['Case Overview — Risk Level "Low Risk"',
         'Scoring service',
         'business_scores.risk_level (derived via score_decision_matrix)',
         'Same as above — score_decision_matrix lookup'],
        ['KYB Background — NAICS Code 812113',
         'A (real-time Fact Engine)',
         'rds_warehouse_public.facts name="naics_code" (JSONB)',
         'Pipeline A Fact Engine winner (Middesk > OC > ZI > Trulioo > EFX > AI)'],
        ['KYB Background — MCC Code 7230',
         'A (real-time Fact Engine)',
         'rds_warehouse_public.facts name="mcc_code"',
         'rel_naics_mcc crosswalk OR direct vendor MCC'],
        ['KYB Background — Number of Employees: 1',
         'A (real-time Fact Engine)',
         'rds_warehouse_public.facts name="employee_count"',
         'ZoomInfo (zi_c_employee_count) or Equifax — via Pipeline A winner'],
        ['KYB Business Registration — SOS data',
         'A (real-time Fact Engine)',
         'integration_data.business_entity_verification (Middesk)',
         'integration-service Middesk live API call (platform_id=16)'],
        ['KYB Watchlists',
         'A (real-time)',
         'integration_data.business_entity_verification.watchlist_hits',
         'integration-service Trulioo KYB call (platform_id=38)'],
        ['KYC — Owner fields (yellow highlight)',
         'n/a — applicant input',
         'data_owners, data_applicants (case-service PostgreSQL)',
         'Submitted by business owner during onboarding form'],
        ['KYC — Match badges',
         'A (real-time IDV/Trulioo)',
         'data_identity_verifications (Plaid IDV)\n'
         'OR integration_data.business_entity_verification (Trulioo person)',
         'integration-service Plaid IDV or Trulioo PSC screening'],
        ['Case Activity feed',
         'n/a — audit',
         'case_audit_logs, case_comments (case-service PostgreSQL)',
         'Written by case-service on every data change and analyst action'],
        ['Worth Score breakdown by category',
         'Scoring service',
         'business_score_factors (case-service PostgreSQL)',
         'manual-score-service — one row per factor'],
        ['Public Records — SOS registration',
         'A (real-time)',
         'integration_data.business_entity_verification (Middesk)',
         'integration-service Middesk (platform_id=16)'],
        ['Banking tab — transactions, balances',
         'A (real-time)',
         'plaid_accounts, plaid_transactions, bank_statements',
         'integration-service Plaid (platform_id=1)'],
        ['customer_files.primary_naics_code (internal)',
         'B (batch Redshift — NOT in UI)',
         'datascience.customer_files (Redshift)',
         'warehouse-service sp_recreate_customer_files() batch job — ZI vs EFX only'],
    ],
    col_widths=[2.8, 1.7, 2.5, 2.6],
    fs=8.5,
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 12 — NEW UI ELEMENTS FROM SCREENSHOTS (DEEP DIVE)
# ════════════════════════════════════════════════════════════════════════════

H1('Section 12 — New UI Elements: KYC No Match, Fraud Report, Addresses, Watchlists', c=PURPLE, pb=False)

H2('12.1 — KYC "No Match" Badge (Marcus Lopez — Home Address)')
callout(
    'The red "No Match" badge on Home Address means the address submitted by the applicant '
    'was compared against official third-party data (Trulioo Comprehensive View) and '
    'the addresses DID NOT match. This is a meaningful KYC risk signal: the person '
    'may not actually live at the submitted address.'
)
table(
    ['Badge', 'Meaning', 'Technical source', 'What triggers it'],
    [
        ['✅ Match (green)',
         'Submitted field value confirmed by Trulioo or Plaid IDV as matching official records',
         'Trulioo DatasourceFields[].Status = "match"\nOR Plaid IDV field comparison',
         'Integration-service: extractAddressMatchStatusFromDatasourceFields()\n'
         'Uses "Comprehensive View" datasource exclusively as primary signal.\n'
         'Returns "match" if all address fields are match or missing.'],
        ['🔴 No Match (red)',
         'Submitted value explicitly contradicted by official records',
         'Trulioo DatasourceFields[].Status = "nomatch" on any address field',
         'If ANY address field (street, city, state, ZIP) returns "nomatch"\n'
         'in the Comprehensive View → overall status = "nomatch"\n'
         '→ address_verification review task status = "failure"\n'
         '→ UI shows red "No Match" badge'],
        ['(no badge)',
         'Field not verified — IDV disabled or verification service returned no data',
         'No IDV record for this owner\nOR IDV_STATUS = PENDING/CANCELED/EXPIRED',
         'Integration settings: identity_verification.status = "INACTIVE"\n'
         'OR Plaid IDV not completed\nOR Trulioo returned no DatasourceFields for this field'],
    ],
    col_widths=[1.4, 2.5, 2.8, 3.0],
)

H3('Where "No Match" is stored and queried')
code_block([
    '-- address_verification review task (Trulioo result):',
    'SELECT key, status, message, sublabel',
    'FROM integration_data.business_entity_review_task',
    "WHERE business_entity_verification_id = '<uuid>'",
    "  AND key = 'address_verification';",
    '-- status = "success" -> Match  |  status = "failure" -> No Match',
    '',
    '-- business_entity_address_source (submitted vs reported addresses):',
    'SELECT address_line_1, city, state, postal_code, submitted, deliverable',
    'FROM integration_data.business_entity_address_source',
    "WHERE business_entity_verification_id = '<uuid>';",
    '-- submitted = true  -> Submitted Address (what the business owner entered)',
    '-- submitted = false -> Reported Address (what Trulioo/Middesk found in registry)',
])

H2('12.2 — Fraud Report (Marcus Lopez: Synthetic Identity Risk Score, Stolen Identity Risk Score)')
callout(
    'The Fraud Report section under KYC is powered exclusively by Plaid IDV (platform_id = 18). '
    'It appears ONLY when Plaid IDV is enabled for the customer. '
    'Plaid runs a "risk check" as part of the identity verification flow, '
    'analyzing behavioral signals and identity abuse patterns.'
)
table(
    ['Fraud Report field', 'Value shown', 'Plaid API field', 'What it means'],
    [
        ['Name', 'MARCUS LOPEZ',
         'risk_check.name',
         'The name extracted and confirmed by Plaid from the submitted ID document'],
        ['User Interactions', 'N/A',
         'risk_check.user_interactions',
         'Behavioral signal: how the user interacted with the IDV flow '
         '(e.g., unusual copy-paste patterns, atypical mouse movements). '
         'N/A = not enough signal or not detected.'],
        ['Fraud Ring', 'N/A',
         'risk_check.fraud_ring_detected',
         'Whether Plaid detected signals that this device/identity has been used '
         'in coordinated fraud ring activity. N/A = not detected.'],
        ['Bot Presence', 'N/A',
         'risk_check.bot_detected',
         'Whether automated/bot behavior was detected during IDV session. '
         'N/A = no bot signal detected.'],
        ['Synthetic Identity Risk Score', '3 (Low Risk)',
         'risk_check.identity_abuse_signals.synthetic_identity.score',
         'Score 0-100: likelihood this identity was synthetically fabricated '
         '(combining real and fake information to create a "Frankenstein identity"). '
         'Score 3 = very low synthetic identity risk. Scores > 70 = High Risk.'],
        ['Stolen Identity Risk Score', '1 (Low Risk)',
         'risk_check.identity_abuse_signals.stolen_identity.score',
         'Score 0-100: likelihood that this identity was stolen from a real person. '
         'Score 1 = very low stolen identity risk. Scores > 70 = High Risk.'],
    ],
    col_widths=[2.4, 1.3, 2.8, 3.1],
)

H3('Email Report (the other section under KYC)')
table(
    ['Email Report field', 'Plaid API field', 'What it means'],
    [
        ['Is Deliverable', 'risk_check.email.is_deliverable',
         'Whether the submitted email address can actually receive email '
         '(i.e., it\'s a real, active mailbox — not a fake or disposable address)'],
        ['Breach Count', 'risk_check.email.breach_count',
         'Number of times this email appeared in known data breach databases '
         '(e.g., Have I Been Pwned). High count = elevated fraud risk.'],
        ['First / Last Breached At', 'risk_check.email.first_breached_at / last_breached_at',
         'Dates when this email was first and most recently seen in a breach'],
        ['Domain Is Free Provider', 'risk_check.email.domain_is_free_provider',
         'Whether the email is from a free provider (Gmail, Yahoo, etc.) vs a business domain. '
         'Free provider emails = slightly higher fraud risk for business applications.'],
        ['Domain Is Disposable', 'risk_check.email.domain_is_disposable',
         'Whether the domain is a known disposable/temporary email service '
         '(e.g., mailinator.com). Disposable email = strong fraud signal.'],
        ['TLD Is Suspicious', 'risk_check.email.top_level_domain_is_suspicious',
         'Whether the top-level domain (e.g. .xyz, .tk) is known to be associated with spam/fraud'],
        ['IP Spam List Count', 'risk_check.ip_spam_list_count',
         'Number of spam/fraud blacklists the applicant\'s IP address appeared on'],
    ],
    col_widths=[2.2, 2.8, 4.6],
)

H3('Fraud Report data flow')
lineage([
    'Plaid IDV session: applicant completes photo ID + selfie',
    '  -> Plaid runs risk_check on submitted identity',
    '  -> Plaid API returns risk_check.identity_abuse_signals.* in IDV result',
    '       |',
    'integration-service/lib/plaid/plaidIdv.ts:',
    '  getApplicantVerificationResponse(owner.id)',
    '  -> extracts risk_check from Plaid IDV response',
    '  -> Maps to Worth AI schema:',
    '     synthetic_identity_risk_score = risk_check.identity_abuse_signals.synthetic_identity.score',
    '     stolen_identity_risk_score    = risk_check.identity_abuse_signals.stolen_identity.score',
    '     fraud_ring_detected           = risk_check.fraud_ring_detected',
    '     bot_detected                  = risk_check.bot_detected',
    '       |',
    'Stored in: integration_data.data_identity_verifications',
    '  (applicant_risk_check_result JSONB column)',
    '       |',
    'Fact Engine sources.ownerVerification reads from data_identity_verifications',
    '  -> Creates owner_verification fact:',
    '     { owner_id: { email_report: {...}, fraud_report: {...} } }',
    '  -> Stored in rds_warehouse_public.facts name="owner_verification"',
    '       |',
    'UI reads owner_verification fact via GET /facts/business/{id}/details',
    '  -> Displays Email Report and Fraud Report panels per owner',
])

H2('12.3 — KYB Contact Information: Submitted vs Reported Addresses')
callout(
    'The Contact Information tab shows multiple address types: '
    '"Submitted Address" (what the business owner typed in the form) and '
    '"Reported Address" (what Trulioo/Middesk found in official registries). '
    'The key insight: these may be DIFFERENT addresses — and the difference matters for risk.'
)

table(
    ['Address type', 'Badge', 'What it means', 'Source', 'Technical field'],
    [
        ['Submitted Address',
         'Business Registration: Verified / Unverified\nGoogle Profile: Verified / Unverified',
         'The address the business owner typed during onboarding.\n'
         '"Verified" = this address was confirmed by Middesk SOS as matching the filing.\n'
         '"Unverified" = Trulioo/Middesk could not confirm this is the SOS-registered address.',
         'data_businesses.address_* (submitted by applicant)\n'
         'Verified via: Middesk review_task key="name" + Trulioo Comprehensive View',
         'integration_data.business_entity_address_source WHERE submitted = true'],
        ['Reported Address (with Deliverable badge)',
         'Business Registration: Verified / Unverified\n(green "Deliverable" tag)',
         'An address Middesk or Trulioo found in official registry as the registered address.\n'
         '"Deliverable" = USPS confirms mail can be delivered to this address.\n'
         '"Verified" = matches what the applicant submitted.\n'
         '"Unverified" = an additional registry address that doesn\'t match submitted.',
         'Middesk addressSources[].submitted = false\n'
         '+ Middesk addressSources[].deliverable = true/false\n'
         '+ Trulioo StandardizedLocations',
         'integration_data.business_entity_address_source WHERE submitted = false'],
        ['Google Profile: Verified',
         'Blue "Verified" badge',
         'The business address was found and confirmed on Google Maps/Google Business.\n'
         'SERP Google Profile integration checked if the address matches Google\'s data.',
         'SERP_GOOGLE_PROFILE (platform_id = 39)\n'
         'caseTabValuesManager.ts: googleProfileStatus = "passed"',
         'integration_data.request_response (platform_id=39)\n'
         'value.google_address compared to submitted address'],
        ['Google Profile: Unverified',
         'Yellow "Unverified" badge',
         'Address was submitted but Google Profile could not confirm it.\n'
         'The business may not have a Google Business listing, or the address differs.',
         'Same source — googleProfileStatus = "missing" (address present but not confirmed)',
         'integration_data.request_response (platform_id=39)'],
    ],
    col_widths=[1.8, 2.0, 2.8, 2.0, 1.6],
)

H3('Why are there multiple Reported Addresses?')
body(
    'Middesk and Trulioo search official SOS databases which may have MULTIPLE '
    'addresses on file for the same entity (e.g., registered agent address, '
    'principal office, mailing address, historical addresses). '
    'All of them are shown. The one marked "Deliverable" is the USPS-confirmed '
    'current deliverable address. "Unverified" reported addresses are additional registry '
    'entries that don\'t match the submitted address — useful for analyst review.'
)

H3('address_source table lineage')
lineage([
    'SOURCES for Submitted Address:',
    '  data_businesses.address_line_1, address_city, address_state, address_postal_code',
    '  (entered by business owner during onboarding form)',
    '',
    'SOURCES for Reported Addresses:',
    '  Middesk API response -> businessEntityVerification.addressSources[]',
    '    { full_address, submitted: false, deliverable: true/false }',
    '  Trulioo API response -> StandardizedLocations[]',
    '    { Address1, City, Province, PostalCode }',
    '',
    'STORAGE:',
    '  integration_data.business_entity_address_source',
    '    business_entity_verification_id  FK',
    '    address_line_1, city, state, postal_code',
    '    submitted (boolean: true=applicant submitted, false=registry reported)',
    '    deliverable (boolean: USPS-confirmed deliverable)',
    '',
    'ADDRESS VERIFICATION STATUS:',
    '  integration_data.business_entity_review_task',
    '    key = "address_verification"',
    '    status = "success" (Verified) | "failure" (Unverified)',
    '  -> Determined by: extractAddressMatchStatusFromDatasourceFields()',
    '     Trulioo Comprehensive View DatasourceFields match/nomatch',
    '',
    'GOOGLE PROFILE VERIFICATION:',
    '  integration_data.request_response (platform_id = 39)',
    '  -> SERP_GOOGLE_PROFILE integration scrapes Google Maps for business address',
    '  -> caseTabValuesManager.ts computes googleProfileStatus:',
    '     "passed" = Google address matches submitted -> "Verified" badge',
    '     "missing" = address present but not confirmed -> "Unverified" badge',
])

H2('12.4 — KYB Business Registration: EIN/Tax ID Verification')
body(
    '"Business Registration ✓ Verified" and "Tax ID Number (EIN): 931667813" '
    'shown in the Business Registration tab for Pizza and a Chef LLC. '
    'This shows the EIN was verified, not just submitted.'
)
table(
    ['Field', 'Source', 'Verification process', 'Storage'],
    [
        ['Business Name', 'data_businesses.name (submitted)',
         'Middesk: review_task key="name"\n'
         'Trulioo: DatasourceFields where FieldName=BusinessName, Status=match',
         'integration_data.business_entity_review_task key="name"'],
        ['Tax ID Number (EIN)', 'data_businesses.tin (submitted, AES-256 encrypted)',
         'Middesk: review_task key="tin"\n'
         '  status="success" = EIN matches IRS records for this business\n'
         '  status="failure" = EIN does not match registry\n'
         'Trulioo: extractRegistrationNumberFromTruliooResponse()\n'
         '  Compares submitted TIN against registry registration number',
         'data_businesses.tin (encrypted)\n'
         'integration_data.business_entity_review_task key="tin"\n'
         '  { status, message, sublabel }'],
        ['Secretary of State Filings ✓ Verified',
         'Middesk SOS API call',
         'Middesk successfully found and confirmed an SOS filing matching\n'
         'the submitted business name and EIN in the state\'s official registry.\n'
         'integration_data.business_entity_verification.status = "success"',
         'integration_data.business_entity_verification\n'
         '(name, entity_type, incorporation_date, status, sos_status)'],
        ['Filing Status: Active', 'Middesk SOS response',
         'SOS registry reports business status as "Active" (in good standing)\n'
         'vs "Inactive", "Dissolved", "Revoked", "Suspended"',
         'integration_data.business_entity_verification.sos_status'],
        ['Entity Jurisdiction: Domestic / Primary',
         'Middesk SOS response',
         '"Domestic" = business is registered in its home state (FL).\n'
         '"Primary" = this is the primary state of registration\n'
         '(vs "Foreign" = registered in a different state than where incorporated).',
         'integration_data.business_entity_verification.entity_jurisdiction_type'],
        ['Registration Date', 'Middesk SOS response',
         'The date the business entity was first registered with the state.\n'
         'Pizza and a Chef LLC: 05/31/2023 (registered in FL).',
         'integration_data.business_entity_verification.incorporation_date'],
        ['Corporate Officers', 'Middesk SOS response',
         'Officially registered officers/managers as listed in the SOS filing.\n'
         '"William A. Bergez, MANAGER & REGISTERED AGENT"\n'
         'Cross-referenced with KYC owners to check for discrepancies.',
         'integration_data.business_entity_verification.officers JSONB'],
    ],
    col_widths=[2.0, 1.8, 3.3, 2.5],
)

H2('12.5 — KYB Business Names: Submitted Names vs Reported Names')
table(
    ['Name type', 'Example', 'Source', 'Purpose'],
    [
        ['Submitted Name (legal)', 'PIZZA AND A CHEF LLC',
         'data_businesses.name — submitted via onboarding form as legal name',
         'Primary submitted entity name'],
        ['Submitted Name (DBA)', 'Pizza and a chef',
         'data_businesses.dba_name — submitted via onboarding as "doing business as"',
         'Alternative trading name submitted by applicant'],
        ['Reported Name', 'ARI\'s Deli & Pizza',
         'Middesk SOS: businessEntityVerification.name\n'
         'Additional DBA names found in registry under same EIN',
         'What official registries have on record — may differ from submitted.\n'
         'Discrepancy between submitted and reported = name mismatch risk signal.'],
        ['Reported Name', 'PIZZA AND A CHEF LLC',
         'Trulioo: DatasourceFields FieldName=BusinessName from Comprehensive View',
         'Confirming name — Trulioo also found same legal name in its registry,\n'
         'which corroborates Middesk\'s SOS data.'],
    ],
    col_widths=[1.8, 2.2, 2.8, 2.8],
)
body(
    '"Business Names ✓ Verified" means the Fact Engine resolved a consistent legal name '
    'across Middesk and Trulioo sources. The name_verification review task status = "success". '
    'The presence of "ARI\'s Deli & Pizza" as a Reported Name is noteworthy — this is '
    'an additional DBA that Worth AI found in the registry, which the applicant did not '
    'disclose. This is visible to analysts for manual review.'
)

H2('12.6 — KYC "Verification Pending" Badge (Leslie Knope)')
callout(
    '"Verification Pending" (yellow badge) is different from "Unverified" and "Verified".\n\n'
    'It means: IDV was triggered for this owner, the applicant STARTED the Plaid IDV flow, '
    'but has NOT yet completed it.\n\n'
    'IDV_STATUS = PENDING (status_id = 2) — the applicant clicked the IDV link, '
    'started the identity session, but the photo ID + selfie upload is incomplete '
    'or awaiting Plaid\'s processing.'
)
table(
    ['KYC owner badge', 'Condition', 'IDV_STATUS value', 'Table field'],
    [
        ['✅ Verified',
         'Plaid IDV flow completed successfully — all required fields matched',
         'SUCCESS (1)',
         'data_identity_verifications.status_id = 1'],
        ['⏳ Verification Pending',
         'IDV flow started but not completed (applicant in progress)',
         'PENDING (2)',
         'data_identity_verifications.status_id = 2'],
        ['❌ Unverified (IDV Disabled)',
         'IDV feature is not active for this customer — never triggered',
         'n/a — no IDV record exists',
         'data_integration_settings.settings.identity_verification.status = "INACTIVE"'],
        ['❌ Unverified (Expired)',
         'IDV link expired before applicant completed the flow',
         'EXPIRED (4)',
         'data_identity_verifications.status_id = 4'],
        ['❌ Unverified (Canceled)',
         'Applicant canceled the IDV session',
         'CANCELED (3)',
         'data_identity_verifications.status_id = 3'],
        ['❌ Unverified (Failed)',
         'IDV completed but identity could not be verified (failed checks)',
         'FAILED (99)',
         'data_identity_verifications.status_id = 99'],
    ],
    col_widths=[2.2, 2.8, 1.5, 3.1],
)

H2('12.7 — KYB Watchlists Tab: Which Lists Are Scanned and How')
body(
    'The Watchlists tab shows screening results for the business AND all owners '
    'against every major sanctions and regulatory watchlist. '
    'Data comes from TWO sources: Trulioo KYB (business screening) and Trulioo PSC (owner screening).'
)

H3('Which watchlists are scanned (from the UI right panel)')
table(
    ['Agency', 'List name', 'List type in Trulioo', 'Risk meaning if hit'],
    [
        ['Office of Foreign Assets Control (OFAC)',
         'Capta List Foreign Sanctions Evaders', 'SANCTIONS',
         'Individual evaded US financial sanctions — extreme risk'],
        ['OFAC', 'Non-SDN Menu-Based Sanctions', 'SANCTIONS',
         'Entity subject to sectoral/menu-based OFAC restrictions'],
        ['OFAC', 'Non-SDN Iranian Sanctions', 'SANCTIONS',
         'Entity linked to Iran sanctions program'],
        ['OFAC', 'Non-SDN Chinese Military-Industrial Complex Companies List', 'SANCTIONS',
         'Entity listed as Chinese military company — US investment restrictions'],
        ['OFAC', 'Non-SDN Palestine Legislative Council List', 'SANCTIONS',
         'Entities linked to Palestinian Legislative Council'],
        ['OFAC', 'Specially Designated Nationals (SDN)', 'SANCTIONS',
         'Highest severity — entity on US blacklist; cannot transact with US persons'],
        ['OFAC', 'Sectoral Sanctions Identifications List (SSI)', 'SANCTIONS',
         'Entity in sanctioned sector (Russian finance, energy, defense)'],
        ['Bureau of Industry and Security (BIS)', 'Entity List', 'SANCTIONS',
         'Entity prohibited from receiving US exports without license'],
        ['BIS', 'Denied Persons List (DPL)', 'SANCTIONS',
         'Entity denied US export privileges'],
        ['BIS', 'Unverified List', 'SANCTIONS',
         'Entity whose bona fides BIS could not verify in prior transactions'],
        ['Various', 'PEP (Politically Exposed Persons)', 'PEP',
         'Individual is or is related to a government official — higher corruption risk'],
        ['Various', 'Adverse Media', 'ADVERSE_MEDIA',
         'Negative news coverage linking entity to crime, fraud, corruption'],
    ],
    col_widths=[2.5, 2.8, 1.5, 3.4],
)

H3('Watchlist data flow')
lineage([
    'BUSINESS SCREENING (KYB):',
    '  integration-service triggers Trulioo KYB API (platform_id=38)',
    '  Trulioo screens business name + address against ALL watchlists above',
    '  -> Returns watchlistResults: TruliooWatchlistHit[]',
    '     { listType, listName, sourceAgencyName, listCountry, url }',
    '  -> transformTruliooHitToWatchlistMetadata() maps hits to WatchlistValueMetadatum',
    '  -> storeBusinessWatchlistResults() writes to:',
    '     integration_data.business_entity_review_task',
    '       key = "watchlist"',
    '       metadata = JSON array of WatchlistValueMetadatum[]',
    '       status = "success" (no hits) | "failure" (hits found)',
    '',
    'OWNER SCREENING (PSC — Persons with Significant Control):',
    '  integration-service triggers Trulioo PSC API (platform_id=42)',
    '  Screens each beneficial owner (>25% ownership) against PEP + sanctions lists',
    '  -> extractWatchlistHitsFromScreenedPeople()',
    '  -> entity_type = WATCHLIST_ENTITY_TYPE.PERSON (vs BUSINESS for KYB)',
    '  -> Same storage: integration_data.business_entity_review_task',
    '',
    'CONSOLIDATED WATCHLIST FACT:',
    '  consolidatedWatchlist.ts: calculateConsolidatedWatchlist()',
    '  Merges KYB hits (BUSINESS type) + PSC hits (PERSON type)',
    '  Deduplication: createWatchlistDedupKey() by title + agency + entity_name',
    '  -> Stored in rds_warehouse_public.facts name="watchlist"',
    '',
    'UI displays:',
    '  "Hits for Lisa\'s Nail Salon" (BUSINESS entity hits)',
    '  "Hits for Lisa\'s Nail Salon, L.l.c." (alternate legal name screened)',
    '  "Hits for Thanh Ngoc Nguyen" (PERSON PSC hits)',
    '  Each: ✓ No Hits | or list of matched watchlist entries',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 13 — CUSTOMER DETAILS PAGE & STANDALONE CASES
# ════════════════════════════════════════════════════════════════════════════

H1('Section 13 — Customer Details Page, Standalone Cases, Tenants', c=PURPLE, pb=False)

H2('13.1 — Customer Details Page (Joannah SB)')
callout(
    'The Customer Details page is the admin\'s view of a specific CUSTOMER (the bank/lender '
    'using Worth AI). It has 4 tabs: Overview, Cases, Businesses, Users. '
    'This is distinct from a Case Details page (which shows a specific business application).'
)

table(
    ['UI element', 'Value', 'Source table', 'Notes'],
    [
        ['Customer name', 'Joannah SB',
         'data_customers.name (case-service PostgreSQL)',
         'Set when the Worth AI admin creates the customer account'],
        ['Customer ID (UUID)', 'Shown with eye+copy icons',
         'data_customers.id',
         'Primary UUID key for this customer. Used in all API calls.'],
        ['Created date', '03/27/2026, 1:52 PM',
         'data_customers.created_at',
         'Timestamp when customer was provisioned in Worth AI'],
        ['Account Type: SANDBOX',
         'Yellow "SANDBOX" badge',
         'data_customers.customer_type = "SANDBOX"',
         'Per Section 8 — set at provisioning, controls which vendor strategies are used'],
        ['Account Status: Active',
         'Green "Active ✓" badge',
         'data_customers.status (or data_customers.is_active)',
         'Active = customer is live and can submit applications.\n'
         'Other states: Inactive, Suspended.'],
        ['Account Owner', 'Joannah Sesebo',
         'data_users linked to data_customers via rel_customer_users',
         'The Worth AI internal user responsible for this customer account'],
        ['April Onboarding Count: 0 New Onboardings',
         'Counter + donut chart',
         'onboarding_schema.data_customer_onboarding_limits\n'
         '(current_count, limit, reset_at)',
         'Tracks how many businesses were onboarded THIS MONTH.\n'
         '"0 New Onboardings" + "no monthly limit set" = no cap configured.\n'
         'When limit is set: displays usage vs limit.\n'
         'Monthly reset: cron job monthly-onboarding-limit-reset\n'
         'resets current_count = 0 every month (run-cron.ts).'],
    ],
    col_widths=[2.2, 1.8, 2.8, 2.8],
)

H2('13.2 — Customer Details → Cases Tab')
body(
    'Shows all cases submitted under this customer. '
    'Each row is a data_cases record linked to this customer via data_cases.customer_id.'
)
table(
    ['Column', 'Source field', 'Notes'],
    [
        ['Case #', 'data_cases.id (UUID, truncated)',
         'Unique case identifier. Links to the Case Details page.'],
        ['Date', 'data_cases.created_at',
         'When the case was created (business submitted application)'],
        ['Type', 'data_cases.case_type_id -> CASE_TYPE enum',
         'ONBOARDING (1) = new business application\n'
         'APPLICATION_EDIT (2) = business updated their application\n'
         'RISK (3) = ongoing risk monitoring case'],
        ['Business Name', 'data_businesses.name WHERE id = data_cases.business_id',
         'The business that submitted this application'],
        ['Status', 'data_cases.status_id -> CASE_STATUS constants',
         'Under Manual Review (4), Auto Approved (6), Archived (9), etc.\n'
         'See Section 7 for full 20-status breakdown'],
        ['Integrations', '"Complete" or "In Progress"',
         'Derived from business_integration_tasks WHERE business_id = X\n'
         'Complete = all tasks finished (completed or failed)\n'
         'In Progress = at least one task still running'],
        ['Assignee', 'data_cases.assigned_to_user_id -> data_users',
         'Analyst assigned to review. "-" = unassigned.'],
    ],
    col_widths=[1.5, 2.5, 5.6],
)

H3('Case status differences seen in screenshot')
table(
    ['Business', 'Status', 'What happened', 'Why this status'],
    [
        ['LADYBIRD ACADEMY OF ST...', 'Under Manual Review',
         'Worth Score was in the medium-risk range\nOR analyst manually triggered review',
         'score_decision_matrix returned decision="UNDER_MANUAL_REVIEW"\n'
         'OR PATCH /cases/{id} { status: "UNDER_MANUAL_REVIEW" }'],
        ['Tech Systems Inc', 'Archived',
         'Case was closed without a final approval/rejection',
         'PATCH /cases/{id} { status: "ARCHIVED" } — manual analyst action\n'
         'OR automated archiving policy after inactivity period'],
        ['Lisa\'s Nail Salon', 'Auto Approved',
         'Worth Score 819/850 met the auto-approve threshold in score_decision_matrix',
         'manual-score-service: score_decision = "AUTO_APPROVED"\n'
         '-> UPDATE data_cases SET status_id = 6 (AUTO_APPROVED)'],
    ],
    col_widths=[2.2, 1.8, 2.8, 2.8],
)

H2('13.3 — Standalone Cases: What They Are and Why They Exist')
callout(
    'The left sidebar has TWO separate navigation items: "Customers" and "Standalone Cases".\n\n'
    'A Standalone Case is a case where data_cases.customer_id IS NULL — '
    'there is NO associated customer/tenant. The business submitted directly '
    'without being invited by a specific customer.\n\n'
    'This is a feature flag gated concept: FEATURE_FLAGS.BEST_60_STANDALONE_CASE'
)

table(
    ['Case type', 'customer_id', 'applicant_id', 'How created', 'Use case'],
    [
        ['Customer case (standard)',
         'SET — UUID of the customer (bank/lender)',
         'SET — UUID of the business owner',
         'Customer invites business via onboarding portal',
         'Normal onboarding flow: bank invites merchant to apply'],
        ['Standalone case',
         'NULL',
         'SET — UUID of the business owner who self-submitted',
         'Business submits directly without customer invitation\n'
         'INSERT INTO data_cases WHERE customer_id IS NULL',
         'Business owner wants Worth AI score without being linked to a specific lender.\n'
         'Used for self-service credit assessment or pre-qualification.'],
    ],
    col_widths=[1.8, 1.5, 1.5, 2.5, 2.3],
)

code_block([
    '-- Find standalone cases:',
    'SELECT * FROM data_cases WHERE customer_id IS NULL;',
    '',
    '-- Standalone cases for a specific business:',
    'SELECT * FROM data_cases',
    "WHERE business_id = '<uuid>' AND customer_id IS NULL AND applicant_id = '<applicant_uuid>';",
    '',
    '-- When a business later gets a customer invitation, both cases co-exist:',
    '-- standalone case (customer_id=NULL) + invited case (customer_id=<uuid>)',
    '-- The scoring pipeline links them via cases_to_link in the score trigger payload.',
])

H2('13.4 — Tenants: What They Are')
callout(
    '"Tenants" in the left sidebar is a SEPARATE concept from "Customers". '
    'A Tenant is a top-level organizational entity — typically the parent organization '
    'that owns and manages multiple Customer accounts.\n\n'
    'Hierarchy: Tenant -> Customer (child) -> Case -> Business\n\n'
    'Example: "Worldpay" might be the Tenant, with "Worldpay KYC 3" and "Worldpay POC KYC rerun" '
    'being two separate Customer accounts (different environments/configurations) owned by the same Tenant.'
)

table(
    ['Concept', 'Worth AI term', 'Table', 'Relationship'],
    [
        ['Top-level org owning Worth AI account', 'Tenant',
         'data_tenants (or equivalent top-level table)',
         'Parent of one or more Customers'],
        ['Bank/lender/fintech using Worth AI', 'Customer',
         'data_customers',
         'Child of Tenant. Has customer_type (SANDBOX/PRODUCTION).\n'
         'Can inherit settings from parent customer via parent_customer_data Kafka payload.'],
        ['Business applying for services', 'Business',
         'data_businesses',
         'Linked to Customer via data_cases.customer_id'],
        ['Integration settings inheritance',
         'Customer hierarchy',
         'integration_data.data_integration_settings',
         'When a child Customer is created with parent_customer_data.parent_id,\n'
         'integration settings are copied from parent to child customer.\n'
         'Source: integration-service/src/messaging/kafka/consumers/handlers/business.ts\n'
         'customerIntegrationSettings.copyCustomerIntegrationSettingsFromParent()'],
    ],
    col_widths=[2.2, 1.5, 2.3, 3.6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 14 — UNANSWERED QUESTIONS: COMPLETE Q&A
# ════════════════════════════════════════════════════════════════════════════

H1('Section 14 — All Questions Worth Asking: Complete Q&A', c=PURPLE, pb=False)
callout(
    'This section covers every important question about Worth AI that was NOT explicitly '
    'in the screenshots — questions that any analyst, engineer, or executive working with '
    'this system should understand. Each question is answered with full technical lineage.'
)

H2('Q1: How does Worth AI know which vendor\'s NAICS code to trust?')
body(
    'The Fact Engine uses a two-part decision system (from integration-service/lib/facts/rules.ts):'
)
bullet('Step 1: factWithHighestConfidence()', ' — If the top source\'s confidence exceeds the next by >0.05, it wins outright')
bullet('Step 2: weightedFactSelector()', ' — If scores are within 0.05, source weight (Middesk=2.0, OC=0.9, ZI=0.8, Trulioo=0.8, EFX=0.7, AI=0.1) breaks the tie')
body('Confidence is produced by:')
bullet('XGBoost entity-matching model', ' for OC, ZI, EFX (probability that the vendor record = same business)')
bullet('XGBoost confidenceScoreMany()', ' or task-based score for Middesk')
bullet('similarity_index / 55', ' (heuristic Levenshtein) for Trulioo and fallback')
bullet('Self-reported HIGH/MED/LOW', ' mapped to 0.70/0.50/0.30 for AI')

H2('Q2: What happens when two vendors give completely different NAICS codes?')
body(
    'This is a source conflict. The Fact Engine still picks a winner via the weight/confidence rules above. '
    'The losing codes appear in facts.alternatives[]. '
    'The UI shows source.confidence for each, allowing an analyst to judge. '
    'If no source reaches a minimum confidence (there is NO minimum cutoff in the code), '
    'the lowest-confidence result still wins. '
    'The analyst can override via PATCH /facts/business/{id}/override/naics_code.'
)

H2('Q3: How long does the entire onboarding pipeline take from submission to Worth Score?')
table(
    ['Stage', 'Typical duration', 'What happens'],
    [
        ['Case creation + Kafka dispatch', '< 1 second',
         'POST /businesses/customers/{id} -> data_cases created -> Kafka BUSINESS_SUBMITTED'],
        ['Middesk SOS lookup', '2–10 seconds',
         'Live API call to Secretary of State registry (varies by state)'],
        ['OpenCorporates + ZoomInfo + Equifax', '1–3 seconds',
         'Redshift queries (pre-loaded data — no external API call)'],
        ['Trulioo KYB + watchlist screening', '3–15 seconds',
         'Live API call — varies by country and number of watchlists'],
        ['NPI lookup (if enabled)', '1–5 seconds',
         'NPPES public API call'],
        ['SERP scraping (if configured)', '5–30 seconds',
         'Web crawling — highly variable depending on site'],
        ['Plaid IDV (if enabled)', 'Minutes to hours',
         'Requires applicant to complete photo ID + selfie flow — human-dependent'],
        ['Score computation (after all integrations)', '1–3 seconds',
         'Kafka GENERATE_AI_SCORE -> calculateBankingScore + calculatePublicRecordsScore + calculateFinancialScore'],
        ['Total (without Plaid IDV)', '~15–60 seconds typical',
         'Depends on slowest integration (usually Trulioo or Middesk for complex businesses)'],
    ],
    col_widths=[2.5, 1.8, 5.3],
)

H2('Q4: What is the complete list of database schemas and who owns them?')
table(
    ['Schema / Database', 'Owner service', 'Key tables'],
    [
        ['public (case-service PostgreSQL)', 'case-service',
         'data_cases, data_businesses, data_applicants, data_owners, data_customers,\n'
         'business_scores, business_score_factors, score_decision_matrix,\n'
         'score_config_history, data_current_scores, rel_case_applicants'],
        ['onboarding_schema (case-service PostgreSQL)', 'case-service',
         'data_customer_onboarding_limits, data_customer_stage_fields_config,\n'
         'rel_customer_setup_countries'],
        ['integration_data (integration-service PostgreSQL)', 'integration-service',
         'business_entity_verification, business_entity_review_task,\n'
         'business_entity_address_source, business_entity_people,\n'
         'business_integration_tasks, request_response, data_integration_settings,\n'
         'data_identity_verifications, data_connections'],
        ['rds_warehouse_public (warehouse-service PostgreSQL + Redshift mirror)', 'warehouse-service',
         'facts (all 217 JSONB facts per business)'],
        ['rds_cases_public (case-service PostgreSQL + Redshift mirror)', 'case-service',
         'data_businesses, data_cases, core_naics_code, core_mcc_code,\n'
         'core_business_industries, rel_naics_mcc'],
        ['datascience (Redshift)', 'warehouse-service + data engineering',
         'customer_files (Pipeline B), ml_model_matches (XGBoost output),\n'
         'zoominfo_standard_ml_2, open_corporates_standard_ml_2,\n'
         'smb_zi_oc_efx_combined, business_score_triggers'],
        ['warehouse (Redshift)', 'data engineering',
         'equifax_us_standardized, equifax_us_latest'],
        ['zoominfo (Redshift)', 'data engineering',
         'comp_standard_global (raw ZoomInfo dump)'],
        ['subscriptions (case-service PostgreSQL)', 'case-service',
         'data_customers (Stripe subscription context)'],
        ['s3 (AWS S3)', 'integration-service (lib/aws)',
         'Documents, bank statements, signed PDFs, uploaded files\n'
         'References in: s3_documents table (integration-service)'],
    ],
    col_widths=[2.8, 2.2, 4.6],
)

H2('Q5: What is the difference between the "facts" table and data_businesses?')
table(
    ['Aspect', 'rds_warehouse_public.facts', 'rds_cases_public.data_businesses'],
    [
        ['Structure', 'Key-value JSONB store\n(business_id, name, value JSONB)',
         'Wide denormalized relational table\n(one row per business, many columns)'],
        ['Content', '217 computed facts with full source lineage,\nconfidence scores, alternatives[], overrides',
         'Basic submitted fields: name, address, tin (encrypted),\ntax_id, industry (FK), website, mobile'],
        ['Data origin', 'Pipeline A Fact Engine — aggregates all vendors\nand selects winner per fact',
         'Submitted directly by applicant during onboarding form\n(NOT enriched by vendor data)'],
        ['NAICS storage', 'name="naics_code" -> value.code (vendor-enriched winner)',
         'industry column = FK to core_business_industries.id\n(sector level only, not 6-digit NAICS)'],
        ['Updated by', 'warehouse-service FactService (Kafka consumer)\n'
         'UPSERT on (business_id, name)',
         'case-service API on form submission and application edits'],
        ['Used for', 'Customer-facing API responses (all enriched data)\nWorth 360 Report',
         'Internal joins, search, basic business metadata'],
    ],
    col_widths=[1.5, 3.8, 4.3],
)

H2('Q6: How is a business matched to a vendor record? (The entity matching problem)')
body(
    'This is the core challenge: when a merchant submits "Lisa\'s Nail Salon, 7554 Louisiana 182, '
    'Morgan City, LA", how does Worth AI find the right ZoomInfo, Equifax, or OC record?'
)
lineage([
    'STEP 1 — Heuristic similarity (Levenshtein):',
    '  For each vendor (ZI, EFX, OC), compute similarity_index against all vendor records:',
    '    similarity_index = (20 - levenshtein(submitted_name, vendor_name))',
    '                     + (20 - levenshtein(submitted_address, vendor_address))',
    '                     + state_match + city_match + zip_match',
    '  Keep top 1,000 candidates per business per vendor.',
    '  Stored in: smb_zoominfo_standardized_joined, smb_equifax_standardized_joined, etc.',
    '',
    'STEP 2 — XGBoost entity matching model (entity_matching_20250127 v1):',
    '  For each candidate pair, compute 33 pairwise features:',
    '    - Jaccard k-gram similarities on name (k=1,2,3,4, word-level)',
    '    - Jaccard k-gram similarities on street, short name',
    '    - Exact match flags: city, ZIP, street number, address line 2',
    '    - Street number distance, block-level match',
    '  Output: zi_probability, efx_probability, oc_probability (0-1)',
    '  Stored in: datascience.ml_model_matches',
    '',
    'STEP 3 — Combined confidence score:',
    '  zi_match_confidence = zi_probability (if >= 0.8 from XGBoost)',
    '                      ELSE similarity_index/55 (if >= 0.8)',
    '                      ELSE 0',
    '  Stored in: datascience.smb_zi_oc_efx_combined',
    '',
    'STEP 4 — Fact Engine uses confidence as source weight:',
    '  integration_data.request_response.confidence = zi_match_confidence',
    '  factWithHighestConfidence() picks the vendor with highest confidence',
    '  -> That vendor\'s industry code (naics, sic, etc.) wins',
])

H2('Q7: What happens if a business has NO match in any vendor database?')
body('This is handled gracefully at each level:')
table(
    ['Level', 'What happens', 'Result in UI'],
    [
        ['ZoomInfo no match', 'similarity_index < 45 for all ZI records\nOR XGBoost probability < 0.8',
         'ZI source: confidence = 0, does not participate in Fact Engine winner selection'],
        ['Equifax no match', 'Same — confidence = 0',
         'EFX source not used for Fact Engine winner'],
        ['OC no match', 'oc_probability < 0.8 and no heuristic match',
         'OC source not used'],
        ['Middesk no SOS record', 'No SOS filing found in any state for this business name + EIN',
         'Business Registration: "Not Verified" / "No records found"\n'
         'NAICS from Middesk: not contributed'],
        ['Trulioo no match', 'Trulioo returns no Comprehensive View data',
         'KYB Watchlists: "No data available"\n'
         'Address: not verified'],
        ['ALL vendors fail', 'No vendor provided NAICS code',
         'AI enrichment (GPT-5-mini) triggers as safety net.\n'
         'If AI also fails -> naics_code = "561499" (All Other Business Support Services)\n'
         'Worth Score: Public Records category scores as minimum (no SOS = risk factor)'],
    ],
    col_widths=[1.8, 3.2, 4.6],
)

H2('Q8: Can a business have multiple cases? How does Worth AI handle that?')
body(
    'YES. One business (data_businesses record) can have MULTIPLE cases — one per customer '
    'who invited them, plus one standalone case. '
    'Example: The same business might apply to Worldpay AND Branch (two separate customers). '
    'Each creates a separate data_cases row with the same business_id but different customer_id.'
)
code_block([
    '-- All cases for a business:',
    'SELECT id AS case_id, customer_id, status_id, created_at, case_type_id',
    'FROM data_cases',
    "WHERE business_id = '<uuid>'",
    'ORDER BY created_at DESC;',
    '',
    '-- One business may have:',
    '-- case 1: customer_id = Worldpay_UUID, status_id = 6 (Auto Approved)',
    '-- case 2: customer_id = Branch_UUID,   status_id = 4 (Under Manual Review)',
    '-- case 3: customer_id = NULL,           status_id = 7 (Score Calculated) <- standalone',
    '',
    '-- The Worth Score is shared:',
    '-- data_current_scores stores ONE score per (business_id, customer_id)',
    '-- Each customer sees their own score (may differ based on their score_decision_matrix)',
])

H2('Q9: How does risk monitoring work? What triggers RISK_ALERT status?')
body(
    'Risk monitoring is a continuous post-onboarding service. '
    'After a business is onboarded (AUTO_APPROVED or MANUALLY_APPROVED), '
    'Worth AI can continue monitoring it for changes in risk signals.'
)
table(
    ['Risk monitoring trigger', 'What causes it', 'Result'],
    [
        ['Scheduled monitoring refresh',
         'Cron job triggers MONITORING_REFRESH score trigger for all monitored businesses',
         'New Worth Score computed with fresh vendor data.\n'
         'If new score drops below threshold -> RISK_ALERT triggered.'],
        ['Subscription refresh',
         'New vendor data received via Kafka SUBSCRIPTION_REFRESH',
         'Score recomputed with updated data (e.g., new ZoomInfo data)'],
        ['Watchlist hit detected',
         'Trulioo ongoing screening returns a new watchlist match\n'
         '(OFAC, BIS, PEP, etc.)',
         'Kafka risk alert event fired -> data_cases.status_id = RISK_ALERT (14)'],
        ['Adverse media hit',
         'Adverse Media integration (platform_id=27) finds new negative news',
         'data_risk_alerts table updated -> case status updated -> analyst notified'],
        ['Analyst-triggered',
         'Admin manually triggers a risk review via UI action',
         'PATCH /cases/{id} { status: "RISK_ALERT" }'],
    ],
    col_widths=[2.5, 3.3, 3.8],
)

H2('Q10: What does the Worth 360 Report contain? Where does each section\'s data come from?')
table(
    ['Report section', 'Data source', 'Table / fact'],
    [
        ['Executive Summary / Risk Level',
         'manual-score-service score result',
         'business_scores.risk_level + score_decision'],
        ['Worth Score breakdown',
         'manual-score-service per-factor results',
         'business_score_factors table'],
        ['Company Profile',
         'Pipeline A Fact Engine winners',
         'rds_warehouse_public.facts (naics_code, legal_name, address, employee_count, revenue...)'],
        ['Business Registration (SOS)',
         'Middesk SOS API',
         'integration_data.business_entity_verification'],
        ['Watchlist screening',
         'Trulioo KYB + PSC',
         'integration_data.business_entity_review_task key="watchlist"'],
        ['Financial Trends',
         'Plaid + Rutter',
         'plaid_transactions + accounting_statements'],
        ['Public Records (NPI, adverse media)',
         'NPI (NPPES) + Adverse Media integration',
         'integration_data.request_response (platform_id=28, 27)'],
        ['Owners (KYC)',
         'Case-service data + Plaid IDV + Trulioo PSC',
         'data_owners + data_identity_verifications + owner_verification fact'],
        ['The report itself (PDF)',
         'Generated by report.ts Kafka handler in case-service',
         'S3 document storage. Kafka event: GENERATE_REPORT'],
    ],
    col_widths=[2.2, 2.5, 4.9],
)

H2('Q11: How does Worth AI handle non-US businesses? (M Supply Guam case)')
callout(
    'The M Supply Guam case (jurisdiction: GU — Guam, US territory) is a good example. '
    'Guam uses US business registration (SOS via Guam DCCA) but is a US territory, '
    'not a state. Worth AI handles this via jurisdiction-aware fact extraction.'
)
table(
    ['Element', 'For US states', 'For territories / international'],
    [
        ['SOS lookup', 'Middesk covers all 50 US states + DC',
         'Middesk also covers US territories (GU, PR, VI, etc.).\n'
         'For non-US: OpenCorporates (global), Canada Open (platform_id=32), Trulioo (global KYB)'],
        ['Address format', 'Standard US format (state code, 5-digit ZIP)',
         'Jurisdiction-specific. Guam: ZIP 96910 is valid US ZIP.\n'
         'International: country code + local postal format'],
        ['NAICS vs SIC', 'NAICS preferred (US standard)',
         'UK businesses: UK SIC 2007 from OC (gb_sic- prefix in industry_code_uids)\n'
         'EU: NACE Rev2 (nace- prefix). Canada: ca_naics- prefix.\n'
         'These are captured in classification_codes fact but NOT yet exposed in UI.'],
        ['Watchlist screening', 'US OFAC lists + global lists',
         'Same global lists (OFAC, BIS, UN, EU) apply regardless of jurisdiction.\n'
         'Trulioo screens against all applicable lists for any country.'],
        ['Worth Score', 'Full score (banking + public records + financial)',
         'May have partial scores if some integrations don\'t cover the territory/country.\n'
         'E.g., no Plaid banking for some international businesses.'],
    ],
    col_widths=[1.8, 2.8, 4.6],
)

H2('Q12: What is the "control_person" designation and how is it set?')
body(
    '"Control Person" is a regulatory term (from FinCEN Customer Due Diligence rules). '
    'It refers to the individual with significant control over the business '
    '(e.g., CEO, Managing Member). Worth AI tracks this separately from beneficial owners.'
)
table(
    ['Designation', 'Who it applies to', 'How it is set', 'Table / field'],
    [
        ['Control Person',
         'The individual with significant managerial control (CEO, President, Managing Member).\n'
         'Required under FinCEN CDD rule for legal entity customers.',
         'Applicant self-designates during onboarding form.\n'
         'Customer config sets: max_control_persons (default 1).',
         'data_owners.is_control_person = true\n'
         'data_owners.title (e.g., "Partner", "CEO", "Managing Member")'],
        ['Beneficial Owner',
         'Any individual owning >= 25% of the business.\n'
         'FinCEN requires collecting all >= 25% owners.',
         'Applicant submits all owners >= 25% during onboarding.\n'
         'Ownership percentage stored in data_owners.ownership_percentage',
         'data_owners.ownership_percentage\n'
         'Validation: max_beneficial_owners config setting'],
        ['Applicant (submitter)',
         'The person who actually filled out and submitted the application.\n'
         'May or may not be a beneficial owner or control person.',
         'The logged-in user at the time of form submission.',
         'data_applicants.id linked to data_cases.applicant_id'],
    ],
    col_widths=[1.8, 2.8, 2.5, 2.5],
)

H2('Q13: How does the EIN/TIN get stored and is it visible to analysts?')
table(
    ['Aspect', 'Detail'],
    [
        ['Storage format', 'AES-256 encrypted at rest in data_businesses.tin column.\n'
         'The raw EIN string is never stored in plaintext in any table.'],
        ['Decryption', 'case-service decryptEin() utility function decrypts for authorized API calls.\n'
         'Only returned to roles with DATA_PERMISSION for TIN.'],
        ['Display in UI', 'Shown partially masked: last 4 digits visible (e.g., ***-**-3874 for SSN,\n'
         '***-667813 for EIN). Full value visible to admin analysts with permissions.'],
        ['Verification', 'Middesk review_task key="tin" confirms if submitted EIN matches IRS + SOS records.\n'
         'status="success" = EIN matches; status="failure" = mismatch.'],
        ['Compliance', 'EIN handling complies with IRS Publication 1075 (tax data security requirements).\n'
         'Separate from SSN which has additional PII protection under GLBA.'],
    ],
    col_widths=[2.0, 7.6],
)

H2('Q14: What is the complete Kafka topic map? Which service writes/reads which topic?')
table(
    ['Kafka topic', 'Who publishes', 'Who consumes', 'Key events'],
    [
        ['facts.v1', 'integration-service (Fact Engine)',
         'warehouse-service (FactService)',
         'CALCULATE_BUSINESS_FACTS, UPDATE_NAICS_CODE, UPDATE_MCC_CODE,\n'
         'UPDATE_INDUSTRY, FACT_OVERRIDE_CREATED_AUDIT, PROCESS_COMPLETION_FACTS'],
        ['ai_scores (AI_SCORES)', 'manual-score-service (taskHandler)',
         'AI scoring service (external)',
         'GENERATE_AI_SCORE — triggers Worth Score computation'],
        ['integration_data', 'integration-service (per task completion)',
         'manual-score-service (taskQueue)',
         'UPDATE_INTEGRATION_DATA_FOR_SCORE — feeds integration results into scoring'],
        ['business_events', 'case-service (on form submit)',
         'integration-service',
         'BUSINESS_SUBMITTED — triggers all integrations'],
        ['case_events', 'case-service (on status changes)',
         'Webhooks service, integration-service',
         'STATUS_CHANGED, MANUALLY_APPROVED, AUTO_APPROVED, etc.'],
        ['risk_alerts', 'integration-service (risk monitoring)',
         'case-service',
         'RISK_ALERT_TRIGGERED — creates RISK_ALERT status on case'],
        ['playground_ai_scores', 'manual-score-service (when BEST_56_PLAYGROUND_ENABLED)',
         'Playground scoring service',
         'Same as AI_SCORES but for experimental scoring environment'],
    ],
    col_widths=[2.2, 2.2, 2.2, 3.0],
)

H2('Q15: What exactly is Pipeline A vs Pipeline B and can they contradict each other?')
callout(
    'YES — they can and DO produce different NAICS codes for the same business. '
    'This is one of the most important architectural facts in Worth AI.'
)
table(
    ['Dimension', 'Pipeline A (Integration Service)', 'Pipeline B (Warehouse Service)'],
    [
        ['Purpose', 'Deliver best classification to customer in real time', 'Internal analytics, risk model training, data exports'],
        ['Sources', 'ALL 6+ vendors: Middesk, OC, ZI, EFX, Trulioo, AI', 'ZoomInfo + Equifax ONLY'],
        ['Winner rule', 'factWithHighestConfidence() -> weightedFactSelector()\n(confidence × weight across all 6 sources)',
         'CASE WHEN zi_match_confidence > efx_match_confidence\n  THEN ZoomInfo ELSE Equifax'],
        ['Customer sees?', 'YES — REST API, Worth 360 Report, Worth AI admin UI', 'NO — internal Redshift only'],
        ['NAICS stored in', 'rds_warehouse_public.facts name="naics_code"', 'datascience.customer_files.primary_naics_code'],
        ['Example divergence', 'Middesk SOS returns NAICS 812113 (Nail Salons) -> wins (weight 2.0)\nPipeline A shows 812113',
         'ZoomInfo has 812990 (Other Personal Care Services) with zi_confidence=0.85\nEFX has 812113 with efx_confidence=0.72\n-> ZI wins -> Pipeline B shows 812990'],
        ['Implication', 'Customer sees Middesk\'s more accurate SOS-verified code',
         'Redshift analytics table has ZoomInfo\'s less accurate code\nRisk model trained on potentially wrong NAICS'],
    ],
    col_widths=[1.8, 3.6, 4.2],
)

H2('Q16: How does the Worth Score change over time? Can it go up or down?')
body('YES — the Worth Score is recalculated on multiple triggers:')
table(
    ['Trigger', 'Constant name', 'When it fires'],
    [
        ['Initial onboarding', 'ONBOARDING_INVITE', 'First submission — standard flow'],
        ['Application edit', 'APPLICATION_EDIT', 'Business updates their application (new data submitted)'],
        ['Manual refresh', 'MANUAL_REFRESH', 'Analyst clicks "Refresh Score" in admin UI'],
        ['Monitoring refresh', 'MONITORING_REFRESH', 'Scheduled cron job for ongoing risk monitoring'],
        ['Subscription refresh', 'SUBSCRIPTION_REFRESH', 'New vendor data available (subscription update)'],
    ],
    col_widths=[2.2, 2.5, 4.9],
)
body(
    'Each recalculation creates a new business_score_triggers row and a new business_scores row. '
    'data_current_scores is updated to point to the latest score. '
    'Historical scores are preserved in business_score_history. '
    'A business that was AUTO_APPROVED at score 750 could drop to RISK_ALERT if a monitoring refresh '
    'returns a score of 450 (new watchlist hit, bank account overdrafts, etc.).'
)

H2('Q17: What is the Integrations Complete / In Progress status based on?')
code_block([
    '-- Check integration completion status:',
    'SELECT platform_id, task_code, status, updated_at',
    'FROM integration_data.business_integration_tasks',
    "WHERE business_id = '<uuid>'",
    'ORDER BY updated_at DESC;',
    '',
    '-- "Integrations Complete" = ALL tasks have status IN ("completed", "failed", "skipped")',
    '-- "In Progress"           = at least one task has status IN ("pending", "in_progress")',
    '',
    '-- Source code: integration-service/src/helpers/integrationsCompletionTracker.ts',
    '-- Also reflected as "Integrations Complete" banner in Case Details header.',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 15 — FULL DATA ENTITY RELATIONSHIP MAP
# ════════════════════════════════════════════════════════════════════════════

H1('Section 15 — Full Data Entity Relationship Map', c=PURPLE, pb=False)

H2('Core Business Object Relationships')
lineage([
    'HIERARCHY (top to bottom):',
    '',
    'data_tenants (top-level org owning Worth AI access)',
    '  |',
    '  +-- data_customers (bank/lender/fintech using Worth AI)',
    '        customer_id, name, customer_type (SANDBOX/PRODUCTION)',
    '        |',
    '        +-- rel_customer_users (which Worth AI users can see this customer)',
    '        |',
    '        +-- data_integration_settings (which vendors enabled, SANDBOX/PROD mode per vendor)',
    '        |',
    '        +-- score_decision_matrix (customer-specific score thresholds)',
    '        |',
    '        +-- onboarding_schema.data_customer_onboarding_limits (monthly onboarding cap)',
    '        |',
    '        +-- data_cases (one per business per customer)',
    '               case_id, business_id, customer_id, status_id, case_type_id',
    '               score_trigger_id -> business_score_triggers -> business_scores',
    '               |',
    '               +-- data_businesses (the merchant being verified)',
    '                    business_id, name, address, tin (encrypted), industry FK',
    '                    |',
    '                    +-- data_owners (beneficial owners >= 25% + control persons)',
    '                    |    owner_id, first_name, last_name, dob, ssn (encrypted)',
    '                    |    ownership_percentage, is_control_person',
    '                    |    |',
    '                    |    +-- data_identity_verifications (Plaid IDV per owner)',
    '                    |         status_id (SUCCESS=1, PENDING=2, FAILED=99)',
    '                    |',
    '                    +-- rds_warehouse_public.facts (217 vendor-enriched facts)',
    '                    |    (business_id, name, value JSONB, received_at)',
    '                    |    Written by Pipeline A Fact Engine',
    '                    |',
    '                    +-- integration_data.business_entity_verification (KYB per vendor)',
    '                    |    (business_id, platform_id, status, incorporation_date,',
    '                    |     entity_type, officers, watchlist_hits, address_sources)',
    '                    |',
    '                    +-- integration_data.request_response (raw API responses)',
    '                    |    (business_id, platform_id, response JSONB, confidence)',
    '                    |',
    '                    +-- business_scores (Worth Score per score trigger)',
    '                         (score_trigger_id, weighted_score_850, risk_level,',
    '                          score_decision, status)',
    '                         |',
    '                         +-- business_score_factors (per-factor breakdown)',
    '                         +-- score_inputs (raw data snapshot for audit)',
    '                         +-- data_current_scores (latest score per business+customer)',
])

H2('Pipeline Responsibility Matrix')
table(
    ['Question', 'Pipeline A', 'Pipeline B', 'Neither (applicant/analyst)'],
    [
        ['Who provides the NAICS code the customer sees?', 'YES', 'NO', 'NO'],
        ['Who drives the Worth Score?', 'Provides input data', 'NO', 'manual-score-service'],
        ['Who decides Auto Approve vs Manual Review?', 'Provides input data', 'NO', 'manual-score-service'],
        ['Who produces the KYB background fields?', 'YES (Fact Engine winner)', 'NO', 'NO'],
        ['Who does watchlist screening?', 'YES (Trulioo KYB)', 'NO', 'NO'],
        ['Who verifies KYC Match badges?', 'YES (Plaid IDV + Trulioo PSC)', 'NO', 'NO'],
        ['Who powers risk model training?', 'NO', 'YES (customer_files)', 'NO'],
        ['Who sets case status to Archived?', 'NO', 'NO', 'Analyst (manual)'],
        ['Who handles banking (Plaid)?', 'YES', 'NO', 'Applicant connects bank'],
        ['Who writes the Worth 360 Report?', 'Provides data', 'NO', 'case-service generates PDF'],
        ['Who tracks monthly onboarding limits?', 'NO', 'NO', 'case-service cron job'],
        ['Who provides ZI vs EFX winner for analytics?', 'NO', 'YES', 'NO'],
    ],
    col_widths=[4.2, 1.5, 1.5, 2.4],
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = ('/workspace/AI-Powered-NAICS-Industry-Classification-Agent/'
       'modeling/Worth_AI_Admin_UI_Lineage.docx')
doc.save(out)
size_kb = round(os.path.getsize(out) / 1024, 1)
print(f'Saved  : {out}')
print(f'Size   : {size_kb} KB')
print('Ready  : Import into Google Docs via File > Import')
