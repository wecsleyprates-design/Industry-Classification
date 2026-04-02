"""
Build a comprehensive, Google-Docs-compatible .docx Worthopedia.

Design principles:
- NO CSS/HTML — only python-docx native styles
- Tables with explicit column widths and light fills
- Code blocks as shaded 1-cell tables (Courier New, light grey bg)
- Callout / warning / gap boxes as shaded 1-cell tables with a left border
- All colour via RGBColor on Run objects (survives Google Docs import)
- Each of the 8 classification facts gets: priority table, post-processing
  code, JSONB example, SQL queries, and full table-lineage section
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
s = doc.sections[0]
s.page_width    = Inches(11.0)   # landscape for wide tables
s.page_height   = Inches(8.5)
s.left_margin   = Inches(0.75)
s.right_margin  = Inches(0.75)
s.top_margin    = Inches(0.65)
s.bottom_margin = Inches(0.65)
s.orientation   = 1  # landscape

PAGE_W = 9.5  # printable width (11 - 2*0.75)

# ── Colour palette ────────────────────────────────────────────────────────────
PURPLE  = RGBColor(0x5B, 0x21, 0xB6)
BLUE    = RGBColor(0x1E, 0x40, 0xAF)
TEAL    = RGBColor(0x04, 0x5F, 0x7E)
GREEN   = RGBColor(0x06, 0x5F, 0x46)
RED     = RGBColor(0x99, 0x1B, 0x1B)
AMBER   = RGBColor(0x78, 0x35, 0x00)
SLATE   = RGBColor(0x33, 0x41, 0x55)
DARK    = RGBColor(0x0F, 0x17, 0x2A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DGREY   = RGBColor(0x9C, 0xA3, 0xAF)


# ════════════════════════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _shade(cell, hex6: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)


def _left_border(cell, hex6: str, sz: int = 28):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    b = OxmlElement('w:left')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), hex6)
    tcB.append(b)
    tcPr.append(tcB)


def _cell_margins(cell, top=60, right=100, bottom=60, left=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('right', right), ('bottom', bottom), ('left', left)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _no_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblB.append(b)
    tblPr.append(tblB)


def _set_repeat_header(row):
    """Mark a table row as a repeated header row."""
    trPr = row._tr.get_or_add_trPr()
    tblH = OxmlElement('w:tblHeader')
    trPr.append(tblH)


# ════════════════════════════════════════════════════════════════════════════
# HIGH-LEVEL CONTENT HELPERS
# ════════════════════════════════════════════════════════════════════════════

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


def H1(text, colour=PURPLE, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(20)
    r.font.color.rgb = colour


def H2(text, colour=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(14)
    r.font.color.rgb = colour


def H3(text, colour=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(11.5)
    r.font.color.rgb = colour


def H4(text, colour=SLATE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(10.5)
    r.font.color.rgb = colour


def body(text, size=10.5, colour=SLATE, space_after=5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.font.color.rgb = colour
    r.bold = bold
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def body_parts(parts, space_after=5):
    """
    parts = list of (text, bold:bool, colour:RGBColor|None)
    Renders multiple styled runs in one paragraph.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, colour in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(10.5)
        r.font.color.rgb = colour or SLATE
    return p


def bullet(prefix: str, text: str, bold_prefix=True, size=10.5, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if prefix:
        rb = p.add_run(prefix)
        rb.bold = bold_prefix
        rb.font.size = Pt(size)
        rb.font.color.rgb = DARK
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = SLATE


def code_block(lines, size=8.5):
    """Shaded 1-cell table containing monospace code."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, 'F1F5F9')
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=80, right=120, bottom=80, left=120)
    first = True
    for line in lines:
        if first:
            cp = cell.paragraphs[0]
            first = False
        else:
            cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        r = cp.add_run(line)
        r.font.name  = 'Courier New'
        r.font.size  = Pt(size)
        r.font.color.rgb = DARK
    spacer(5)


def callout(text, bg='EFF6FF', border='1D4ED8', text_colour=None, size=10):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, bg)
    _left_border(cell, border, sz=28)
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=60, right=120, bottom=60, left=160)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(0)
    r = cp.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = text_colour or RGBColor.from_string(border)
    spacer(7)


def warn(text):
    callout(text, bg='FFFBEB', border='D97706',
            text_colour=RGBColor(0x78, 0x35, 0x00))


def gap_box(text):
    callout(text, bg='FEE2E2', border='DC2626',
            text_colour=RGBColor(0x7F, 0x1D, 0x1D))


def ok_box(text):
    callout(text, bg='F0FDF4', border='059669',
            text_colour=RGBColor(0x06, 0x5F, 0x46))


def data_table(headers, rows, col_widths=None, hdr_bg='DBEAFE', alt_bg='F8FAFF',
               font_size=9.5):
    """
    Multi-column table. Header row = light blue. Alternating light-grey rows.
    Values starting with YES/NO/WARN/GAP get automatic colour.
    """
    ncols = len(headers)
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    hrow = tbl.rows[0]
    _set_repeat_header(hrow)
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]
        _shade(cell, hdr_bg)
        _cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = DARK

    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                _shade(cell, alt_bg)
            _cell_margins(cell)
            p = cell.paragraphs[0]
            v = str(val)
            if v.startswith('YES') or v.startswith('✅'):
                rgb, bld = GREEN, True
            elif v.startswith('NO') or v.startswith('❌'):
                rgb, bld = RED, True
            elif v.startswith('WARN') or v.startswith('⚠️'):
                rgb, bld = AMBER, True
            elif v.startswith('GAP') or v.startswith('🔴'):
                rgb, bld = RED, True
            else:
                rgb, bld = SLATE, False
            r = p.add_run(v)
            r.font.size  = Pt(font_size)
            r.font.color.rgb = rgb
            r.bold = bld

    if col_widths:
        for ri2 in range(len(tbl.rows)):
            for ci2, w in enumerate(col_widths):
                tbl.rows[ri2].cells[ci2].width = Inches(w)
    spacer(8)


def lineage_box(lines):
    """Monospace lineage / ASCII-art diagram in a teal-bordered box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _shade(cell, 'F0F9FF')
    _left_border(cell, '0284C7', sz=28)
    cell.width = Inches(PAGE_W)
    _cell_margins(cell, top=80, right=120, bottom=80, left=140)
    first = True
    for line in lines:
        if first:
            cp = cell.paragraphs[0]
            first = False
        else:
            cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        r = cp.add_run(line)
        r.font.name  = 'Courier New'
        r.font.size  = Pt(8.5)
        r.font.color.rgb = DARK
    spacer(8)


# ════════════════════════════════════════════════════════════════════════════
# SHARED PRIORITY TABLE BUILDER FOR FACTS 1–8
# ════════════════════════════════════════════════════════════════════════════

PRIORITY_HEADERS = ['Priority', 'Source', 'Redshift / API field', 'Weight',
                    'Confidence method', 'Pipeline']

NAICS_PRIORITY_ROWS = [
    ['1st',   'Middesk',
     'SOS filing NAICS (case-service API call)',
     '2.0 — heaviest',
     'XGBoost confidenceScoreMany() OR 0.15 + 0.20 per passing KYB task\n'
     '(name-match, TIN, address, SOS lookup)',
     'A only'],
    ['2nd',   'OpenCorporates',
     'industry_code_uids → parse prefix us_naics-XXXXXX\n'
     'Table: open_corporates_standard_ml_2',
     '0.9',
     'XGBoost oc_probability\n(ml_model_matches, entity_matching_20250127 v1)',
     'A only'],
    ['3rd',   'ZoomInfo',
     'zi_c_naics6 (6-digit NAICS)\nTable: zoominfo_standard_ml_2',
     '0.8',
     'XGBoost zi_probability OR\nheuristic similarity_index/55 fallback',
     'A + B'],
    ['4th',   'Trulioo',
     'sicCode — 4-digit triggers POLLUTED flag\n(Live API — not in Redshift batch)',
     '0.8',
     'Heuristic ONLY:\nmatch.index / MAX_CONFIDENCE_INDEX (55)\nNo XGBoost for Trulioo.',
     'A only'],
    ['5th',   'Equifax',
     'efx_primnaicscode (6-digit NAICS)\nefx_primsic (4-digit SIC)\nTable: equifax_us_standardized',
     '0.7',
     'XGBoost efx_probability OR\nheuristic similarity_index/55 fallback',
     'A + B'],
    ['6th',   'GPT-5-mini AI',
     'AI-generated code (Live LLM call)\nplatform_id = 31 in request_response',
     '0.1 — lowest',
     'Self-reported in JSON output:\nHIGH → ~0.70 / MED → ~0.50 / LOW → ~0.30',
     'A only'],
    ['Last',  'Hardcoded fallback',
     '"561499" (All Other Business Support Services)',
     'n/a',
     'Validation failure fallback:\ncode not found in core_naics_code table',
     'A only'],
]

MCC_PRIORITY_ROWS = [
    ['Path A — best',
     'Middesk / Trulioo (live API)',
     'MCC returned directly in API response JSON\nplatform_id 16 (Middesk) or 38 (Trulioo)',
     'Inherits source weight',
     'Same as source (XGBoost / task-based)',
     'A only'],
    ['Path B — most common',
     'NAICS → MCC crosswalk',
     'rel_naics_mcc table: naics_code → mcc_code\ncase-service PostgreSQL',
     'n/a (deterministic lookup)',
     'n/a — exact table join',
     'A + B'],
    ['Path C — fallback',
     'GPT-5-mini AI',
     'mcc_code field in AI enrichment JSON response\nplatform_id = 31',
     '0.1',
     'Self-reported HIGH/MED/LOW',
     'A only'],
    ['Last',
     'NULL / missing',
     'No MCC stored if no path succeeded',
     'n/a',
     'n/a',
     'Both'],
]


# ════════════════════════════════════════════════════════════════════════════
# HELPER: RENDER ONE COMPLETE FACT SECTION
# ════════════════════════════════════════════════════════════════════════════

def fact_section(
    number,          # int  1..8
    name,            # str  'naics_code'
    pipeline_tag,    # str  '[PIPELINE A — Fact Engine winner]'
    what_it_is,      # str  long description
    source_note,     # str  extra note below priority table (or None)
    priority_rows,   # list of rows for priority table
    priority_cols,   # col widths list
    post_lines,      # list of str — post-processing code block
    jsonb_lines,     # list of str — jsonb example block
    query_sections,  # list of (heading_str, [code lines])
    lineage_lines,   # list of str — lineage / ASCII box
    extra_notes=None # list of (callout_fn, text) extra notes at end
):
    H2(f'Fact {number}: {name}')
    body(pipeline_tag, size=10, colour=PURPLE, bold=True, space_after=3)
    body(what_it_is, space_after=6)

    H4('Winner / Source Priority')
    data_table(PRIORITY_HEADERS, priority_rows, col_widths=priority_cols,
               font_size=9)

    if source_note:
        body(source_note, size=9.5, colour=SLATE, space_after=6)

    H4('Post-processing / Validation')
    code_block(post_lines, size=8.5)

    H4('JSONB Payload stored in rds_warehouse_public.facts')
    code_block(jsonb_lines, size=8.5)

    H4('SQL Queries')
    for (q_heading, q_lines) in query_sections:
        H4(q_heading)
        code_block(q_lines, size=8.5)

    H4('Table Lineage')
    lineage_box(lineage_lines)

    if extra_notes:
        for fn, text in extra_notes:
            fn(text)

    spacer(10)


# ════════════════════════════════════════════════════════════════════════════
# ░░░░░░░░░░░░░░░░░░  DOCUMENT CONTENT START  ░░░░░░░░░░░░░░░░░░░░░░░░░░░░░
# ════════════════════════════════════════════════════════════════════════════

# ── COVER ─────────────────────────────────────────────────────────────────────
spacer(20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Industry Classification Facts')
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = PURPLE
p.paragraph_format.space_after = Pt(5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Complete End-to-End Reference — Worth AI Classification Engine')
r.font.size = Pt(14); r.font.color.rgb = BLUE
p.paragraph_format.space_after = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Every source, table, SQL query, JSONB payload, Kafka event, and API endpoint '
    'explained from raw vendor data to customer-facing fact.\n'
    'Content verified against SIC-UK-Codes and Entity-Matching source code.'
)
r.font.size = Pt(10.5); r.font.color.rgb = SLATE
p.paragraph_format.space_after = Pt(20)

H2('Document Contents')
data_table(
    ['Tab', 'Section', 'What it covers'],
    [
        ['1', 'Overview — Two Pipelines',
         'Goals, sources, confidence methods, customer visibility, when they differ'],
        ['2', 'Full Timeline (Day 0 → Fact Delivered)',
         'Step-by-step: vendor loading, similarity tables, XGBoost, winner rule, '
         'Kafka, API response'],
        ['3', 'Pipeline B — Batch Redshift Deep Dive',
         'Entity matching XGBoost, zi/efx_match_confidence formula, customer_table.sql, '
         'all scenarios, gaps'],
        ['4', 'Pipeline A — Real-Time Fact Engine Deep Dive',
         '6 sources, 6 winner rules, AI enrichment triggers, all scenarios'],
        ['5', 'The 8 Classification Facts (full detail)',
         'Each fact: priority table, post-processing code, JSONB payload, SQL queries, '
         'table lineage diagram'],
        ['6', 'API & Storage Reference',
         'All endpoints, Kafka events, Redshift schema, query templates'],
    ],
    col_widths=[0.4, 2.2, 6.9],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TAB 1 — OVERVIEW
# ════════════════════════════════════════════════════════════════════════════

H1('Tab 1 — Overview: The Two Pipelines', colour=PURPLE)
callout(
    'Worth AI runs TWO completely separate pipelines for industry classification. '
    'They share some raw vendor input tables but have different goals, logic, '
    'outputs, and audiences.\n\n'
    'CRITICAL: The customer ONLY ever sees Pipeline A output. '
    'Pipeline B is internal Redshift analytics — never exposed via API.',
    bg='EDE9FE', border='5B21B6', text_colour=RGBColor(0x2E, 0x10, 0x65)
)

H2('Pipeline A — Integration Service (Real-Time)')
data_table(
    ['Attribute', 'Detail'],
    [
        ['Full name',     'integration-service (Node.js / TypeScript)'],
        ['Goal',
         'Deliver the best possible industry classification and KYB facts '
         'to the customer in real time, using ALL available sources.'],
        ['Trigger',       'Every business submission — POST /businesses/customers/{customerID}'],
        ['Sources',       'ZoomInfo, Equifax, OpenCorporates, Middesk, Trulioo, AI/GPT-5-mini (6+)'],
        ['Confidence',
         'XGBoost entity-matching model (OC, ZI, EFX)\n'
         '+ XGBoost task-based (Middesk via confidenceScoreMany)\n'
         '+ Heuristic similarity_index/55 (Trulioo fallback)\n'
         '+ GPT self-reported HIGH/MED/LOW (AI)'],
        ['Winner rule',   'factWithHighestConfidence() → weightedFactSelector() → manualOverride()'],
        ['Output table',  'rds_warehouse_public.facts — JSONB key-value, 217 facts, 1 row per fact per business'],
        ['Also writes to','integration_data.request_response (all raw API responses + confidence)'],
        ['Customer sees?','YES — REST API, Worth 360 Report, Worth AI UI'],
    ],
    col_widths=[1.8, 7.7],
)

H2('Pipeline B — Warehouse Service (Batch Redshift)')
data_table(
    ['Attribute', 'Detail'],
    [
        ['Full name',    'warehouse-service stored procedures (SQL)'],
        ['Goal',
         'Build a wide denormalized analytics table in Redshift for '
         'risk model training, bulk data exports, and internal analytics.'],
        ['Trigger',      'Scheduled batch job — NOT triggered per business submission'],
        ['Sources',      'ZoomInfo + Equifax ONLY (historical design limitation — see Tab 3)'],
        ['Confidence',
         'XGBoost entity-matching model (zi_probability, efx_probability)\n'
         '+ Heuristic similarity_index/55 fallback\n'
         'Combined into: zi_match_confidence and efx_match_confidence'],
        ['Winner rule',
         'CASE WHEN zi_match_confidence > efx_match_confidence THEN ZoomInfo ELSE Equifax'],
        ['Output table', 'datascience.customer_files — wide denormalized, one row per business'],
        ['Customer sees?','NO — internal Redshift only, zero API exposure'],
        ['Internal use', 'Risk model training, analytics dashboards, data exports'],
    ],
    col_widths=[1.8, 7.7],
)

H2('Shared Input Tables (both pipelines read these)')
data_table(
    ['Source', 'Raw table', 'Standardised table', 'Key industry column', 'Loaded by'],
    [
        ['ZoomInfo',
         'zoominfo.comp_standard_global',
         'datascience.zoominfo_standard_ml_2',
         'zi_c_naics6 (6-digit NAICS 2022)',
         'Bulk file ingestion, regular cadence'],
        ['Equifax',
         'warehouse.equifax_us_latest',
         'warehouse.equifax_us_standardized',
         'efx_primnaicscode (6-digit NAICS)\nefx_primsic (4-digit SIC)',
         'WARN: bulk ingestion — cadence unknown (source code note)'],
        ['OpenCorporates',
         'datascience.open_corporates_data',
         'datascience.open_corporates_standard_ml_2',
         'industry_code_uids\n(pipe-delimited: us_naics-541110|gb_sic-62012|nace-J6201)',
         'Bulk file ingestion'],
        ['Liberty',
         'dev.liberty.einmst_20260218\neinmst_15_5mn\neinmst_5_3m_remaining',
         '(not standardised into unified pipeline)',
         'NAICS, SIC columns',
         'Bulk file ingestion'],
    ],
    col_widths=[1.3, 2.1, 2.5, 2.1, 1.5],
)
gap_box(
    'GAP: Middesk and Trulioo have NO pre-loaded Redshift tables — they are live API '
    'calls only. Liberty has Redshift tables but was never joined into Pipeline B SQL. '
    'OC has a standardised Redshift table AND oc_probability from XGBoost, but '
    'customer_table.sql never parses industry_code_uids for NAICS selection.'
)

H2('Confidence Computation per Source — Summary')
data_table(
    ['Source', 'Platform ID', 'XGBoost?', 'Confidence field', 'Fallback', 'Weight'],
    [
        ['OpenCorporates', '23', 'YES — oc_probability',
         'ml_model_matches.oc_probability',
         'similarity_index/55 if probability < 0.8 threshold', '0.9'],
        ['Middesk', '16', 'YES — confidenceScoreMany()',
         'Sum of task scores: 0.15 base + 0.20 per passing task',
         'task-based score (no heuristic fallback)', '2.0'],
        ['ZoomInfo', '24', 'YES — zi_probability',
         'ml_model_matches.zi_probability',
         'similarity_index/55 if probability < 0.8 threshold', '0.8'],
        ['Equifax', '17', 'YES — efx_probability',
         'ml_model_matches.efx_probability',
         'similarity_index/55 if probability < 0.8 threshold', '0.7'],
        ['Trulioo', '38', 'NO — heuristic ONLY',
         'match.index / MAX_CONFIDENCE_INDEX (55)',
         'No fallback — heuristic is the only method', '0.8'],
        ['AI (GPT-5-mini)', '31', 'NO — self-reported',
         'confidence field in JSON response:\nHIGH→0.70 / MED→0.50 / LOW→0.30',
         'Defaults to 0.10 if missing', '0.1'],
    ],
    col_widths=[1.5, 1.1, 1.4, 2.5, 2.5, 0.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TAB 2 — FULL TIMELINE
# ════════════════════════════════════════════════════════════════════════════

H1('Tab 2 — Full Timeline: Day 0 to Fact Delivered', colour=PURPLE, page_break=False)
callout(
    'This section is the complete story — from raw vendor data loaded offline '
    'in Redshift, through XGBoost entity matching, through the real-time Fact Engine, '
    'to a customer-visible JSON response. Read this first to understand where '
    'every table in Tabs 3–6 fits.'
)

H2('Step 1 — Vendor Bulk Data Loaded into Redshift (Offline, before any submission)')
body(
    'Data engineering team loads vendor files into Redshift on a scheduled cadence. '
    'These tables are the foundation of both Pipeline A and Pipeline B.'
)
lineage_box([
    'ZoomInfo file delivery',
    '  -> zoominfo.comp_standard_global  (raw)',
    '  -> datascience.zoominfo_standard_ml_2  (standardised — name/address normalised)',
    '     Key column: zi_c_naics6 (6-digit NAICS 2022)',
    '',
    'Equifax file delivery',
    '  -> warehouse.equifax_us_latest  (raw)',
    '  -> warehouse.equifax_us_standardized  (standardised)',
    '     Key columns: efx_primnaicscode (6-digit NAICS), efx_primsic (4-digit SIC)',
    '',
    'OpenCorporates file delivery',
    '  -> datascience.open_corporates_data  (raw)',
    '  -> datascience.open_corporates_standard_ml_2  (standardised)',
    '     Key column: industry_code_uids  (e.g. us_naics-541110|gb_sic-62012|nace-J6201)',
    '',
    'Liberty file delivery',
    '  -> dev.liberty.einmst_20260218, einmst_15_5mn, einmst_5_3m_remaining  (raw only)',
    '     NOT standardised — never joined into any classification rule',
])

H2('Step 2 — Heuristic Similarity Scoring (Pipeline B only)')
body(
    'Before XGBoost runs, Worth AI finds the top candidate vendor records for each '
    'submitted business using Levenshtein string distance scoring. '
    'This produces ranked candidate lists — the top candidates are then scored by XGBoost.'
)
H3('similarity_index formula (from smb_zoominfo_standardized_joined.sql)')
code_block([
    'similarity_index =',
    '    (20 - levenshtein(lower(submitted_business_name),',
    '                       lower(vendor_name)))          -- max 20 pts',
    '  + (20 - levenshtein(lower(submitted_street),',
    '                       lower(vendor_street)))         -- max 20 pts',
    '  + CASE WHEN submitted_state   = vendor_state   THEN 1 ELSE 0 END',
    '  + CASE WHEN submitted_city    = vendor_city    THEN 1 ELSE 0 END',
    '  + CASE WHEN submitted_zipcode = vendor_zipcode THEN 1 ELSE 0 END',
    '                                                      -- max ~43-45 pts',
    '',
    'MAX_CONFIDENCE_INDEX = 55   (defined in integration-service/lib/facts/sources.ts)',
    'normalised_confidence = similarity_index / 55',
    '',
    'Minimum threshold: similarity_index >= 45  (= 0.818 confidence)',
    'Top 1,000 candidates per business kept, ranked by similarity_index.',
])
data_table(
    ['Table created', 'Vendor', 'Source Redshift table', 'Rows kept'],
    [
        ['datascience.smb_zoominfo_standardized_joined', 'ZoomInfo',
         'datascience.zoominfo_standard_ml_2', 'Top 1,000 candidates per business'],
        ['datascience.smb_equifax_standardized_joined', 'Equifax',
         'warehouse.equifax_us_standardized', 'Top 1,000 candidates per business'],
        ['datascience.smb_open_corporate_standardized_joined', 'OC',
         'open_corporates_standard_ml_2', 'Top 1,000 candidates per business'],
    ],
    col_widths=[3.1, 1.2, 2.8, 2.4],
)

H2('Step 3 — XGBoost Entity Matching Model (Pipeline B only)')
callout(
    '"Is this submitted business the same real-world legal entity as this vendor record?"  '
    'Model: entity_matching_20250127 v1 — trained on Worth AI labelled match pairs.'
)
H3('33 pairwise features per candidate pair')
data_table(
    ['Feature group', 'Features included'],
    [
        ['Name similarity',
         'Jaccard k-gram on business name: k=1, k=2, k=3, k=4, word-level (5 features)\n'
         'Jaccard k-gram on street name: k=1, k=2, k=3, k=4 (4 features)\n'
         'Jaccard k-gram on short name (city-stripped): k=1..4 (4 features)'],
        ['Address exact match',
         'city exact match (binary)\nZIP code exact match (binary)\n'
         'street number exact match (binary)\naddress line 2 exact match (binary)'],
        ['Street number distance',
         '|submitted_street_num - vendor_street_num| (absolute distance)\n'
         'Block-level match: submitted_num//100 == vendor_num//100 (binary)'],
        ['Remaining',
         'Additional features to reach 33 total (DBA name, state, other address fields)'],
    ],
    col_widths=[2.0, 7.5],
)
H3('XGBoost output (datascience.ml_model_matches)')
code_block([
    'SELECT business_id,',
    '       zi_probability,    -- P(submitted biz == ZoomInfo record)  0.0 to 1.0',
    '       efx_probability,   -- P(submitted biz == Equifax record)   0.0 to 1.0',
    '       oc_probability     -- P(submitted biz == OC record)        0.0 to 1.0',
    'FROM datascience.ml_model_matches',
    "WHERE business_id = '<uuid>';",
])

H2('Step 4 — Building zi_match_confidence and efx_match_confidence')
body(
    'smb_zi_oc_efx_combined.sql merges XGBoost probabilities with heuristic fallback '
    'to produce the two confidence scores that customer_table.sql uses in its winner rule.'
)
code_block([
    '-- From smb_zi_oc_efx_combined.sql (Pipeline B, verified from source code):',
    '',
    'zi_match_confidence =',
    '  CASE',
    "    WHEN m.zi_probability  IS NOT NULL THEN m.zi_probability          -- Tier 1: XGBoost",
    "    WHEN s.similarity_index/55.0 >= 0.8 THEN s.similarity_index/55.0 -- Tier 2: heuristic",
    '    ELSE 0',
    '  END,',
    '',
    'efx_match_confidence =',
    '  CASE',
    "    WHEN m.efx_probability IS NOT NULL THEN m.efx_probability         -- Tier 1: XGBoost",
    "    WHEN s.similarity_index/55.0 >= 0.8 THEN s.similarity_index/55.0 -- Tier 2: heuristic",
    '    ELSE 0',
    '  END',
    '',
    'FROM datascience.smb_zoominfo_standardized_joined s',
    'LEFT JOIN datascience.ml_model_matches m USING (business_id)',
])

H2('Step 5 — Winner-Takes-All SQL Rule (Pipeline B only)')
code_block([
    '-- From datascience/customer_table.sql (actual production SQL):',
    '',
    'primary_naics_code =',
    '  CASE',
    '    WHEN COALESCE(zi_match_confidence, 0) > COALESCE(efx_match_confidence, 0)',
    '      THEN z.zi_c_naics6                -- ZoomInfo wins:  ALL ZI firmographic fields',
    '    ELSE e.efx_primnaicscode             -- Equifax wins:   ALL EFX firmographic fields',
    '  END',
    '',
    '-- The SAME comparison controls every field in customer_files:',
    '-- company_name, address, city, state, ZIP, country, employee_count,',
    '-- revenue, website_url, affiliate_parent_flag, phone, SIC code, etc.',
    '',
    '-- Output: datascience.customer_files (rebuilt from scratch each batch run)',
    '-- Called by: sp_recreate_customer_files()',
])

H2('Step 6 — Business Submitted (Pipeline A starts, real-time)')
body(
    'A customer calls POST /businesses/customers/{customerID}. '
    'integration-service creates a business_id UUID and fires 6 parallel source lookups.'
)
lineage_box([
    'POST /businesses/customers/{customerID}',
    '  |',
    '  +--> Middesk API call     (platform_id=16) -- live SOS registry lookup',
    '  +--> OpenCorporates query (platform_id=23) -- Redshift pre-loaded table',
    '  +--> ZoomInfo query       (platform_id=24) -- Redshift pre-loaded table',
    '  +--> Equifax query        (platform_id=17) -- Redshift pre-loaded table',
    '  +--> Trulioo API call     (platform_id=38) -- live KYB API',
    '  +--> SERP / AI enrichment (platform_id=31) -- conditional: fires if < 3 vendor codes',
    '  |',
    '  All responses stored in: integration_data.request_response',
    '  (business_id, platform_id, response JSONB, confidence)',
])

H2('Step 7 — Fact Engine Selects Winner')
body('Rules applied in order (from integration-service/lib/facts/rules.ts):')
data_table(
    ['Rule', 'Name', 'Logic'],
    [
        ['1', 'manualOverride()',
         'If analyst set an override via PATCH /facts/.../override/{factName} → '
         'override ALWAYS wins immediately. Skips all other rules.'],
        ['2', 'factWithHighestConfidence()',
         'Sort all source candidates by confidence (0–1). '
         'If |top_confidence - second_confidence| > 0.05 (WEIGHT_THRESHOLD) → top wins outright.'],
        ['3', 'weightedFactSelector()',
         'Tie-break when gap <= 0.05: highest source weight wins.\n'
         'Middesk (2.0) > OC (0.9) > ZI (0.8) = Trulioo (0.8) > EFX (0.7) > AI (0.1)'],
        ['4', 'AI safety net',
         'Fires when fewer than 3 vendor sources returned a NAICS code.\n'
         'AI result enters competition with weight 0.1 — rarely wins outright.'],
        ['5', '561499 fallback',
         'If winning code is NOT found in core_naics_code table → replace with "561499".'],
    ],
    col_widths=[0.4, 2.0, 7.1],
)

H2('Step 8 — Kafka → Storage → Customer API')
lineage_box([
    'Fact Engine selects winner',
    '  |',
    '  +--> Kafka: publish to facts.v1 topic',
    '  |      Events:',
    '  |        CALCULATE_BUSINESS_FACTS    (initial calculation)',
    '  |        UPDATE_NAICS_CODE           (naics_code fact changed)',
    '  |        PROCESS_COMPLETION_FACTS    (all verifications complete)',
    '  |        FACT_OVERRIDE_CREATED_AUDIT (analyst override set)',
    '  |',
    '  +--> warehouse-service FactService.consume() subscribes to facts.v1',
    '  |      UPSERT into:',
    '  |        rds_warehouse_public.facts',
    '  |          (business_id, name="naics_code", value=JSONB, received_at)',
    '  |          INSERT ON CONFLICT (business_id, name) DO UPDATE SET value = EXCLUDED.value',
    '  |',
    '  +--> case-service Kafka handler (UPDATE_NAICS_CODE event)',
    '  |      UPDATE rds_cases_public.data_businesses',
    '  |        SET naics_id = (SELECT id FROM core_naics_code WHERE code = "722511")',
    '  |',
    '  +--> Customer API response',
    '         GET /facts/business/{id}/details',
    '           -> returns full JSONB from rds_warehouse_public.facts',
    '              {code, description, source.confidence, source.platformId, alternatives[]}',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TAB 3 — PIPELINE B DEEP DIVE
# ════════════════════════════════════════════════════════════════════════════

H1('Tab 3 — Pipeline B: Batch Redshift (Complete Deep Dive)', colour=PURPLE, page_break=False)

H2('Why Only ZoomInfo and Equifax?')
callout(
    'Pipeline B was built when only ZoomInfo and Equifax had clean pre-loaded '
    'Redshift tables with numeric NAICS fields. OC, Liberty, Middesk, and Trulioo '
    'were integrated into Pipeline A later and were never backfilled into Pipeline B SQL.'
)
data_table(
    ['Source', 'Has XGBoost score in Redshift?', 'Has NAICS in Redshift?',
     'In Pipeline B?', 'What is missing / why'],
    [
        ['OpenCorporates',
         'YES — oc_probability in ml_model_matches',
         'YES — industry_code_uids in open_corporates_standard_ml_2',
         'NO',
         'industry_code_uids is a pipe-delimited string.\n'
         'customer_table.sql never parses it to extract us_naics-XXXXXX entries.'],
        ['Liberty',
         'NO — never run through XGBoost entity matching',
         'YES — NAICS/SIC columns in dev.liberty.einmst_*',
         'NO',
         'Never joined into smb_zi_oc_efx_combined or customer_table.sql'],
        ['Middesk',
         'YES — confidenceScoreMany() in integration-service',
         'NO — live SOS API call only, no Redshift bulk table',
         'NO',
         'Live API source — no pre-loaded Redshift table for batch processing'],
        ['Trulioo',
         'NO — heuristic only (match.index/55)',
         'NO — live KYB API call only',
         'NO',
         'Live API source — not available for batch processing'],
    ],
    col_widths=[1.4, 2.2, 2.1, 0.9, 3.0],
)
gap_box(
    'IMPROVEMENT OPPORTUNITY: OC\'s oc_match_confidence already exists in '
    'smb_zi_oc_efx_combined. The only change needed is parsing industry_code_uids '
    'in customer_table.sql to extract us_naics-XXXXXX entries. '
    'OC weight = 0.9 (highest vendor weight) — adding it would most benefit '
    'non-US companies where OC has superior registry coverage.'
)

H2('All Pipeline B Scenarios')
data_table(
    ['Situation', 'What is stored in customer_files', 'Customer impact'],
    [
        ['zi_match_confidence > efx_match_confidence',
         'ALL fields from ZoomInfo row:\n'
         'zi_c_naics6, zi_c_name, zi_address, zi_employee_count,\n'
         'zi_revenue, zi_website_url, etc.',
         'ZoomInfo data used for analytics, risk scoring, and exports'],
        ['efx_match_confidence >= zi_match_confidence',
         'ALL fields from Equifax row:\n'
         'efx_primnaicscode, efx_name, efx_address,\n'
         'efx_employee_count, efx_revenue, etc.',
         'Equifax data used for analytics, risk scoring, and exports'],
        ['Both confidences = 0 (no match found)',
         'COALESCE fallback to existing stored value, else NULL.\n'
         'primary_naics_code = NULL.',
         'No industry code for this business in analytics table'],
        ['OC had a stronger match than ZI/EFX',
         'WARN: OC ignored — oc_match_confidence exists but SQL never reads it for NAICS',
         'Weaker ZI or EFX NAICS stored instead of better OC NAICS'],
        ['Liberty/Middesk/Trulioo had best match',
         'WARN: Ignored — live API sources not in batch pipeline',
         'Same ZI vs EFX winner — no benefit from superior live sources'],
        ['customer_files rebuilt on schedule',
         'sp_recreate_customer_files() TRUNCATES and rebuilds from scratch',
         'Latest batch result always overwrites previous — no history kept'],
    ],
    col_widths=[2.4, 3.6, 3.5],
)

H2('Pipeline B — Full Table Lineage')
lineage_box([
    'SOURCES (Redshift pre-loaded)',
    '  datascience.zoominfo_standard_ml_2         (zi_c_naics6)',
    '  warehouse.equifax_us_standardized           (efx_primnaicscode)',
    '  datascience.open_corporates_standard_ml_2   (industry_code_uids) [used in Step 2 only]',
    '       |',
    '       v',
    'STEP 1: Heuristic similarity scoring',
    '  datascience.smb_zoominfo_standardized_joined    (similarity_index, ranked top 1000)',
    '  datascience.smb_equifax_standardized_joined     (similarity_index, ranked top 1000)',
    '  datascience.smb_open_corporate_standardized_joined  (similarity_index, ranked top 1000)',
    '       |',
    '       v',
    'STEP 2: XGBoost entity matching model',
    '  entity_matching_20250127 v1  runs on candidate pairs',
    '  -> datascience.ml_model_matches',
    '       (zi_probability, efx_probability, oc_probability)',
    '       |',
    '       v',
    'STEP 3: Build match tables (stored procedures)',
    '  datascience.efx_matches_custom_inc_ml',
    '  datascience.oc_matches_custom_inc_ml',
    '  datascience.zoominfo_matches_custom_inc_ml',
    '       |',
    '       v',
    'STEP 4: Combine confidence scores',
    '  datascience.smb_zi_oc_efx_combined',
    '       (zi_match_confidence, efx_match_confidence)',
    '  Logic: XGBoost probability > heuristic similarity_index/55 fallback',
    '       |',
    '       v',
    'STEP 5: Winner rule SQL (customer_table.sql)',
    '  sp_recreate_customer_files() calls customer_table.sql',
    '  WHEN zi_match_confidence > efx_match_confidence -> ZoomInfo wins',
    '  ELSE -> Equifax wins',
    '       |',
    '       v',
    'OUTPUT: datascience.customer_files',
    '  (primary_naics_code, zi_match_confidence, efx_match_confidence,',
    '   match_confidence, all firmographic fields from the winning vendor)',
    '  Used by: risk model training, analytics dashboards, bulk data exports',
    '  NEVER exposed to customer API',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TAB 4 — PIPELINE A DEEP DIVE
# ════════════════════════════════════════════════════════════════════════════

H1('Tab 4 — Pipeline A: Real-Time Fact Engine (Complete Deep Dive)',
   colour=PURPLE, page_break=False)

H2('The 6 Sources — Field-Level Detail')
data_table(
    ['Source', 'Platform ID', 'Industry code field', 'How fetched', 'Weight'],
    [
        ['OpenCorporates', '23',
         'industry_code_uids\nPipe-delimited multi-taxonomy string:\n'
         'us_naics-541110|gb_sic-62012|nace-J6201|ca_naics-541110',
         'Redshift query on open_corporates_standard_ml_2\n(pre-loaded, not live API)',
         '0.9'],
        ['Middesk', '16',
         'NAICS code from SOS (Secretary of State) official filing',
         'Live Middesk API call per business\nNo Redshift pre-load',
         '2.0'],
        ['ZoomInfo', '24',
         'zi_c_naics6 — 6-digit NAICS 2022\nzi_c_sic — 4-digit SIC\nzi_c_industry — text label',
         'Redshift query on zoominfo_standard_ml_2\n(pre-loaded, not live API)',
         '0.8'],
        ['Equifax', '17',
         'efx_primnaicscode — 6-digit NAICS\nefx_primsic — 4-digit SIC\n'
         'efx_secnaics1..4 — secondary NAICS codes\nefx_naicdesc — text description',
         'Redshift query on equifax_us_standardized\n(pre-loaded, not live API)',
         '0.7'],
        ['Trulioo', '38',
         'sicCode — may be 4-digit (POLLUTED flag triggered if so)\n'
         'naicsCode (if present)\nindustryCode (fallback)',
         'Live Trulioo KYB API call per business\nNo Redshift pre-load',
         '0.8'],
        ['AI (GPT-5-mini)', '31',
         'naics_code (6-digit NAICS 2022)\nnaics_description\n'
         'uk_sic_code (if GB company)\nmcc_code\nmcc_description',
         'Live OpenAI API call (conditional)\n'
         'Stored in integration_data.request_response (platform_id=31)',
         '0.1'],
    ],
    col_widths=[1.4, 1.0, 2.8, 2.7, 0.6],
)

H2('AI Enrichment — Full Trigger Logic and Flow')
H3('When AI fires (from aiNaicsEnrichment.ts)')
code_block([
    '// Trigger conditions (all must be true):',
    'maximumSources: 3   -- fewer than 3 vendor sources returned a NAICS code',
    'minimumSources: 1   -- at least 1 source returned something (context for AI)',
    "ignoreSources: ['AINaicsEnrichment']  -- AI does not count itself",
    '',
    '// AI does NOT fire when:',
    '//  - 3+ vendor sources already have NAICS codes (high confidence, no need)',
    '//  - 0 sources returned anything (no context to provide to GPT)',
    '//  - manualOverride() already selected a winner',
])

H3('What AI receives as input (prompt context)')
data_table(
    ['Input field', 'Source', 'Purpose'],
    [
        ['business_name', 'Submitted application', 'Primary identifier'],
        ['address (street, city, state, ZIP, country)', 'Submitted application', 'Jurisdiction + locality context'],
        ['dba_names', 'Submitted application (if provided)', 'Alternative trading names'],
        ['existing_codes_from_sources',
         'Other vendor results already collected',
         'Correction context — AI uses these to validate its output'],
        ['website URL', 'Submitted or scraped', 'AI calls web_search tool to fetch and read the website'],
        ['website_summary', 'GPT web_search output', 'AI\'s own summary of what the business does'],
    ],
    col_widths=[2.2, 2.5, 4.8],
)

H3('What AI returns (structured JSON)')
code_block([
    '{',
    '  "naics_code":        "722511",',
    '  "naics_description": "Full-Service Restaurants",',
    '  "uk_sic_code":       "56101",   // only if country = GB',
    '  "uk_sic_description":"Licensed restaurants",',
    '  "mcc_code":          "5812",',
    '  "mcc_description":   "Eating Places, Restaurants",',
    '  "confidence":        "HIGH",    // HIGH / MED / LOW',
    '  "reasoning":         "Website confirms full-service dine-in restaurant...",',
    '  "tools_used":        ["web_search"],',
    '  "website_summary":   "Joe\'s Pizza serves New York-style pizza..."',
    '}',
])

H3('AI post-processing (executePostProcessing in aiNaicsEnrichment.ts)')
code_block([
    '// After AI returns naics_code:',
    'const naicsInfo = await internalGetNaicsCode(response.naics_code)',
    '',
    'if (!naicsInfo?.[0]?.naics_label) {',
    '  // Code not in core_naics_code table (hallucination or invalid code)',
    '  await removeNaicsCode(taskId, response)',
    '  response.naics_code = "561499"   // All Other Business Support Services',
    '  // Re-dispatch task-complete event for re-processing',
    '}',
    '',
    '// Confidence mapping:',
    'confidence = response.confidence === "HIGH" ? 0.70',
    '           : response.confidence === "MED"  ? 0.50',
    '           : 0.30  // LOW or missing',
])

H2('The 6 Winner-Selection Rules — Code-Level Detail')
H3('WEIGHT_THRESHOLD constant')
code_block([
    "// From integration-service/lib/facts/constants.ts:",
    'const WEIGHT_THRESHOLD = 0.05',
    '',
    '// Source weights (from lib/facts/sources.ts):',
    'const SOURCE_WEIGHTS = {',
    '  MIDDESK:         2.0,',
    '  OPENCORPORATES:  0.9,',
    '  ZOOMINFO:        0.8,',
    '  TRULIOO:         0.8,',
    '  EQUIFAX:         0.7,',
    '  AI_ENRICHMENT:   0.1,',
    '}',
    '',
    '// MAX_CONFIDENCE_INDEX (Trulioo heuristic denominator):',
    'const MAX_CONFIDENCE_INDEX = 55',
])
H3('Rule 1 — manualOverride()')
code_block([
    '// From lib/facts/rules.ts:',
    'function manualOverride(candidates: FactCandidate[]): FactCandidate | null {',
    '  const override = candidates.find(c => c.override != null)',
    '  return override ?? null  // null = no override set, continue to next rule',
    '}',
    '',
    '// Analyst sets override via:',
    'PATCH /facts/business/{id}/override/naics_code',
    '// Payload: { value: "722511", reason: "Analyst verified via SOS lookup" }',
    '// Stored in facts JSONB with: "override": { value, reason, setBy, setAt }',
])
H3('Rule 2 — factWithHighestConfidence()')
code_block([
    '// From lib/facts/rules.ts:',
    'function factWithHighestConfidence(candidates: FactCandidate[]): FactCandidate | null {',
    '  const sorted = candidates.sort((a, b) => b.confidence - a.confidence)',
    '  const [top, second] = sorted',
    '  if (!second) return top  // only one candidate -> wins immediately',
    '  if (top.confidence - second.confidence > WEIGHT_THRESHOLD) {',
    '    return top   // clear winner (gap > 0.05)',
    '  }',
    '  return null    // too close -> go to weightedFactSelector()',
    '}',
    '',
    '// confidence resolved as:',
    "// factConfidence = fact.confidence ?? fact.source?.confidence ?? 0.1",
])
H3('Rule 3 — weightedFactSelector() tie-break')
code_block([
    '// From lib/facts/rules.ts:',
    'function weightedFactSelector(candidates: FactCandidate[]): FactCandidate {',
    '  // Sort by (confidence * weight) descending',
    '  return candidates.sort((a, b) => {',
    '    const scoreA = a.confidence * SOURCE_WEIGHTS[a.source]',
    '    const scoreB = b.confidence * SOURCE_WEIGHTS[b.source]',
    '    return scoreB - scoreA',
    '  })[0]',
    '}',
    '',
    '// Effective priority when all confidences are close:',
    '// Middesk (conf*2.0) > OC (conf*0.9) > ZI=Trulioo (conf*0.8) > EFX (conf*0.7) > AI (conf*0.1)',
])

H2('Pipeline A — All Scenarios')
data_table(
    ['Situation', 'What gets stored in rds_warehouse_public.facts', 'What customer sees'],
    [
        ['3+ sources return NAICS codes and agree',
         'Highest-confidence code in facts (name="naics_code").\n'
         'All alternatives in alternatives[] array.\n'
         'data_businesses.naics_id updated via Kafka UPDATE_NAICS_CODE.',
         'Strong classification with high confidence.\n'
         'All source confidences visible in alternatives[].'],
        ['Sources disagree (different NAICS codes)',
         'Highest-confidence source wins.\n'
         'Disagreeing codes appear in alternatives[].',
         'Winner code + confidence.\n'
         'Alternatives show the disagreement — low consensus visible.'],
        ['Only 1 source returned NAICS (<3 sources)',
         'That source wins.\n'
         'AI enrichment also fires — adds context or confirms.\n'
         'AI result stored in request_response (platform_id=31).',
         'Code with possibly low confidence.\n'
         'AI result appears in alternatives if it ran.'],
        ['No source returned any NAICS (0 sources)',
         'AI enrichment fires.\n'
         'If valid: AI code stored.\n'
         'If invalid / hallucinated: 561499 stored via removeNaicsCode().',
         'AI-generated code with LOW confidence, OR\n'
         '561499 placeholder — signals "no reliable classification".'],
        ['Winning code NOT in core_naics_code table',
         'removeNaicsCode() called.\n'
         'Replaced with "561499" hardcoded fallback.',
         '561499 — analyst override recommended.'],
        ['All sources have very low confidence (e.g. 0.05)',
         'Highest available wins — NO minimum confidence cutoff.\n'
         'Low confidence IS stored and visible.',
         'Code shown with visibly low confidence.\n'
         'No automatic suppression — analyst should verify.'],
        ['Analyst set manual override before fact engine runs',
         'manualOverride() wins in Rule 1.\n'
         'Override stored with: override:{value, reason, setBy, setAt}.',
         'Override code with override field populated.\n'
         'Visible as analyst decision in API response.'],
        ['AI also fails / hallucinates AND no vendor code',
         '561499 stored as absolute last resort.',
         '561499 — requires analyst review and override.'],
    ],
    col_widths=[2.6, 3.4, 3.5],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TAB 5 — THE 8 CLASSIFICATION FACTS (FULL DETAIL)
# ════════════════════════════════════════════════════════════════════════════

H1('Tab 5 — The 8 Classification Facts (Full Detail)', colour=PURPLE, page_break=False)
callout(
    'Each of the 8 facts is stored in rds_warehouse_public.facts:\n'
    '  Schema: (id SERIAL, business_id UUID, name VARCHAR, value JSONB, received_at TIMESTAMPTZ)\n'
    '  Primary key: (business_id, name)  — one row per fact per business, always latest value.\n'
    '  Write pattern: INSERT ... ON CONFLICT (business_id, name) DO UPDATE SET value = EXCLUDED.value\n\n'
    'Some facts also denormalize into rds_cases_public.data_businesses as integer FK columns '
    '(naics_id, mcc_id, industry) pointing to lookup tables in the case-service PostgreSQL database.',
    bg='EDE9FE', border='5B21B6', text_colour=RGBColor(0x2E, 0x10, 0x65)
)

# ────────────────────────────────────────────────────────────────────────────
# FACT 1: naics_code
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 1: naics_code')
body_parts([
    ('[PIPELINE A — Fact Engine winner]  ', True, PURPLE),
    ('Also derived in Pipeline B as primary_naics_code (internal only)', False, SLATE),
])
body(
    'The single best 6-digit NAICS 2022 code for this business\'s primary economic activity. '
    'Selected by the Fact Engine from all source candidates using the 6-rule priority chain. '
    'Validated against core_naics_code lookup table — invalid codes replaced with 561499.'
)

H3('Winner Selection Priority Table')
data_table(
    ['Priority', 'Source', 'Redshift / API field', 'Weight', 'Confidence method',
     'Pipeline'],
    NAICS_PRIORITY_ROWS,
    col_widths=[0.6, 1.4, 2.5, 1.5, 2.7, 0.8],
    font_size=9,
)

H3('Trulioo POLLUTED flag detail')
callout(
    'Trulioo sometimes returns a 4-digit SIC code when a 6-digit NAICS is expected. '
    'The integration-service detects this: if len(str(sic_code)) == 4, the result is '
    'flagged as POLLUTED. A polluted Trulioo result still participates in selection '
    'but with its raw 4-digit code — the Fact Engine may still choose it if no better '
    'source is available.',
    bg='FFFBEB', border='D97706', text_colour=RGBColor(0x78, 0x35, 0x00)
)

H3('Post-processing / Validation (executePostProcessing in aiNaicsEnrichment.ts)')
code_block([
    '// After ANY winner is selected (not only AI — applies to all sources):',
    'const naicsInfo = await internalGetNaicsCode(winner.naics_code)',
    '// Looks up: SELECT * FROM core_naics_code WHERE code = winner.naics_code',
    '',
    'if (!naicsInfo?.[0]?.naics_label) {',
    '    // Code does not exist in core_naics_code table',
    '    // (data error, vendor returned invalid code, or AI hallucination)',
    '    await removeNaicsCode(taskId, response)',
    '    response.naics_code = "561499"  // All Other Business Support Services',
    '    // Re-dispatch PROCESS_COMPLETION_FACTS Kafka event',
    '}',
])

H3('JSONB Payload stored in rds_warehouse_public.facts (name="naics_code")')
code_block([
    '{',
    '  "code":        "541511",',
    '  "description": "Custom Computer Programming Services",',
    '  "source": {',
    '    "confidence": 0.95,',
    '    "platformId": 16          // 16 = Middesk',
    '  },',
    '  "override": null,           // set if analyst manually overrode',
    '  "alternatives": [',
    '    { "value": "541110", "source": 23, "confidence": 0.89 },  // OC',
    '    { "value": "541512", "source": 24, "confidence": 0.72 }   // ZoomInfo',
    '  ]',
    '}',
    '',
    '// When analyst override set:',
    '{',
    '  "code":     "541511",',
    '  "override": {',
    '    "value":  "722511",',
    '    "reason": "Analyst verified via SOS lookup",',
    '    "setBy":  "analyst@worth.ai",',
    '    "setAt":  "2026-03-15T10:22:00Z"',
    '  }',
    '}',
])

H3('SQL Queries — Retrieving naics_code')
H4('A. From facts table (JSONB, raw — always the latest value)')
code_block([
    '-- Full JSONB fact for one business:',
    'SELECT value, received_at',
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'naics_code';",
    '',
    "-- Extract just the code string from JSONB:",
    "SELECT value->>'code' AS naics_code,",
    "       value->'source'->>'confidence' AS confidence,",
    "       value->'source'->>'platformId' AS platform_id,",
    "       received_at",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'naics_code';",
])
H4('B. Via case-service denormalized table with decoded label')
code_block([
    'SELECT',
    '  b.id                AS business_id,',
    '  n.code              AS naics_code,',
    '  n.label             AS naics_description,',
    '  n.sector            AS naics_sector,      -- 2-digit sector code',
    '  n.subsector         AS naics_subsector     -- 3-digit subsector code',
    'FROM dev.rds_cases_public.data_businesses b',
    'JOIN dev.rds_cases_public.core_naics_code n',
    '  ON n.id = b.naics_id                    -- FK: naics_id -> core_naics_code.id',
    "WHERE b.id = '<uuid>';",
])
H4('C. All source alternatives for a business')
code_block([
    "-- Expand the alternatives[] array from JSONB:",
    'SELECT',
    "  jsonb_array_elements(value->'alternatives') AS alt",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'naics_code';",
    '',
    '-- With platform name decoded:',
    'SELECT',
    "  alt->>'value'      AS alternative_naics,",
    "  alt->>'confidence' AS alt_confidence,",
    '  CASE alt->>\'source\'',
    "    WHEN '16' THEN 'Middesk'",
    "    WHEN '17' THEN 'Equifax'",
    "    WHEN '23' THEN 'OpenCorporates'",
    "    WHEN '24' THEN 'ZoomInfo'",
    "    WHEN '31' THEN 'AI'",
    "    WHEN '38' THEN 'Trulioo'",
    "  END AS source_name",
    'FROM dev.rds_warehouse_public.facts,',
    "     jsonb_array_elements(value->'alternatives') AS alt",
    "WHERE business_id = '<uuid>'",
    "  AND name = 'naics_code';",
])
H4('D. Pipeline B — customer_files (batch analytics, ZI/EFX only)')
code_block([
    'SELECT',
    '  business_id,',
    '  primary_naics_code,      -- winner from ZI vs EFX rule',
    '  zi_match_confidence,     -- XGBoost zi_probability or heuristic fallback',
    '  efx_match_confidence,    -- XGBoost efx_probability or heuristic fallback',
    '  match_confidence         -- GREATEST(zi_match_confidence, efx_match_confidence)',
    'FROM datascience.customer_files',
    "WHERE business_id = '<uuid>';",
])

H3('Table Lineage for naics_code')
lineage_box([
    'INPUT SOURCES:',
    '  [OC]     open_corporates_standard_ml_2.industry_code_uids',
    '              -> parse entries matching prefix "us_naics-"',
    '  [Middesk] live SOS API response.naics_code',
    '  [ZI]     zoominfo_standard_ml_2.zi_c_naics6',
    '  [EFX]    equifax_us_standardized.efx_primnaicscode',
    '  [Trulioo] live KYB API response.sicCode (POLLUTED flag if 4-digit)',
    '  [AI]     OpenAI GPT-5-mini response.naics_code (conditional)',
    '       |',
    '       v',
    'INTERMEDIATE STORE (raw API responses):',
    '  integration_data.request_response',
    '    (business_id, platform_id, response JSONB, confidence)',
    '       |',
    '       v',
    'FACT ENGINE (integration-service/lib/facts/rules.ts):',
    '  manualOverride() -> factWithHighestConfidence() -> weightedFactSelector()',
    '  -> AI safety net -> 561499 validation fallback',
    '       |',
    '  Kafka: facts.v1  CALCULATE_BUSINESS_FACTS / UPDATE_NAICS_CODE',
    '       |',
    '       +----------> PRIMARY STORAGE:',
    '       |              rds_warehouse_public.facts',
    '       |                (business_id, name="naics_code", value=JSONB)',
    '       |',
    '       +----------> DENORMALISED FK STORAGE:',
    '       |              rds_cases_public.data_businesses.naics_id',
    '       |                -> FK to rds_cases_public.core_naics_code.id',
    '       |                   (id, code "541511", label, sector, subsector)',
    '       |',
    '       +----------> CUSTOMER API:',
    '                     GET /facts/business/{id}/details -> value JSONB',
    '                     GET /businesses/customers/{id}  -> naics_code flat field',
    '                     GET /facts/business/{id}/kyb    -> naics_code in KYB set',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 2: naics_description
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 2: naics_description')
body_parts([('[PIPELINE A — derived from naics_code winner]', True, PURPLE)])
body(
    'The human-readable English label for the winning NAICS code. '
    'Derived by looking up the winning naics_code in the core_naics_code table '
    'immediately after Fact 1 is resolved. '
    'Stored separately in facts for convenience so API clients do not need a '
    'secondary lookup.'
)

H3('How it is produced')
code_block([
    '// In integration-service, after naics_code winner is selected:',
    'const label = await internalGetNaicsCode(winner.naics_code)',
    '//   SELECT label FROM core_naics_code WHERE code = "541511"',
    '//   -> "Custom Computer Programming Services"',
    '',
    '// Published to Kafka facts.v1 as a separate fact alongside naics_code',
    '// Kafka event: CALCULATE_BUSINESS_FACTS',
])

H3('JSONB Payload (name="naics_description")')
code_block([
    '{',
    '  "description": "Custom Computer Programming Services"',
    '}',
])

H3('SQL Queries')
H4('A. From facts table')
code_block([
    'SELECT value->>\'description\' AS naics_description',
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'naics_description';",
])
H4('B. Or derive from naics_code + core_naics_code join')
code_block([
    '-- More flexible — also gives sector/subsector:',
    'SELECT n.code, n.label AS description, n.sector, n.subsector',
    'FROM dev.rds_cases_public.core_naics_code n',
    "WHERE n.code = '541511';",
])

H3('Table Lineage for naics_description')
lineage_box([
    'rds_cases_public.core_naics_code  (id, code, label, sector, subsector)',
    '   LOOKUP: core_naics_code WHERE code = winner.naics_code',
    '       |',
    '       v',
    'rds_warehouse_public.facts',
    '  (business_id, name="naics_description", value={"description":"..."})',
    '       |',
    '  NOTE: NOT stored in data_businesses — only in facts table.',
    '  Consumer: GET /facts/business/{id}/details  (alongside naics_code)',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 3: mcc_code
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 3: mcc_code')
body_parts([('[PIPELINE A — three resolution paths, priority order]', True, PURPLE)])
body(
    '4-digit Merchant Category Code (MCC) — the Visa/Mastercard standard used for '
    'interchange rate determination, fraud rules, and transaction risk scoring. '
    'There are three resolution paths, attempted in priority order.'
)

H3('Resolution Priority Table')
data_table(
    ['Path', 'Source', 'Field / Table', 'Weight', 'Confidence method', 'Pipeline'],
    MCC_PRIORITY_ROWS,
    col_widths=[1.3, 1.5, 2.5, 1.3, 2.0, 0.9],
    font_size=9,
)

H3('Path B — NAICS to MCC Crosswalk (rel_naics_mcc)')
code_block([
    '-- Case-service PostgreSQL crosswalk table:',
    'SELECT mcc_code, mcc_description',
    'FROM rds_cases_public.rel_naics_mcc',
    "WHERE naics_code = '722511';",
    '-- Returns: mcc_code=5812, mcc_description="Eating Places, Restaurants"',
    '',
    '-- The crosswalk is a many-to-one mapping:',
    '-- Multiple 6-digit NAICS codes can map to the same 4-digit MCC',
    '-- Example: 722511, 722513, 722514, 722515 all -> 5812',
    '',
    '-- Limitation: some businesses span multiple MCCs.',
    '-- The crosswalk always returns only the primary MCC for the NAICS code.',
])

H3('JSONB Payload (name="mcc_code")')
code_block([
    '{',
    '  "code":        "5812",',
    '  "description": "Eating Places, Restaurants",',
    '  "source": {',
    '    "confidence": 0.95,',
    '    "platformId": 16,',
    '    "path":       "direct"   // "direct" | "crosswalk" | "ai"',
    '  },',
    '  "override": null',
    '}',
])

H3('SQL Queries')
H4('A. From facts table')
code_block([
    "SELECT value->>'code' AS mcc_code,",
    "       value->>'description' AS mcc_description,",
    "       value->'source'->>'path' AS resolution_path",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'mcc_code';",
])
H4('B. Denormalized via data_businesses')
code_block([
    'SELECT b.id, m.code AS mcc_code, m.label AS mcc_description',
    'FROM dev.rds_cases_public.data_businesses b',
    'JOIN dev.rds_cases_public.core_mcc_code m',
    '  ON m.id = b.mcc_id',
    "WHERE b.id = '<uuid>';",
])
H4('C. Check crosswalk for a NAICS code')
code_block([
    'SELECT naics_code, mcc_code',
    'FROM rds_cases_public.rel_naics_mcc',
    "WHERE naics_code IN ('722511','722513','722514','722515');",
])

H3('Table Lineage for mcc_code')
lineage_box([
    'PATH A: Direct from vendor API',
    '  Middesk (platform_id=16) or Trulioo (platform_id=38)',
    '  -> mcc_code field in API response JSON',
    '  -> integration_data.request_response',
    '',
    'PATH B: NAICS->MCC crosswalk (most common)',
    '  rds_cases_public.core_naics_code (winning naics_code)',
    '  -> JOIN rds_cases_public.rel_naics_mcc ON naics_code',
    '  -> mcc_code derived deterministically',
    '',
    'PATH C: AI enrichment',
    '  OpenAI GPT-5-mini response.mcc_code (platform_id=31)',
    '  -> integration_data.request_response',
    '',
    'MERGE: Fact Engine selects best MCC via same priority rules',
    '       |',
    '       v',
    'rds_warehouse_public.facts  (name="mcc_code", value=JSONB)',
    'rds_cases_public.data_businesses.mcc_id  -> FK to core_mcc_code.id',
    '       |',
    '  Customer API: GET /facts/business/{id}/details -> mcc_code fact',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 4: mcc_description
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 4: mcc_description')
body_parts([('[PIPELINE A — derived from mcc_code]', True, PURPLE)])
body(
    'Human-readable English label for the MCC code. '
    'Derived by looking up the winning mcc_code in core_mcc_code table '
    'immediately after Fact 3 (mcc_code) is resolved. '
    'Stored separately in facts for convenience.'
)

H3('How it is produced')
code_block([
    '// After mcc_code winner selected:',
    'const mccInfo = await internalGetMccCode(winner.mcc_code)',
    '//   SELECT label FROM core_mcc_code WHERE code = "5812"',
    '//   -> "Eating Places, Restaurants"',
    '',
    '// Published to Kafka facts.v1 alongside mcc_code',
])

H3('JSONB Payload (name="mcc_description")')
code_block([
    '{',
    '  "description": "Eating Places, Restaurants"',
    '}',
])

H3('SQL Queries')
H4('A. From facts table')
code_block([
    "SELECT value->>'description' AS mcc_description",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'mcc_description';",
])
H4('B. From core_mcc_code lookup table')
code_block([
    'SELECT code, label AS description',
    'FROM rds_cases_public.core_mcc_code',
    "WHERE code = '5812';",
])

H3('Table Lineage for mcc_description')
lineage_box([
    'rds_cases_public.core_mcc_code  (id, code, label)',
    '   LOOKUP WHERE code = winner.mcc_code',
    '       |',
    '       v',
    'rds_warehouse_public.facts  (name="mcc_description", value={"description":"..."})',
    '  NOT stored in data_businesses — only in facts table.',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 5: mcc_code_found
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 5: mcc_code_found')
body_parts([('[PIPELINE A — boolean provenance flag]', True, PURPLE)])
body(
    'Boolean flag indicating HOW the MCC code was obtained. '
    'true = a vendor API directly returned an MCC code (higher quality, '
    'more precise because the vendor\'s classification system uses MCC natively). '
    'false = MCC was derived via the NAICS->MCC crosswalk table (approximation, '
    'since the crosswalk maps one NAICS to one MCC even if the business spans '
    'multiple MCC categories).'
)

H3('When mcc_code_found = true vs false')
data_table(
    ['Value', 'Condition', 'Source', 'Reliability'],
    [
        ['true',
         'Middesk or Trulioo API response contained an mcc_code field directly',
         'Live API (platform_id 16 or 38)',
         'High — vendor natively classifies by MCC'],
        ['true',
         'GPT-5-mini AI returned mcc_code in enrichment response',
         'AI enrichment (platform_id 31)',
         'Medium — AI inferred from business activity'],
        ['false',
         'MCC derived from NAICS->MCC crosswalk (rel_naics_mcc)',
         'Crosswalk lookup (no external source)',
         'Medium — approximation, one-to-one mapping only'],
        ['false / NULL',
         'No NAICS code available to crosswalk, and no vendor returned MCC',
         'n/a',
         'No MCC assigned'],
    ],
    col_widths=[0.8, 3.5, 2.3, 2.9],
)

H3('JSONB Payload (name="mcc_code_found")')
code_block([
    '{ "value": true }',
    '// or',
    '{ "value": false }',
])

H3('SQL Queries')
code_block([
    "SELECT value->>'value' AS mcc_code_found",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'mcc_code_found';",
    '',
    '-- Join with mcc_code to see both together:',
    'SELECT',
    "  f1.value->>'code'  AS mcc_code,",
    "  f2.value->>'value' AS mcc_found_directly",
    'FROM dev.rds_warehouse_public.facts f1',
    'JOIN dev.rds_warehouse_public.facts f2',
    '  ON f1.business_id = f2.business_id',
    "WHERE f1.business_id = '<uuid>'",
    "  AND f1.name = 'mcc_code'",
    "  AND f2.name = 'mcc_code_found';",
])

H3('Table Lineage for mcc_code_found')
lineage_box([
    'RESOLUTION PATH TRACKING:',
    '  Path A (direct API) -> mcc_code_found = true',
    '    Middesk (platform_id=16) -> API response.mcc_code exists',
    '    Trulioo (platform_id=38) -> API response.mcc_code exists',
    '  Path B (crosswalk)  -> mcc_code_found = false',
    '    rel_naics_mcc lookup derived mcc_code',
    '  Path C (AI)         -> mcc_code_found = true',
    '    AI returned mcc_code field',
    '       |',
    '       v',
    'rds_warehouse_public.facts  (name="mcc_code_found", value={"value":true|false})',
    '  NOT in data_businesses.',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 6: mcc_code_from_naics
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 6: mcc_code_from_naics')
body_parts([('[PIPELINE A — crosswalk-derived MCC, always computed]', True, PURPLE)])
body(
    'The MCC code derived specifically via the NAICS->MCC crosswalk table '
    '(rel_naics_mcc), regardless of whether a better direct MCC was found. '
    'Always computed when a naics_code winner is known. '
    'When mcc_code_found = false, mcc_code equals mcc_code_from_naics. '
    'When mcc_code_found = true, mcc_code_from_naics provides a secondary '
    'reference — useful for validation and crosswalk quality checks.'
)

H3('How it is produced')
code_block([
    '// Always computed after naics_code is resolved:',
    'const crosswalkMcc = await getMccFromNaics(winner.naics_code)',
    '//   SELECT mcc_code FROM rel_naics_mcc WHERE naics_code = "722511"',
    '//   -> "5812"',
    '',
    '// Published to Kafka as a separate fact regardless of other MCC facts.',
])

H3('JSONB Payload (name="mcc_code_from_naics")')
code_block([
    '{',
    '  "code":        "5812",',
    '  "naics_input": "722511",',
    '  "crosswalk":   "rel_naics_mcc"',
    '}',
])

H3('SQL Queries')
H4('A. From facts table')
code_block([
    "SELECT value->>'code' AS mcc_from_naics,",
    "       value->>'naics_input' AS source_naics",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'mcc_code_from_naics';",
])
H4('B. Compare direct MCC vs crosswalk-derived MCC')
code_block([
    'SELECT',
    "  f1.value->>'code'  AS mcc_direct,",
    "  f2.value->>'code'  AS mcc_from_naics,",
    "  CASE WHEN f1.value->>'code' = f2.value->>'code'",
    "    THEN 'AGREE' ELSE 'DISAGREE' END AS mcc_agreement",
    'FROM dev.rds_warehouse_public.facts f1',
    'JOIN dev.rds_warehouse_public.facts f2',
    '  ON f1.business_id = f2.business_id',
    "WHERE f1.business_id = '<uuid>'",
    "  AND f1.name = 'mcc_code'",
    "  AND f2.name = 'mcc_code_from_naics';",
])

H3('Table Lineage for mcc_code_from_naics')
lineage_box([
    'rds_cases_public.core_naics_code (winning naics_code from Fact 1)',
    '   JOIN rds_cases_public.rel_naics_mcc ON naics_code',
    '   -> mcc_code (deterministic crosswalk result)',
    '       |',
    '       v',
    'rds_warehouse_public.facts  (name="mcc_code_from_naics", value=JSONB)',
    '  NOT in data_businesses.',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 7: industry
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 7: industry')
body_parts([('[PIPELINE A — sector-level label for UI display]', True, PURPLE)])
body(
    'High-level sector category label, less granular than NAICS but more '
    'human-readable for customer-facing UI. Derived from the NAICS sector code '
    '(first 2 digits of the 6-digit NAICS code) via a lookup in '
    'core_business_industries. '
    'Also enriched by vendor text labels (ZoomInfo zi_c_industry, Trulioo industry '
    'text field, OC category labels) when available.'
)

H3('Derivation sources in priority order')
data_table(
    ['Source', 'Field / Method', 'Example', 'Notes'],
    [
        ['NAICS sector lookup',
         'First 2 digits of naics_code\n-> core_business_industries lookup',
         '722511 -> "72" -> "Food Service"',
         'Most reliable — derived from verified NAICS winner'],
        ['ZoomInfo',
         'zi_c_industry text field',
         '"Restaurant"',
         'Vendor free-text, less standardised'],
        ['Trulioo',
         'industry field in KYB response',
         '"FOOD_AND_BEVERAGE"',
         'Vendor-specific category system, may differ'],
        ['OpenCorporates',
         'Category labels from industry_code_uids taxonomy names',
         '"I — Accommodation and food service activities"',
         'Based on NACE/UK SIC sector labels'],
    ],
    col_widths=[2.0, 2.8, 2.0, 2.7],
)

H3('JSONB Payload (name="industry")')
code_block([
    '{',
    '  "value":  "Food Service",',
    '  "sector": "72",',
    '  "source": {',
    '    "confidence": 0.95,',
    '    "platformId": 16,',
    '    "derivedFrom": "naics_code"',
    '  }',
    '}',
])

H3('SQL Queries')
H4('A. From facts table')
code_block([
    "SELECT value->>'value' AS industry,",
    "       value->>'sector' AS naics_sector",
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'industry';",
])
H4('B. Via data_businesses FK')
code_block([
    'SELECT b.id, i.name AS industry, i.sector_code',
    'FROM dev.rds_cases_public.data_businesses b',
    'JOIN dev.rds_cases_public.core_business_industries i',
    '  ON i.id = b.industry',
    "WHERE b.id = '<uuid>';",
])
H4('C. Browse the core_business_industries lookup table')
code_block([
    'SELECT sector_code, name AS industry_label',
    'FROM rds_cases_public.core_business_industries',
    'ORDER BY sector_code;',
    '',
    '-- Example rows:',
    '-- 11 -> Agriculture, Forestry, Fishing and Hunting',
    '-- 22 -> Utilities',
    '-- 44 -> Retail Trade',
    '-- 52 -> Finance and Insurance',
    '-- 72 -> Food Service',
    '-- 81 -> Other Services (except Public Administration)',
])

H3('Table Lineage for industry')
lineage_box([
    'naics_code winner (Fact 1) -> first 2 digits = sector_code',
    '  -> rds_cases_public.core_business_industries',
    '     (id, sector_code "72", name "Food Service")',
    '',
    'ALSO enriched from:',
    '  ZoomInfo: zoominfo_standard_ml_2.zi_c_industry',
    '  Trulioo:  API response.industry',
    '  OC:       industry_code_uids taxonomy labels',
    '       |',
    '  Fact Engine merges and selects best label via same priority rules',
    '       |',
    '       v',
    'rds_warehouse_public.facts  (name="industry", value=JSONB)',
    'rds_cases_public.data_businesses.industry -> FK to core_business_industries.id',
    '       |',
    '  Customer API: GET /facts/business/{id}/details -> industry fact',
])

spacer(10)

# ────────────────────────────────────────────────────────────────────────────
# FACT 8: classification_codes
# ────────────────────────────────────────────────────────────────────────────

H2('Fact 8: classification_codes')
body_parts([('[PIPELINE A — raw multi-taxonomy snapshot — KNOWN GAP: downstream unused]',
             True, PURPLE)])
body(
    'A complete multi-taxonomy snapshot of ALL raw industry codes from ALL sources, '
    'captured BEFORE any winner selection. '
    'Includes NAICS, UK SIC, NACE, Canadian NAICS, and Equifax secondary codes. '
    'This is the richest classification fact — but it is currently an orphaned fact: '
    'no Kafka handler, API endpoint, PDF report, or database column reads it downstream. '
    'UK SIC and NACE codes for international companies exist in the system but are '
    'silently discarded.'
)
gap_box(
    'CONFIRMED GAP: classification_codes is written to rds_warehouse_public.facts '
    'correctly. It is never read by any downstream consumer. '
    'UK SIC code gb_sic-56101 (Licensed restaurants) exists for UK companies in OC\'s '
    'industry_code_uids but is discarded after being captured here. '
    'This is the primary gap blocking multi-taxonomy / international classification '
    'capabilities from reaching customers. '
    'Fixing it requires: (1) a Kafka handler for UPDATE_CLASSIFICATION_CODES event, '
    '(2) a new API endpoint or extending /facts/business/{id}/details, '
    '(3) a new data_businesses column or a separate classification_codes table.'
)

H3('All taxonomy systems captured')
data_table(
    ['Taxonomy', 'Code prefix in OC industry_code_uids', 'Example', 'Country coverage'],
    [
        ['US NAICS 2022',    'us_naics-',   'us_naics-541511',   'United States'],
        ['UK SIC 2007',      'gb_sic-',     'gb_sic-62012',      'United Kingdom'],
        ['NACE Rev 2',       'nace-',       'nace-J6201',        'European Union'],
        ['Canadian NAICS',   'ca_naics-',   'ca_naics-541511',   'Canada'],
        ['US SIC 1987',      'us_sic-',     'us_sic-7372',       'United States'],
        ['Equifax secondary','(direct field)','efx_secnaics1..4', 'United States only'],
        ['Trulioo raw SIC',  '(direct field)','sicCode from API', 'Multiple (may be POLLUTED)'],
    ],
    col_widths=[1.8, 2.3, 2.0, 3.4],
)

H3('JSONB Payload (name="classification_codes") — full example')
code_block([
    '{',
    '  "codes": [',
    '    {',
    '      "system":     "NAICS",',
    '      "code":       "722511",',
    '      "label":      "Full-Service Restaurants",',
    '      "source":     "Middesk",',
    '      "platformId": 16,',
    '      "confidence": 0.95,',
    '      "primary":    true',
    '    },',
    '    {',
    '      "system":     "NAICS",',
    '      "code":       "722511",',
    '      "source":     "OpenCorporates",',
    '      "platformId": 23,',
    '      "confidence": 0.89,',
    '      "taxonomy":   "us_naics",',
    '      "raw_entry":  "us_naics-722511"',
    '    },',
    '    {',
    '      "system":     "uk_sic",',
    '      "code":       "56101",',
    '      "label":      "Licensed restaurants",',
    '      "source":     "OpenCorporates",',
    '      "platformId": 23,',
    '      "confidence": 0.89,',
    '      "raw_entry":  "gb_sic-56101"',
    '    },',
    '    {',
    '      "system":     "nace",',
    '      "code":       "I5610",',
    '      "label":      "Restaurants and mobile food service activities",',
    '      "source":     "OpenCorporates",',
    '      "platformId": 23,',
    '      "confidence": 0.89,',
    '      "raw_entry":  "nace-I5610"',
    '    },',
    '    {',
    '      "system":     "NAICS",',
    '      "code":       "722513",',
    '      "source":     "ZoomInfo",',
    '      "platformId": 24,',
    '      "confidence": 0.72',
    '    },',
    '    {',
    '      "system":     "NAICS",',
    '      "code":       "722515",',
    '      "source":     "Equifax",',
    '      "platformId": 17,',
    '      "confidence": 0.68',
    '    },',
    '    {',
    '      "system":     "NAICS",',
    '      "code":       "722513",',
    '      "source":     "Equifax secondary",',
    '      "field":      "efx_secnaics2",',
    '      "confidence": 0.68',
    '    }',
    '  ]',
    '}',
])

H3('SQL Queries')
H4('A. Retrieve the full classification_codes fact')
code_block([
    'SELECT value, received_at',
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name = 'classification_codes';",
])
H4('B. Expand codes array — one row per taxonomy code')
code_block([
    'SELECT',
    "  code_entry->>'system'     AS taxonomy,",
    "  code_entry->>'code'       AS industry_code,",
    "  code_entry->>'label'      AS label,",
    "  code_entry->>'source'     AS source,",
    "  code_entry->>'confidence' AS confidence",
    'FROM dev.rds_warehouse_public.facts,',
    "     jsonb_array_elements(value->'codes') AS code_entry",
    "WHERE business_id = '<uuid>'",
    "  AND name = 'classification_codes'",
    "ORDER BY (code_entry->>'confidence')::float DESC;",
])
H4('C. Count businesses with UK SIC codes captured (currently orphaned)')
code_block([
    '-- How many businesses have a UK SIC code in classification_codes?',
    'SELECT COUNT(DISTINCT business_id) AS businesses_with_uk_sic',
    'FROM dev.rds_warehouse_public.facts,',
    "     jsonb_array_elements(value->'codes') AS code_entry",
    "WHERE name = 'classification_codes'",
    "  AND code_entry->>'system' = 'uk_sic';",
    '',
    '-- GAP: this count shows UK SIC data captured but never exposed to customers.',
])

H3('Table Lineage for classification_codes')
lineage_box([
    'INPUT: ALL raw vendor codes, pre-winner-selection',
    '',
    '  OpenCorporates:',
    '    open_corporates_standard_ml_2.industry_code_uids',
    '    -> parse ALL pipe-delimited entries:',
    '       us_naics-541511 | gb_sic-62012 | nace-J6201 | ca_naics-541511',
    '',
    '  ZoomInfo:',
    '    zoominfo_standard_ml_2.zi_c_naics6  (primary)',
    '',
    '  Equifax:',
    '    equifax_us_standardized.efx_primnaicscode  (primary)',
    '    equifax_us_standardized.efx_secnaics1      (secondary 1)',
    '    equifax_us_standardized.efx_secnaics2      (secondary 2)',
    '    equifax_us_standardized.efx_secnaics3      (secondary 3)',
    '    equifax_us_standardized.efx_secnaics4      (secondary 4)',
    '',
    '  Trulioo: live API response.sicCode',
    '  Middesk: live SOS API response.naics_code',
    '  AI:      GPT response.naics_code',
    '       |',
    '  Aggregated BEFORE Fact Engine winner selection',
    '       |',
    '       v',
    'rds_warehouse_public.facts  (name="classification_codes", value=JSONB array)',
    '',
    '  DOWNSTREAM: !! ORPHANED — no consumers !!',
    '  NOT in data_businesses.',
    '  NOT exposed via any API endpoint.',
    '  UK SIC / NACE / ISIC codes captured here are DISCARDED.',
    '',
    '  REQUIRED to fix:',
    '    1. Kafka handler for UPDATE_CLASSIFICATION_CODES',
    '    2. New API endpoint or extend /facts/business/{id}/details',
    '    3. New DB column or separate classification_codes table in case-service',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TAB 6 — API & STORAGE
# ════════════════════════════════════════════════════════════════════════════

H1('Tab 6 — API & Storage Reference', colour=PURPLE, page_break=False)

H2('All Database Tables Referenced in This Document')
data_table(
    ['Table', 'Database / Schema', 'Type', 'Purpose', 'Written by'],
    [
        ['rds_warehouse_public.facts',
         'PostgreSQL (warehouse-service) + Redshift mirror',
         'Key-value store',
         'All 217 facts as JSONB. One row per (business_id, name). Always latest value.',
         'warehouse-service FactService (Kafka consumer)'],
        ['rds_cases_public.data_businesses',
         'PostgreSQL (case-service)',
         'Denormalized business row',
         'naics_id, mcc_id, industry — integer FKs to lookup tables.',
         'case-service Kafka handler (UPDATE_NAICS_CODE)'],
        ['rds_cases_public.core_naics_code',
         'PostgreSQL (case-service)',
         'Lookup / reference',
         'id, code (6-digit), label, sector (2-digit), subsector (3-digit).',
         'DB migrations — static reference data'],
        ['rds_cases_public.core_mcc_code',
         'PostgreSQL (case-service)',
         'Lookup / reference',
         'id, code (4-digit), label.',
         'DB migrations — static reference data'],
        ['rds_cases_public.rel_naics_mcc',
         'PostgreSQL (case-service)',
         'Crosswalk / mapping',
         'naics_code -> mcc_code mapping table (many NAICS to one MCC).',
         'DB migrations — static reference data'],
        ['rds_cases_public.core_business_industries',
         'PostgreSQL (case-service)',
         'Lookup / reference',
         "sector_code (2-digit), name (e.g. '72' -> 'Food Service').",
         'DB migrations — static reference data'],
        ['integration_data.request_response',
         'PostgreSQL (integration-service)',
         'Audit / raw store',
         'All raw vendor API responses + confidence, per business per platform.',
         'integration-service — one row per API call'],
        ['datascience.customer_files',
         'Redshift',
         'Wide analytics table',
         'Pipeline B output. primary_naics_code, zi/efx_match_confidence, all firmographic.',
         'warehouse-service sp_recreate_customer_files()'],
        ['datascience.ml_model_matches',
         'Redshift',
         'XGBoost output',
         'zi_probability, efx_probability, oc_probability per business.',
         'XGBoost entity_matching_20250127 v1 batch job'],
        ['datascience.smb_zi_oc_efx_combined',
         'Redshift',
         'Intermediate combined',
         'zi_match_confidence, efx_match_confidence (XGBoost + heuristic merged).',
         'smb_zi_oc_efx_combined.sql stored procedure'],
        ['datascience.zoominfo_standard_ml_2',
         'Redshift',
         'Standardised vendor',
         'ZoomInfo standardised — zi_c_naics6, names, addresses.',
         'Bulk file ingestion ETL'],
        ['warehouse.equifax_us_standardized',
         'Redshift',
         'Standardised vendor',
         'Equifax standardised — efx_primnaicscode, efx_primsic, names, addresses.',
         'Bulk file ingestion ETL'],
        ['datascience.open_corporates_standard_ml_2',
         'Redshift',
         'Standardised vendor',
         'OC standardised — industry_code_uids (multi-taxonomy pipe-delimited).',
         'Bulk file ingestion ETL'],
    ],
    col_widths=[2.5, 2.0, 1.3, 2.9, 0.8],
    font_size=8.5,
)

H2('API Endpoints That Expose Classification Facts')
data_table(
    ['Endpoint', 'Method', 'Returns naics_code?', 'Access level', 'What is returned'],
    [
        ['GET /facts/business/{id}/details',
         'GET', 'YES — primary endpoint',
         'Admin / Customer / Applicant',
         'Full JSONB fact: code, description, source.confidence, source.platformId, alternatives[]'],
        ['GET /facts/business/{id}/kyb',
         'GET', 'YES — via KYB fact set',
         'Admin / Customer / Applicant',
         'naics_code + mcc_code + industry as part of KYB verification package'],
        ['GET /businesses/customers/{id}',
         'GET', 'YES — simplified flat fields',
         'Admin / Customer',
         'Flat response: naics_code + naics_title strings (no source lineage, no alternatives)'],
        ['GET /facts/business/{id}/all',
         'GET', 'YES — all 217 facts',
         'Admin ONLY',
         "All 217 facts including naics_code. Note in code: 'intentionally not cached, leaks information'"],
        ['PATCH /businesses/{id}',
         'PATCH', 'NO — stripped from body',
         'n/a',
         'naics_code is removed from the request body — cannot be set via PATCH. Use override endpoint.'],
        ['PATCH /facts/business/{id}/override/naics_code',
         'PATCH', 'YES — write path',
         'Admin / Customer',
         'Sets analyst override. Payload: {value, reason}. Override always wins future Fact Engine runs.'],
        ['GET /facts/business/{id}/classification-codes',
         'GET', 'GAP: not implemented',
         'n/a',
         'This endpoint does NOT exist. classification_codes fact is never exposed via API.'],
    ],
    col_widths=[3.0, 0.7, 1.5, 1.4, 2.9],
    font_size=8.5,
)

H2('Kafka Events for Classification Facts')
data_table(
    ['Event', 'Topic', 'Triggered by', 'Consumers', 'Effect'],
    [
        ['CALCULATE_BUSINESS_FACTS', 'facts.v1',
         'Business submitted', 'warehouse-service FactService',
         'Initial fact calculation for all 217 facts including naics_code, mcc_code, industry'],
        ['UPDATE_NAICS_CODE', 'facts.v1',
         'naics_code fact changes', 'case-service Kafka handler',
         'UPDATE data_businesses SET naics_id = (SELECT id FROM core_naics_code WHERE code=winner)'],
        ['UPDATE_MCC_CODE', 'facts.v1',
         'mcc_code fact changes', 'case-service Kafka handler',
         'UPDATE data_businesses SET mcc_id = (SELECT id FROM core_mcc_code WHERE code=winner)'],
        ['UPDATE_INDUSTRY', 'facts.v1',
         'industry fact changes', 'case-service Kafka handler',
         'UPDATE data_businesses SET industry = (SELECT id FROM core_business_industries WHERE ...)'],
        ['FACT_OVERRIDE_CREATED_AUDIT', 'facts.v1',
         'Analyst sets override via PATCH', 'warehouse-service + audit service',
         'Override stored in facts JSONB with audit fields. override: {value, reason, setBy, setAt}'],
        ['PROCESS_COMPLETION_FACTS', 'facts.v1',
         'All verifications complete', 'warehouse-service FactService',
         'Final fact recalculation with full context — may change winner after Trulioo/Middesk complete'],
        ['UPDATE_CLASSIFICATION_CODES', 'facts.v1',
         'classification_codes captured', 'GAP: no consumer',
         'GAP: this event is never subscribed to. classification_codes is captured but discarded.'],
    ],
    col_widths=[2.4, 0.9, 1.8, 1.9, 2.5],
    font_size=8.5,
)

H2('Master Redshift Query — All 8 Classification Facts for One Business')
code_block([
    '-- Retrieve all 8 classification facts in one query:',
    'SELECT',
    '  name,',
    "  value->>'code'        AS code,",
    "  value->>'description' AS description,",
    "  value->>'value'       AS boolean_value,",
    "  value->'source'->>'confidence'  AS confidence,",
    "  value->'source'->>'platformId'  AS platform_id,",
    '  received_at',
    'FROM dev.rds_warehouse_public.facts',
    "WHERE business_id = '<uuid>'",
    "  AND name IN (",
    "    'naics_code',",
    "    'naics_description',",
    "    'mcc_code',",
    "    'mcc_description',",
    "    'mcc_code_found',",
    "    'mcc_code_from_naics',",
    "    'industry',",
    "    'classification_codes'",
    '  )',
    'ORDER BY name;',
])

H2('Full Denormalized Join — All Facts Decoded from case-service Lookup Tables')
code_block([
    '-- Complete business classification with all decoded labels:',
    'SELECT',
    '  b.id                                        AS business_id,',
    '  n.code                                      AS naics_code,',
    '  n.label                                     AS naics_description,',
    '  n.sector                                    AS naics_sector_2digit,',
    '  n.subsector                                 AS naics_subsector_3digit,',
    '  m.code                                      AS mcc_code,',
    '  m.label                                     AS mcc_description,',
    '  r.mcc_code                                  AS mcc_code_from_naics,',
    '  i.name                                      AS industry_sector_label,',
    '  i.sector_code                               AS industry_sector_code',
    'FROM dev.rds_cases_public.data_businesses b',
    'LEFT JOIN dev.rds_cases_public.core_naics_code n',
    '       ON n.id = b.naics_id',
    'LEFT JOIN dev.rds_cases_public.core_mcc_code m',
    '       ON m.id = b.mcc_id',
    'LEFT JOIN dev.rds_cases_public.rel_naics_mcc r',
    '       ON r.naics_code = n.code',
    'LEFT JOIN dev.rds_cases_public.core_business_industries i',
    '       ON i.id = b.industry',
    "WHERE b.id = '<uuid>';",
])

H2('Pipeline A vs Pipeline B — NAICS Code Agreement Check')
code_block([
    '-- Find businesses where Pipeline A and Pipeline B disagree on NAICS:',
    'SELECT',
    '  f.business_id,',
    "  f.value->>'code'              AS pipeline_a_naics,   -- Fact Engine winner",
    "  f.value->'source'->>'platformId' AS pipeline_a_source,",
    '  cf.primary_naics_code::text   AS pipeline_b_naics,   -- ZI/EFX winner',
    '  cf.zi_match_confidence,',
    '  cf.efx_match_confidence',
    'FROM dev.rds_warehouse_public.facts f',
    'JOIN datascience.customer_files cf',
    '  ON cf.business_id = f.business_id',
    "WHERE f.name = 'naics_code'",
    "  AND f.value->>'code' != cf.primary_naics_code::text",
    'ORDER BY cf.match_confidence DESC',
    'LIMIT 500;',
    '',
    '-- To quantify the disagreement rate:',
    'SELECT',
    '  COUNT(*) AS total_businesses,',
    "  SUM(CASE WHEN f.value->>'code' = cf.primary_naics_code::text THEN 1 ELSE 0 END) AS agree,",
    "  SUM(CASE WHEN f.value->>'code' != cf.primary_naics_code::text THEN 1 ELSE 0 END) AS disagree,",
    "  ROUND(100.0 * SUM(CASE WHEN f.value->>'code' != cf.primary_naics_code::text",
    '                         THEN 1 ELSE 0 END) / COUNT(*), 2) AS disagree_pct',
    'FROM dev.rds_warehouse_public.facts f',
    'JOIN datascience.customer_files cf ON cf.business_id = f.business_id',
    "WHERE f.name = 'naics_code';",
])

# ════════════════════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════════════════════
out = ('/workspace/AI-Powered-NAICS-Industry-Classification-Agent/'
       'modeling/Industry_Classification_Worthopedia.docx')
doc.save(out)
size_kb = round(os.path.getsize(out) / 1024, 1)
print(f'Saved  : {out}')
print(f'Size   : {size_kb} KB')
print('Ready  : Upload to Google Docs via File > Import or drag-and-drop')
